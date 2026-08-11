from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "Nguyễn Đức Thịnh_2305CT2084_CT07PM"
OUTPUT = ROOT / "docs" / "Bao_cao_do_an_HomeX_POS.docx"
LOGO = ROOT / "docs" / "images" / "logo_hvu.png"
ADVISOR = "ThS. Nguyễn Thanh Tiến"
TITLE = "XÂY DỰNG VÀ TRIỂN KHAI WEBSITE POS CHO CỬA HÀNG ĐỒ GIA DỤNG HOMEX"
TITLE_DISPLAY = "Xây dựng và triển khai website POS cho cửa hàng đồ gia dụng HomeX"


def field(paragraph, instruction: str, result: str = ""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = result
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))
    return run


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.left_margin, section.right_margin = Cm(3), Cm(2)
    section.top_margin = section.bottom_margin = Cm(2.5)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header.paragraphs[0].text = ""
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_before = normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(1)
    for level in range(1, 5):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.size = Pt(14 if level == 1 else 13)
        style.font.bold = level <= 3
        style.font.italic = level == 3
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(6), Pt(3)
        style.paragraph_format.keep_with_next = True
    front_style = doc.styles.add_style("Front Matter Title", WD_STYLE_TYPE.PARAGRAPH) if "Front Matter Title" not in doc.styles else doc.styles["Front Matter Title"]
    front_style.base_style = normal
    front_style.font.name = "Times New Roman"
    front_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    front_style.font.size, front_style.font.bold = Pt(14), True
    front_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front_style.paragraph_format.first_line_indent = Cm(0)
    front_style.paragraph_format.space_before, front_style.paragraph_format.space_after = Pt(6), Pt(3)
    front_style.paragraph_format.keep_with_next = True
    front_ppr = front_style._element.get_or_add_pPr()
    front_outline = front_ppr.find(qn("w:outlineLvl"))
    if front_outline is None:
        front_outline = OxmlElement("w:outlineLvl")
        front_ppr.append(front_outline)
    front_outline.set(qn("w:val"), "9")

    for name, alignment in (("Figure Caption", WD_ALIGN_PARAGRAPH.CENTER), ("Table Caption", WD_ALIGN_PARAGRAPH.LEFT)):
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH) if name not in doc.styles else doc.styles[name]
        style.base_style = normal
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size, style.font.bold = Pt(13), True
        style.paragraph_format.alignment = alignment
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(6 if name == "Figure Caption" else 0)
        style.paragraph_format.space_after = Pt(0 if name == "Figure Caption" else 6)
        style.paragraph_format.keep_with_next = name == "Table Caption"
        caption_ppr = style._element.get_or_add_pPr()
        caption_outline = caption_ppr.find(qn("w:outlineLvl"))
        if caption_outline is None:
            caption_outline = OxmlElement("w:outlineLvl")
            caption_ppr.append(caption_outline)
        caption_outline.set(qn("w:val"), "9")

    list_style = doc.styles.add_style("Table of Figures", WD_STYLE_TYPE.PARAGRAPH) if "Table of Figures" not in doc.styles else doc.styles["Table of Figures"]
    list_style.base_style = normal
    list_style.font.name = "Times New Roman"
    list_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    list_style.font.size = Pt(12)
    list_style.paragraph_format.first_line_indent = Cm(0)
    list_style.paragraph_format.space_before = Pt(0)
    list_style.paragraph_format.space_after = Pt(0)
    list_style.paragraph_format.line_spacing = 1.0
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    doc.settings._element.append(update)


def heading(doc: Document, text: str, level: int, new_page: bool = False):
    paragraph = doc.add_heading(text, level=level)
    if new_page:
        paragraph.paragraph_format.page_break_before = True
    return paragraph


def front_heading(doc: Document, text: str, new_page: bool = True):
    paragraph = doc.add_paragraph(text, style="Front Matter Title")
    paragraph.paragraph_format.page_break_before = new_page
    return paragraph


def body(doc: Document, text: str, indent: bool = True, italic: bool = False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(1 if indent else 0)
    run = paragraph.add_run(text)
    run.font.name, run.font.size, run.font.italic = "Times New Roman", Pt(13), italic
    return paragraph


def bullets(doc: Document, items: Iterable[str], numbered: bool = False) -> None:
    for index, item in enumerate(items, 1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.8)
        paragraph.paragraph_format.first_line_indent = Cm(-0.5)
        paragraph.add_run((f"{index}. " if numbered else "- ") + item)


_bookmark_id = 0


def caption_number(paragraph, kind: str, number: str) -> str:
    global _bookmark_id
    prefix, ordinal = number.rsplit(".", 1)
    sequence = "Bang" if kind == "Bảng" else "Hinh"
    bookmark_name = f"{sequence.lower()}_{prefix}_{ordinal}".replace("-", "_")
    _bookmark_id += 1
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(_bookmark_id))
    bookmark_start.set(qn("w:name"), bookmark_name)
    paragraph._p.append(bookmark_start)
    paragraph.add_run(f"{kind} {prefix}.")
    instruction = f"SEQ {sequence} \\* ARABIC"
    if ordinal == "1":
        instruction += " \\r 1"
    field(paragraph, instruction, ordinal)
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(_bookmark_id))
    paragraph._p.append(bookmark_end)
    return bookmark_name


def body_with_reference(doc: Document, before: str, bookmark: str, display: str, after: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(1)
    paragraph.add_run(before)
    field(paragraph, f"REF {bookmark} \\h", display)
    paragraph.add_run(after)
    return paragraph


def table(doc: Document, number: str | None, title: str | None, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float] | None = None, font_size: float = 11) -> None:
    if number is not None and title is not None:
        caption = doc.add_paragraph(style="Table Caption")
        caption_number(caption, "Bảng", number)
        caption.add_run(f". {title}")
    elif number is not None or title is not None:
        raise ValueError("number and title must both be provided or both be None")
    obj = doc.add_table(rows=1, cols=len(headers))
    obj.style, obj.alignment, obj.autofit = "Table Grid", WD_TABLE_ALIGNMENT.CENTER, False
    for index, value in enumerate(headers):
        cell = obj.rows[0].cells[index]
        cell.text = value
        if widths:
            cell.width = Cm(widths[index])
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "D9E2F3")
        cell._tc.get_or_add_tcPr().append(shade)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            for run in paragraph.runs:
                run.font.bold, run.font.size = True, Pt(font_size)
    header_props = obj.rows[0]._tr.get_or_add_trPr()
    header_repeat = OxmlElement("w:tblHeader")
    header_repeat.set(qn("w:val"), "true")
    header_props.append(header_repeat)
    no_split_header = OxmlElement("w:cantSplit")
    header_props.append(no_split_header)
    for data in rows:
        row = obj.add_row()
        no_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(no_split)
        cells = row.cells
        for index, value in enumerate(data):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cells[index].width = Cm(widths[index])
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    run.font.name, run.font.size = "Times New Roman", Pt(font_size)



def picture(doc: Document, path: str, number: str, title: str, max_width: float = 15.7, max_height: float = 17.5) -> None:
    source = ASSETS / path
    if not source.exists():
        body(doc, f"[Không tìm thấy hình {source.name}]", False, True)
        return
    with Image.open(source) as image:
        ratio = image.width / image.height
    width, height = max_width, max_width / ratio
    if height > max_height:
        height, width = max_height, max_height * ratio
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(source), width=Cm(width), height=Cm(height))
    caption = doc.add_paragraph(style="Figure Caption")
    caption_number(caption, "Hình", number)
    caption.add_run(f". {title}")


def centered(doc: Document, text: str, size: float = 13, bold: bool = False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(text)
    run.font.name, run.font.size, run.font.bold = "Times New Roman", Pt(size), bold
    return paragraph


def add_page_number(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    paragraph = section.header.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field(paragraph, "PAGE")
    section.header_distance = Cm(1)
    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:start"), "1")
    section._sectPr.append(pg_num)


def cover(doc: Document) -> None:
    centered(doc, "BỘ GIÁO DỤC VÀ ĐÀO TẠO", 14, True)
    centered(doc, "TRƯỜNG ĐẠI HỌC HÙNG VƯƠNG TP. HỒ CHÍ MINH", 14, True)
    logo_paragraph = centered(doc, "---------------------------\n", 14)
    if LOGO.exists():
        logo_paragraph.add_run().add_picture(str(LOGO), width=Cm(5.9478), height=Cm(5.9478))
    centered(doc, "THỰC TẬP NGHỀ NGHIỆP 1\n", 16, True)
    centered(doc, TITLE + "\n\n", 23, True)
    info = doc.add_table(rows=4, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ("Người hướng dẫn:", "Người thực hiện:", "Mã số sinh viên:", "Lớp:")
    values = (ADVISOR, "NGUYỄN ĐỨC THỊNH", "2305CT2084", "CT07PM")
    for row, (label, value) in enumerate(zip(labels, values)):
        info.cell(row, 0).text, info.cell(row, 1).text = label, value
        for cell in info.rows[row].cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                for run in paragraph.runs:
                    run.font.name, run.font.size, run.font.bold = "Times New Roman", Pt(14), True
        info.cell(row, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    info._tbl.tblPr.append(borders)
    centered(doc, "\nTHÀNH PHỐ HỒ CHÍ MINH, NĂM 2026", 14, True)


def front_matter(doc: Document) -> None:
    front_heading(doc, "LỜI CẢM ƠN")
    body(doc, "Để hoàn thành báo cáo Thực tập nghề nghiệp 1, em xin chân thành cảm ơn Ban Giám hiệu, Khoa Kỹ thuật - Công nghệ và quý thầy cô Trường Đại học Hùng Vương Thành phố Hồ Chí Minh đã truyền đạt kiến thức, tạo điều kiện học tập và hỗ trợ em trong quá trình thực hiện đề tài.")
    body(doc, f"Đặc biệt, em xin gửi lời cảm ơn sâu sắc đến giảng viên hướng dẫn {ADVISOR}. Những góp ý về phương pháp phân tích, thiết kế hệ thống và cách trình bày khoa học là cơ sở để em hoàn thiện đề tài “{TITLE_DISPLAY}”.")
    body(doc, "Mặc dù đã nỗ lực đối chiếu giữa yêu cầu, mã nguồn và kết quả triển khai, báo cáo khó tránh khỏi thiếu sót. Em kính mong nhận được ý kiến góp ý của quý thầy cô để hệ thống tiếp tục được hoàn thiện.")
    front_heading(doc, "LỜI CAM ĐOAN")
    body(doc, f"Em xin cam đoan đề tài “{TITLE_DISPLAY}” là kết quả nghiên cứu, phân tích và phát triển của cá nhân em dưới sự hướng dẫn của {ADVISOR}. Nội dung báo cáo phản ánh mã nguồn và kết quả triển khai của dự án HomeX POS; các tài liệu, công nghệ và ý tưởng tham khảo đều được trích dẫn.")
    body(doc, "Các số liệu minh họa, sơ đồ và hình ảnh giao diện được sử dụng trung thực nhằm mô tả hệ thống. Em chịu trách nhiệm về nội dung báo cáo và các vấn đề liên quan đến quyền tác giả trong phạm vi công việc của mình.")
    sign = doc.add_paragraph("\nTP. Hồ Chí Minh, ngày 11 tháng 08 năm 2026\nNgười cam đoan\n\n\nNGUYỄN ĐỨC THỊNH")
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign.paragraph_format.first_line_indent = Cm(0)
    front_heading(doc, "MỤC LỤC")
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    field(paragraph, 'TOC \\o "1-3" \\h \\z', "Đang cập nhật mục lục...")
    front_heading(doc, "DANH MỤC CÁC TỪ VIẾT TẮT")
    table(doc, None, None, ("Viết tắt", "Tiếng Anh", "Tiếng Việt"), (
        ("AI", "Artificial Intelligence", "Trí tuệ nhân tạo"), ("API", "Application Programming Interface", "Giao diện lập trình ứng dụng"),
        ("CORS", "Cross-Origin Resource Sharing", "Chia sẻ tài nguyên khác nguồn"),
        ("CSDL", "Database", "Cơ sở dữ liệu"), ("ERD", "Entity Relationship Diagram", "Sơ đồ thực thể - liên kết"),
        ("JWT", "JSON Web Token", "Mã thông báo web dạng JSON"), ("ORM", "Object-Relational Mapping", "Ánh xạ đối tượng - quan hệ"),
        ("POS", "Point of Sale", "Hệ thống điểm bán hàng"), ("QR", "Quick Response", "Mã phản hồi nhanh"),
        ("REST", "Representational State Transfer", "Kiểu kiến trúc dịch vụ web"), ("SKU", "Stock Keeping Unit", "Mã đơn vị lưu kho"),
        ("UI/UX", "User Interface/User Experience", "Giao diện/trải nghiệm người dùng"), ("UML", "Unified Modeling Language", "Ngôn ngữ mô hình hóa thống nhất"),
        ("HMAC", "Hash-based Message Authentication Code", "Mã xác thực thông điệp bằng hàm băm"),
        ("ID", "Identifier", "Mã định danh"), ("JSON", "JavaScript Object Notation", "Định dạng dữ liệu JSON"),
        ("LLM", "Large Language Model", "Mô hình ngôn ngữ lớn"),
        ("MAE", "Mean Absolute Error", "Sai số tuyệt đối trung bình"), ("MAPE", "Mean Absolute Percentage Error", "Sai số phần trăm tuyệt đối trung bình"),
        ("NCC", "Supplier", "Nhà cung cấp"), ("PO", "Purchase Order", "Phiếu nhập hàng"),
        ("RMSE", "Root Mean Squared Error", "Căn sai số bình phương trung bình"), ("SDK", "Software Development Kit", "Bộ công cụ phát triển phần mềm"),
        ("SMTP", "Simple Mail Transfer Protocol", "Giao thức gửi thư điện tử"),
        ("URL", "Uniform Resource Locator", "Địa chỉ tài nguyên"), ("VAT", "Value Added Tax", "Thuế giá trị gia tăng"),
        ("WAPE", "Weighted Absolute Percentage Error", "Sai số phần trăm tuyệt đối có trọng số"),
    ), (2.1, 6.2, 7.1), 11)
    for title, sequence in (("DANH MỤC BẢNG", "Bang"), ("DANH MỤC HÌNH ẢNH", "Hinh")):
        front_heading(doc, title)
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0)
        field(paragraph, f'TOC \\h \\z \\c "{sequence}"', "Đang cập nhật danh mục...")

def opening(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.left_margin, section.right_margin = Cm(3), Cm(2)
    section.top_margin = section.bottom_margin = Cm(2.5)
    add_page_number(section)
    heading(doc, "MỞ ĐẦU", 1)
    topics = (
        ("Lý do chọn đề tài", (
            "Chuyển đổi số trong bán lẻ không chỉ thay thế sổ sách bằng phần mềm mà còn yêu cầu liên kết chặt chẽ giữa bán hàng, thanh toán, tồn kho, khách hàng và báo cáo quản trị. Cửa hàng đồ gia dụng có danh mục hàng hóa rộng, mức giá và thời hạn bảo hành khác nhau; vì vậy sai lệch ở một giao dịch có thể ảnh hưởng đồng thời đến tiền quỹ, số lượng tồn, quyền lợi bảo hành và dữ liệu doanh thu.",
            "Trong quy trình thủ công, nhân viên phải tra cứu giá, ghi hóa đơn, tính tiền và cập nhật kho ở nhiều nơi. Cách làm này dễ phát sinh nhập sai mã hàng, bán vượt tồn, bỏ sót khuyến mãi, khó truy vết người thao tác và mất nhiều thời gian tổng hợp báo cáo. Từ nhu cầu đó, đề tài xây dựng HomeX POS như một hệ thống web dùng tại quầy, tập trung dữ liệu và tự động hóa luồng nghiệp vụ từ khi chọn hàng đến khi hoàn tất hóa đơn.",
        )),
        ("Lịch sử và định hướng nghiên cứu", (
            "Các hệ thống POS hiện đại phát triển từ chức năng ghi nhận giao dịch thành nền tảng quản trị bán lẻ tích hợp. Trong phạm vi học phần, tài liệu giao đề tài yêu cầu sản phẩm không dừng ở CRUD mà phải thể hiện đầy đủ logic tạo đơn, thanh toán, biến động tồn kho, báo cáo, phân quyền và ít nhất một nội dung nâng cao.",
            "HomeX POS được định hướng theo kiến trúc client-server, giao tiếp REST API, cơ sở dữ liệu quan hệ và giao diện đáp ứng. Điểm tập trung là mô hình hóa quy trình bán hàng nhất quán, bảo toàn dữ liệu qua giao dịch CSDL và tạo trải nghiệm thao tác nhanh cho hai vai trò Admin và Cashier.",
        )),
        ("Đối tượng và phạm vi nghiên cứu", (
            "Đối tượng nghiên cứu gồm quy trình bán hàng tại quầy, quản lý hàng hóa và các kỹ thuật xây dựng hệ thống web POS. Người dùng trực tiếp là quản trị viên/chủ cửa hàng và thu ngân; khách hàng tương tác gián tiếp qua thanh toán và trang tra cứu bảo hành công khai.",
            "Phạm vi triển khai là một cửa hàng đồ gia dụng quy mô nhỏ hoặc vừa, một kho và hai vai trò nội bộ. Hệ thống chưa giải quyết chuỗi chi nhánh, thương mại điện tử đầy đủ, vận chuyển phức tạp hoặc chế độ offline hoàn toàn. AI đóng vai trò hỗ trợ quyết định; người dùng vẫn là chủ thể phê duyệt đề xuất.",
        )),
    )
    for title, paragraphs in topics:
        heading(doc, title, 2)
        for paragraph in paragraphs:
            body(doc, paragraph)
    heading(doc, "Mục tiêu nghiên cứu", 2)
    bullets(doc, (
        "Phân tích tác nhân, quy tắc nghiệp vụ, yêu cầu chức năng và yêu cầu phi chức năng của cửa hàng đồ gia dụng.",
        "Thiết kế kiến trúc frontend-backend, mô hình PostgreSQL và bộ sơ đồ UML/ERD mô tả hệ thống.",
        "Xây dựng POS, quản lý sản phẩm, kho, khách hàng, đơn hàng, bảo hành, ca, khuyến mãi, VAT và báo cáo.",
        "Tích hợp PayOS, quét mã từ điện thoại, mã vạch CODE128 và trợ lý AI gợi ý sản phẩm/nhập hàng.",
        "Kiểm tra khả năng build, các luồng chính và đánh giá hạn chế để đề xuất hướng phát triển.",
    ))
    heading(doc, "Phương pháp nghiên cứu", 2)
    bullets(doc, (
        "Nghiên cứu tài liệu chính thức của công nghệ và hướng dẫn đề tài.",
        "Phân tích nghiệp vụ, mô hình hóa UML/ERD và thiết kế hợp đồng API.",
        "Thực nghiệm phát triển theo từng phân hệ, tích hợp dịch vụ ngoài và dữ liệu mẫu.",
        "Kiểm thử chấp nhận, kiểm tra tĩnh/build, benchmark tái lập và đối chiếu kết quả với yêu cầu.",
    ))
    heading(doc, "Đóng góp và bố cục báo cáo", 2)
    body(doc, "Đề tài tạo ra hệ thống POS liên kết quanh giao dịch bán hàng cốt lõi. Các điểm nổi bật gồm đồng bộ quét mã từ điện thoại, PayOS có xác minh webhook và chống xử lý lặp, bảo hành theo chi tiết đơn, ca làm việc, nhật ký và AI có phương án dự phòng. Báo cáo gồm sáu chương: tổng quan đề tài; cơ sở lý thuyết; phân tích và thiết kế hệ thống; xây dựng hệ thống; kiểm thử và đánh giá; kết luận và hướng phát triển.")


def chapter1(doc: Document) -> None:
    heading(doc, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI", 1, True)
    heading(doc, "1.1. Bối cảnh và bài toán", 2)
    body(doc, "Cửa hàng đồ gia dụng kinh doanh nhiều nhóm sản phẩm như dụng cụ nhà bếp, đồ vệ sinh, thiết bị điện nhỏ và vật dụng tiện ích. Mỗi sản phẩm có SKU, giá vốn, giá bán, nhà cung cấp, số lượng tồn tối thiểu và chính sách bảo hành khác nhau. Một ngày vận hành còn bao gồm mở ca, nhận nhiều phương thức thanh toán, xử lý khuyến mãi, in hóa đơn, nhập hàng, điều chỉnh kho và đối soát cuối ca.")
    body(doc, "Nếu mỗi nghiệp vụ dùng một bảng tính hoặc sổ riêng, dữ liệu dễ không nhất quán. Nhân viên khó biết tồn thực, chủ cửa hàng không có báo cáo tức thời và khách hàng khó tra cứu bảo hành. Giải pháp cần tập trung dữ liệu, giảm nhập lại, kiểm soát quyền và giữ dấu vết giao dịch.")
    table(doc, "1.1", "Vấn đề của quy trình thủ công", ("Vấn đề", "Tác động", "Giải pháp HomeX POS"), (
        ("Tra cứu thủ công", "Thanh toán chậm, chọn nhầm hàng", "Tìm theo tên/SKU/barcode và hiển thị tồn"),
        ("Hóa đơn rời rạc", "Khó đối soát", "Liên kết Order, OrderDetail và Payment"),
        ("Kho không đồng bộ", "Sai lệch hoặc bán vượt tồn", "Kiểm tra tồn và ghi StockTransaction"),
        ("Bảo hành giấy", "Khó tra cứu, dễ thất lạc", "Mã bảo hành và trang tra cứu công khai"),
        ("Báo cáo thủ công", "Thiếu thông tin kịp thời", "Dashboard, doanh thu và lợi nhuận"),
        ("Thiếu dấu vết", "Khó tìm nguyên nhân sai lệch", "AuditLog và webhook log"),
    ), (4.0, 4.5, 6.9))
    heading(doc, "1.2. Giải pháp đề xuất", 2)
    body(doc, "HomeX POS là ứng dụng web gồm frontend Next.js và backend Express, kết nối PostgreSQL qua Prisma ORM. Frontend cung cấp màn hình theo nghiệp vụ; backend xác thực, phân quyền, validation và thực thi transaction. Giải pháp ưu tiên luồng bán hàng: mở ca, tìm/quét sản phẩm, giỏ hàng, khách hàng, khuyến mãi, đơn nháp, thanh toán, cập nhật kho, điểm và bảo hành.")
    heading(doc, "1.3. Đối tượng sử dụng", 2)
    table(doc, "1.2", "Tác nhân và trách nhiệm", ("Tác nhân", "Trách nhiệm", "Giới hạn"), (
        ("Admin", "Quản lý toàn hệ thống, kho, người dùng, báo cáo và cấu hình", "Yêu cầu vai trò ADMIN"),
        ("Cashier", "Mở/đóng ca, bán hàng, đơn, khách hàng và bảo hành", "Không vào trang quản trị nhạy cảm"),
        ("Khách hàng", "Thanh toán, nhận hóa đơn, tra cứu bảo hành", "Không vào dashboard nội bộ"),
        ("PayOS", "Tạo liên kết thanh toán và gửi webhook", "Hệ thống ngoài, cần cấu hình khóa"),
        ("Dịch vụ AI", "Gemini là chính; Groq là dự phòng", "Kết quả chỉ hỗ trợ; nếu cả hai lỗi thì dùng cách tính cục bộ"),
    ), (3.0, 7.3, 5.1))
    heading(doc, "1.4. Phạm vi chức năng", 2)
    table(doc, "1.3", "Các phân hệ", ("Nhóm", "Chức năng"), (
        ("Xác thực", "Đăng nhập, JWT, khóa tài khoản, role guard"),
        ("Bán hàng", "Giỏ, đơn nháp, khuyến mãi, tiền mặt, PayOS, hóa đơn, quét mã"),
        ("Hàng hóa/kho", "Sản phẩm, danh mục, NCC, barcode, nhập, điều chỉnh, cảnh báo"),
        ("Khách hàng", "Thông tin, điểm, hạng và lịch sử mua"),
        ("Sau bán", "Bảo hành điện tử, trả hàng và hóa đơn VAT"),
        ("Vận hành", "Ca, thông báo, cài đặt và nhật ký"),
        ("Quản trị", "Dashboard, doanh thu, lợi nhuận và top sản phẩm"),
    ), (4.2, 11.2))
    heading(doc, "1.5. Tiêu chí thành công", 2)
    bullets(doc, (
        "Checkout tạo đơn, thanh toán và biến động kho nhất quán; không bán vượt tồn khi cấu hình cấm.",
        "Admin/Cashier thấy đúng chức năng và API chặn truy cập ngoài quyền.",
        "Các phân hệ liên kết với giao dịch thực tế, có thông báo và khả năng truy vết.",
        "Giao diện dùng tốt trên desktop; quét mã và tra cứu công khai dùng trên mobile.",
        "Mã nguồn có cấu trúc, dữ liệu mẫu, hướng dẫn cài đặt và build được.",
    ))


def chapter2(doc: Document) -> None:
    heading(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT", 1, True)

    heading(doc, "2.1. Hệ thống POS và chu trình giao dịch bán lẻ", 2)
    body(doc, "POS (Point of Sale) là điểm ghi nhận giao dịch bán hàng và đồng thời là nút liên kết giữa hàng hóa, thanh toán, tồn kho, khách hàng, ca làm việc và báo cáo. Vì vậy, một hệ thống POS không thể chỉ lưu một hóa đơn độc lập. Mỗi lần checkout phải bảo đảm giá bán, giảm giá, phương thức thanh toán, người bán, số lượng tồn và chứng từ hậu mãi cùng phản ánh một sự kiện nghiệp vụ duy nhất.")
    body(doc, "Trong HomeX POS, đơn hàng được tạo ở trạng thái DRAFT để người dùng kiểm tra giỏ hoặc chờ thanh toán PayOS. Khi tiền mặt được xác nhận hoặc PayOS trả kết quả hợp lệ, backend chuyển Payment sang PAID, Order sang COMPLETED, trừ Product.stockQuantity, ghi StockTransaction, cộng điểm khách hàng và tạo Warranty cho các dòng có warrantyMonths lớn hơn 0. Giá bán và giá vốn được chụp vào OrderDetail để hóa đơn lịch sử không thay đổi khi Product được cập nhật sau này.")
    table(doc, "2.1", "Ánh xạ khái niệm POS vào HomeX POS", ("Khái niệm", "Cách áp dụng", "Dữ liệu liên quan"), (
        ("Giỏ và đơn chờ", "Lưu DRAFT trước checkout hoặc trong khi chờ QR", "Order, OrderDetail"),
        ("Thanh toán", "Tiền mặt hoặc chuyển khoản PayOS", "Payment, PaymentWebhookLog"),
        ("Tồn kho", "Kiểm tra trước bán, trừ trong transaction", "Product, StockTransaction"),
        ("Khách hàng", "Lưu lịch sử, điểm và hạng", "Customer, Order"),
        ("Hậu mãi", "Bảo hành theo dòng, hoàn trả theo chứng từ", "Warranty, ReturnOrder"),
        ("Đối soát", "Gắn đơn với ca và nhật ký thao tác", "Shift, AuditLog"),
    ), (3.0, 7.0, 5.4), 10.5)

    heading(doc, "2.2. Giao dịch cơ sở dữ liệu và tính toàn vẹn", 2)
    body(doc, "Một checkout là thao tác nhiều bước. Nếu hệ thống đã trừ kho nhưng chưa ghi thanh toán, hoặc đã cộng điểm nhưng đơn vẫn DRAFT, dữ liệu sẽ rơi vào trạng thái không nhất quán. Thuộc tính nguyên tử của transaction yêu cầu toàn bộ thay đổi cùng thành công hoặc cùng rollback; tính nhất quán được tăng cường bằng khóa ngoại, unique, enum và validation; tính cô lập hạn chế hai request cạnh tranh; tính bền vững bảo đảm dữ liệu đã commit không mất khi tiến trình kết thúc.")
    body(doc, "HomeX POS dùng prisma.$transaction cho checkout, hoàn trả, nhập hàng và xử lý webhook. Riêng webhook PayOS còn khóa dòng Payment bằng SELECT ... FOR UPDATE. Khi hai thông báo cho cùng giao dịch đến đồng thời, request thứ hai phải chờ request thứ nhất; sau đó nó nhìn thấy trạng thái PAID và chỉ ghi DUPLICATE, không trừ kho hoặc cộng điểm lần nữa. Đây là idempotency ở mức hiệu ứng nghiệp vụ, không phụ thuộc việc trình duyệt còn mở.")
    body(doc, "Tiền tệ dùng Decimal(12,2) thay vì số thực nhị phân để giảm sai số. OrderDetail lưu unitPrice, lineTotal và unitCost tại thời điểm bán. StockTransaction lưu lượng biến động có dấu và tham chiếu chứng từ nguồn, nhờ đó số tồn hiện tại có thể được đối chiếu với lịch sử nhập, bán, điều chỉnh và phục hồi.")

    heading(doc, "2.3. Kiến trúc client-server và REST API", 2)
    body(doc, "Kiến trúc client-server tách giao diện khỏi nơi thực thi quy tắc và lưu dữ liệu. Frontend gửi HTTP request và nhận JSON; backend xác thực token, kiểm tra vai trò, validation, thực thi transaction rồi mới trả kết quả. Cách tách này giúp giao diện web nội bộ, trang tra cứu công khai và điện thoại quét mã dùng chung một nguồn dữ liệu mà không được truy cập trực tiếp PostgreSQL.")
    body(doc, "REST tổ chức chức năng theo tài nguyên: /api/products cho sản phẩm, /api/orders cho đơn, /api/payments cho thanh toán và /api/inventory cho kho. GET dùng đọc, POST tạo hoặc kích hoạt nghiệp vụ, PATCH cập nhật một phần và DELETE/xóa mềm thay đổi trạng thái. Mã HTTP và cấu trúc success/message/data giúp frontend xử lý loading, lỗi validation, không đủ quyền và lỗi nghiệp vụ nhất quán.")
    table(doc, "2.2", "Luồng xử lý một REST request", ("Bước", "Thành phần", "Vai trò trong dự án"), (
        ("1", "Next.js/Axios", "Gửi token, payload và hiển thị trạng thái"),
        ("2", "Express middleware", "CORS, JSON, demo mode, authenticateToken"),
        ("3", "authorizeRoles/Zod", "Kiểm tra quyền và hình dạng dữ liệu"),
        ("4", "Route/service", "Thực thi quy tắc bán, kho, AI hoặc tích hợp"),
        ("5", "Prisma/PostgreSQL", "Truy vấn, transaction và ràng buộc"),
        ("6", "errorMiddleware", "Chuẩn hóa phản hồi lỗi cho frontend"),
    ), (1.4, 4.3, 9.7), 10.5)

    heading(doc, "2.4. TypeScript, React và Next.js trong frontend", 2)
    body(doc, "TypeScript bổ sung hệ thống kiểu cho JavaScript, giúp phát hiện sớm sai hợp đồng dữ liệu [1]. Trong HomeX POS, các kiểu Product, Order, Payment, Warranty và DTO phản hồi API được dùng lại giữa service và component. Kiểu không thay thế validation runtime, nhưng giảm lỗi như đọc nhầm field, truyền sai trạng thái hoặc bỏ sót trường nullable khi phát triển nhiều phân hệ.")
    body(doc, "React 19 biểu diễn giao diện bằng component [2]. HomeX POS dùng state để quản lý giỏ, modal, bộ lọc, trạng thái PayOS và kết quả AI; context quản lý xác thực/ngôn ngữ; component dùng chung quản lý bảng, badge trạng thái, dialog xác nhận và hóa đơn in. Việc tách component giúp POS và trang quản trị giữ hành vi nhất quán.")
    body(doc, "Next.js 16 App Router ánh xạ thư mục app thành 28 trang [3]. Route group phân tách dashboard, auth và public; dynamic route /invoice/[orderCode] phục vụ hóa đơn; /mobile-scan tối ưu cho điện thoại. Production build kiểm tra import, TypeScript, khả năng tạo trang tĩnh/động và là một bằng chứng kỹ thuật trong Chương 5. Tailwind CSS, Radix UI, TanStack Table, React Hook Form và Recharts lần lượt hỗ trợ layout, component tương tác, bảng dữ liệu, biểu mẫu và biểu đồ.")

    heading(doc, "2.5. Express và tổ chức backend", 2)
    body(doc, "Express 5 tổ chức xử lý request qua middleware và router [4]. Backend HomeX POS hiện có 117 khai báo HTTP trong các router và ba endpoint trực tiếp ở tệp index.ts. Mỗi phân hệ có route riêng; logic tích hợp dùng service khi cần tái sử dụng, ví dụ payos.service.ts, sales-assistant.service.ts và email.service.ts. catchAsync chuyển lỗi promise về errorMiddleware, tránh lặp try/catch ở mọi endpoint.")
    body(doc, "authenticateToken đọc Authorization Bearer, xác minh JWT và gắn thông tin người dùng vào request. authorizeRoles áp dụng quyền ADMIN/CASHIER ở server; việc ẩn menu ở frontend chỉ là lớp hỗ trợ trải nghiệm. Zod parse body/query trước khi truy cập CSDL. Cấu trúc này làm rõ trust boundary: mọi giá, tồn, trạng thái thanh toán và quyền cuối cùng đều do backend quyết định.")

    heading(doc, "2.6. PostgreSQL và Prisma ORM", 2)
    body(doc, "PostgreSQL cung cấp giao dịch, khóa dòng, kiểu Decimal, khóa ngoại và chỉ mục phù hợp dữ liệu bán lẻ [5]. Prisma ánh xạ schema thành client có kiểu, hỗ trợ migration và transaction [6]. Schema hiện có 22 model, bao phủ dữ liệu nền, giao dịch, kho, hậu mãi và giám sát.")
    body(doc, "Các ràng buộc unique được đặt cho email, SKU, barcode, orderCode, providerOrderCode và warrantyCode. Order-Payment là một-một trong phạm vi hiện tại; Order-OrderDetail là một-nhiều; Product-StockTransaction và Supplier-PurchaseOrder mô tả lịch sử vận hành. Xóa mềm bằng status được ưu tiên với dữ liệu nền để không phá tham chiếu lịch sử.")
    table(doc, "2.3", "Ràng buộc dữ liệu tiêu biểu", ("Ràng buộc", "Ví dụ", "Lỗi được ngăn chặn"), (
        ("UNIQUE", "Product.sku, Payment.providerOrderCode", "Trùng mã sản phẩm hoặc mã giao dịch"),
        ("FOREIGN KEY", "OrderDetail.productId", "Dòng hàng tham chiếu sản phẩm không tồn tại"),
        ("ENUM", "OrderStatus, PaymentStatus", "Trạng thái ngoài tập hợp cho phép"),
        ("DECIMAL", "Giá vốn, giá bán, tổng tiền", "Sai số số thực khi tính tiền"),
        ("TRANSACTION", "Checkout, return, webhook", "Cập nhật một phần gây lệch dữ liệu"),
        ("ROW LOCK", "Payment khi nhận webhook", "Hai webhook cùng hoàn tất một giao dịch"),
    ), (3.0, 5.7, 7.0), 10.5)

    heading(doc, "2.7. Xác thực, phân quyền và validation", 2)
    body(doc, "Mật khẩu được băm bằng bcrypt; backend không lưu mật khẩu rõ. Khi đăng nhập đúng và tài khoản ACTIVE, server ký JWT chứa định danh và vai trò [7]. Các endpoint nội bộ xác minh token, sau đó kiểm tra role. Cơ chế này phù hợp ứng dụng demo, nhưng production cần bổ sung refresh token, thu hồi phiên, rate limit, CORS cụ thể và quản lý secret tập trung.")
    body(doc, "Zod kiểm tra email, số điện thoại, số lượng, giá, trạng thái và payload nghiệp vụ [8]. Prisma sinh truy vấn tham số hóa, giảm rủi ro SQL Injection. Tuy nhiên, validation không chỉ là kiểm tra kiểu: backend còn tải lại sản phẩm, tính lại subtotal, giới hạn giảm giá, kiểm tra ca, tồn và trạng thái đối tượng thay vì tin dữ liệu frontend.")

    heading(doc, "2.8. PayOS, webhook và mô hình tin cậy", 2)
    body(doc, "PayOS cung cấp API tạo payment link/QR và gửi webhook khi giao dịch thay đổi [9]. HomeX POS tạo Order DRAFT và Payment PENDING, sau đó dùng providerOrderCode duy nhất để tạo payment link. QR chỉ là phương tiện hướng dẫn chuyển tiền; việc frontend hiển thị trang return hoặc thông báo thành công không có quyền chuyển Payment sang PAID.")
    body(doc, "Webhook là server-to-server callback và có thể đến chậm, đến trước khi client polling, được gửi lặp hoặc bị giả mạo. HomeX POS gọi payOS.webhooks.verify để tính lại HMAC-SHA256 từ data bằng checksum key và so với signature. Sau xác minh, backend dùng orderCode từ webhook tìm Payment, đối chiếu số tiền và trạng thái Order/Payment trước khi hoàn tất trong transaction.")
    body(doc, "Idempotency có nghĩa cùng một sự kiện được xử lý nhiều lần nhưng hiệu ứng cuối không thay đổi. Hệ thống kiểm tra event reference đã PROCESSED, khóa dòng Payment, kiểm tra PAID và ghi PROCESSED/DUPLICATE/UNMATCHED/FAILED. Nếu webhook chậm hoặc bị mất, endpoint trạng thái chủ động truy vấn PayOS khi local còn PENDING. Vì nghiệp vụ nằm ở backend, đơn vẫn hoàn tất nếu khách đã trả tiền rồi đóng trình duyệt.")
    table(doc, "2.4", "Các lớp kiểm soát webhook PayOS", ("Nguy cơ", "Kiểm soát", "Kết quả mong muốn"), (
        ("Payload giả mạo", "SDK xác minh signature bằng checksum key", "Từ chối trước khi tra cứu giao dịch"),
        ("Replay/gửi lặp", "event reference, row lock và trạng thái PAID", "Không trừ kho/cộng điểm lần hai"),
        ("Sai số tiền", "So Payment.amount với webhook amount", "FAILED và chuyển kiểm tra thủ công"),
        ("Sai mã đơn", "Tra providerOrderCode và đối chiếu trong transaction", "UNMATCHED hoặc từ chối"),
        ("Client đóng", "Webhook xử lý độc lập frontend", "Backend vẫn hoàn tất đơn"),
        ("Webhook trễ/mất", "Polling và truy vấn trực tiếp PayOS", "Đồng bộ lại trạng thái có căn cứ"),
    ), (3.1, 7.1, 5.5), 10.2)

    heading(doc, "2.9. Phân loại chức năng hỗ trợ thông minh", 2)
    body(doc, "HomeX POS không coi mọi chức năng có LLM là một mô hình AI cùng loại. Các dịch vụ AI dùng Gemini làm lựa chọn chính và Groq làm phương án dự phòng qua giao diện tương thích của thư viện OpenAI [10], [11]. Dự án không sử dụng OPENAI_API_KEY; thư viện OpenAI ở đây chỉ là công cụ gửi yêu cầu theo cùng định dạng. Riêng nhập hàng, số lượng đề xuất được tính bằng công thức xác định; LLM chỉ diễn giải và phân loại sau phép tính, không được dự án huấn luyện và không tạo dự báo gốc.")
    table(doc, "2.5", "Phân loại chức năng hỗ trợ thông minh", ("Chức năng", "Bản chất", "Đầu vào", "Đầu ra/giới hạn"), (
        ("Trợ lý bán hàng", "LLM xếp hạng trên tập ứng viên đã lọc", "Nhu cầu, ngân sách, giỏ", "Tối đa 5 ID có thật; thu ngân quyết định"),
        ("Ước lượng nhập hàng", "Công thức theo luật; LLM diễn giải", "Bán 7/30 ngày, tồn, minStock, danh mục, tháng", "Đề xuất tham khảo; không phải mô hình được huấn luyện"),
        ("Làm giàu barcode/ảnh", "Tra API và vision hỗ trợ điền biểu mẫu", "Barcode, dữ liệu ngoài, ảnh nhãn", "Admin kiểm tra trước khi lưu; có thể sai"),
    ), (3.0, 4.6, 4.4, 4.3), 9.3)

    heading(doc, "2.10. AI gợi ý bán hàng", 2)
    body(doc, "sales-assistant.service.ts chỉ tải sản phẩm ACTIVE và stockQuantity > 0, lọc theo ngân sách, sau đó chấm điểm sơ bộ theo khuyến mãi, bảo hành, tồn cao, danh mục trong giỏ và từ khóa nhu cầu. Tối đa 25 ứng viên được gửi cho mô hình. Prompt cấm tạo sản phẩm, giảm giá hoặc ID ngoài danh sách.")
    body(doc, "Sau khi nhận JSON, backend tra từng productId trong candidatesToSend và lấy lại tên, giá, tồn từ CSDL; nội dung AI không được quyền sửa các trường này. Nếu Gemini lỗi, quá 15 giây hoặc trả JSON không dùng được, hệ thống thử Groq; nếu vẫn không thành công, thuật toán chấm điểm cục bộ trả tối đa năm sản phẩm. Vì vậy tiêu chí an toàn quan trọng hơn câu văn thuyết phục: ID phải tồn tại, sản phẩm phải còn hàng, giá phải trong ngân sách và người thu ngân vẫn là người phê duyệt.")

    heading(doc, "2.11. Ước lượng nhu cầu và đề xuất nhập hàng theo luật", 2)
    body(doc, "Endpoint /api/inventory/ai-forecast hiện giữ tên kỹ thuật để tương thích giao diện, nhưng về phương pháp đây là bộ ước lượng nhu cầu theo luật, không phải mô hình dự báo được huấn luyện. Dự án không có tập training, không học tham số và không fine-tune LLM. Dữ liệu lịch sử chỉ là các OrderDetail thuộc đơn COMPLETED hoặc đã PAID trong hai cửa sổ 7 và 30 ngày; đơn CANCELLED hoặc thanh toán REFUNDED bị loại.")
    body(doc, "Các đại lượng hiện thực gồm avgDailySales30 = soldLast30Days/30; avgDailySales7 = soldLast7Days/7; trendRatio = avgDailySales7/avgDailySales30; categoryTrendRatio được tính tương tự trên danh mục; seasonBoost là hệ số nghiệp vụ cố định theo tháng và từ khóa danh mục. Công thức predictedDailySales = avgDailySales30 × trendRatio × categoryTrendRatio × seasonBoost. Nhu cầu kỳ phân tích bằng predictedDailySales × số ngày; tồn an toàn bằng 20% nhu cầu; lượng đề xuất là phần thiếu lớn hơn giữa minStock và nhu cầu cộng tồn an toàn.", False)
    body(doc, "Hệ số mùa vụ từ 1,10 đến 1,35 là luật do dự án đặt, không phải tham số học từ nhiều năm. Thuật toán hiện không nhận kế hoạch khuyến mãi tương lai. Nếu 7 và 30 ngày đều không có bán, hệ thống chỉ nhập bù đến minStock khi tồn hiện tại đang thấp hơn minStock; nếu tồn đã đủ thì đề xuất bằng 0. Độ tin cậy được đặt LOW vì chưa có tín hiệu bán. Gemini hoặc Groq chỉ diễn giải kết quả đã tính và không tự tạo PurchaseOrder.")

    heading(doc, "2.11.1. Baseline và chỉ số đánh giá", 3)
    body(doc, "Đánh giá chuỗi thời gian phải tách quá khứ và tương lai theo thứ tự thời gian. HomeX POS dùng rolling-origin: tại mỗi mốc chỉ dùng dữ liệu trước mốc để dự đoán kỳ sau. Hai baseline là trung bình trượt 7 ngày và trung bình trượt 30 ngày. Baseline đơn giản giúp xác định công thức nhiều hệ số có thực sự tốt hơn việc dùng tốc độ bán gần đây hay không.")
    table(doc, "2.6", "Chỉ số sai số dùng trong backtest", ("Chỉ số", "Công thức/ý nghĩa", "Lưu ý với dữ liệu HomeX"), (
        ("MAE", "Trung bình |dự đoán - thực tế|", "Đơn vị sản phẩm; dễ diễn giải"),
        ("RMSE", "Căn bậc hai trung bình bình phương sai số", "Phạt nặng sai số lớn"),
        ("WAPE", "Tổng sai số tuyệt đối / tổng nhu cầu", "Phù hợp hơn khi nhiều ngày bằng 0"),
        ("MAPE", "Trung bình sai số phần trăm", "Chỉ tính actual > 0; phải công bố số mẫu"),
        ("Bias", "Trung bình dự đoán - thực tế", "Dương là xu hướng nhập dư; âm là thiếu"),
    ), (2.6, 7.5, 5.3), 9.2)
    body(doc, "Một công thức không được coi là chính xác chỉ vì trả JSON hợp lệ hoặc số lượng không âm. Nếu MAE/RMSE cao hơn baseline, kết luận đúng là công thức chưa đủ bằng chứng để dùng như dự báo tự động; nó chỉ nên là cảnh báo và đề xuất tham khảo.")

    heading(doc, "2.12. Làm giàu sản phẩm theo barcode và ảnh", 2)
    body(doc, "Luồng làm giàu ưu tiên sản phẩm đã có trong HomeX POS. Nếu chưa có, backend gọi tuần tự các nguồn có giao diện được tài liệu hóa, hợp nhất trường không rỗng và chuẩn hóa category/supplier. Nguồn HTML không chính thức không nằm trong chuỗi tra cứu của phiên bản báo cáo này. Các khóa dịch vụ ở backend; frontend chỉ nhận dữ liệu gợi ý.")
    table(doc, "2.7", "Nguồn dữ liệu barcode của phiên bản đánh giá", ("Nguồn", "Cách dùng/trạng thái", "Dữ liệu và quyền sử dụng"), (
        ("HomeX POS", "Ưu tiên đầu tiên; barcode trùng trả bản ghi nội bộ", "Dữ liệu do cửa hàng quản lý; confidence 1"),
        ("UPCitemdb", "Endpoint trial lookup; không cần key; giới hạn 100 request/ngày [12]", "Tên, brand, loại, giá/ảnh/mô tả nếu có; quyền giới hạn, không bảo đảm chính xác"),
        ("Barcode Lookup v3", "API key phía server; đang cấu hình trong môi trường đánh giá [13]", "Product Data theo thuê bao; ảnh/mô tả có thể là nội dung bên thứ ba"),
        ("Open Products Facts", "API v2 cho hàng tổng quát; dự án Open Facts còn thử nghiệm [15]", "Dữ liệu ODbL/DBCL; ảnh CC BY-SA; cần ghi nguồn/share-alike [14]"),
        ("Open Food Facts", "Nguồn phụ cho barcode thực phẩm [14]", "Dữ liệu cộng đồng có thể thiếu/sai; cùng nghĩa vụ giấy phép"),
        ("Barcode Spider", "Mã tích hợp tùy chọn; không có key trong môi trường đánh giá [16]", "Chỉ dùng khi có thuê bao; quyền ảnh/nội dung bên thứ ba do người dùng chịu trách nhiệm"),
        ("Gemini/Groq", "Gemini là chính; Groq dự phòng cho chữ và ảnh", "Không phải nguồn sự thật; Admin phải kiểm tra trước khi lưu"),
    ), (3.0, 6.2, 6.2), 8.5)
    body(doc, "Barcode Lookup quy định Product Data theo quyền sử dụng giới hạn của thuê bao, không bảo đảm luôn có kết quả hoặc chính xác; quyền đối với ảnh/mô tả của bên thứ ba không tự động được cấp [13]. UPCitemdb cũng cung cấp dữ liệu as-is và đặt trách nhiệm không xâm phạm quyền bên thứ ba lên người sử dụng [12]. Với Open Facts, việc tái sử dụng phải ghi nguồn và tuân thủ ODbL/DBCL; ảnh có giấy phép CC BY-SA và vẫn có thể chứa yếu tố đồ họa thuộc quyền khác [14]. Vì vậy bản demo lưu URL/giá trị sau khi Admin duyệt, còn triển khai thương mại phải rà lại hợp đồng, attribution và chính sách lưu trữ của từng nguồn.")

    heading(doc, "2.12.1. Dữ liệu thiếu, xung đột và bước duyệt", 3)
    table(doc, "2.8", "Kiểm soát chất lượng dữ liệu làm giàu", ("Tình huống", "Xử lý hiện tại", "Giới hạn/rủi ro"), (
        ("API timeout/404/429", "Bắt lỗi, trả null và chuyển nguồn kế tiếp", "Có thể không tìm được dù barcode hợp lệ"),
        ("Nguồn trả thiếu trường", "Hợp nhất trường có ích; trả missingFields", "Không được coi trường suy đoán là dữ liệu gốc"),
        ("Nguồn mâu thuẫn", "Ưu tiên theo thứ tự lookup và chỉ điền ô trống", "HYBRID chưa lưu provenance tới từng field"),
        ("Không có danh tính đáng tin", "Chỉ đọc ảnh nếu có URL; ảnh không rõ trả object rỗng", "Vision có thể nhầm biến thể/brand"),
        ("AI bù thuộc tính", "Sanitize với dữ liệu external; không đổi name đã tin cậy", "Giá, tồn, bảo hành chỉ là gợi ý"),
        ("Lưu sản phẩm", "Admin xem form đã điền, sửa nếu cần rồi bấm Lưu", "Endpoint enrich không tự tạo hoặc cập nhật Product"),
    ), (3.5, 7.0, 4.9), 8.8)
    body(doc, "Điểm confidence 0,58–0,84 trong service là trọng số quy ước theo loại nguồn, chưa được hiệu chuẩn thành xác suất đúng. AI vẫn có thể tạo sai tên, thương hiệu, danh mục hoặc thông số. Benchmark 12 ảnh ở Chương 5 cho thấy tên đúng 33,3% và danh mục đúng 50%; do đó mọi trường từ AI phải được trình bày như gợi ý, không ghi đè dữ liệu đã có và không tự lưu.")

    heading(doc, "2.13. Mã vạch và quét từ xa", 2)
    body(doc, "CODE128 phù hợp SKU chữ-số và được JsBarcode sinh để in tem. Máy quét USB hoạt động như bàn phím; với điện thoại, POS tạo sessionId và QR chứa URL /mobile-scan. html5-qrcode đọc camera [17], gửi barcode về /api/pos/remote-scan, còn màn hình POS polling /remote-scan-poll/:sessionId để nhận mã. Session có thời hạn và chỉ chuyển dữ liệu barcode, không chuyển quyền thanh toán sang điện thoại.")

    heading(doc, "2.14. Logging, khả năng phục hồi và giới hạn", 2)
    body(doc, "Hệ thống dùng AuditLog cho thao tác người dùng, StockTransaction cho biến động tồn và PaymentWebhookLog cho callback PayOS. Theo hướng kiểm soát bảo mật ứng dụng [18], ba loại log trả lời ba câu hỏi khác nhau: ai thao tác, tồn thay đổi vì chứng từ nào và cổng thanh toán đã gửi gì. Các service ngoài đều có lỗi kiểm soát; AI có fallback cục bộ, PayOS có truy vấn trạng thái dự phòng, email là tùy chọn.")
    body(doc, "Các cơ chế trên phù hợp môi trường đồ án và demo, chưa thay thế observability production. Hệ thống chưa có distributed tracing, hàng đợi webhook, retry worker, metric/alert tập trung hoặc secret manager. Những giới hạn này được đưa vào đánh giá và hướng phát triển thay vì giả định dịch vụ ngoài luôn khả dụng.")

    heading(doc, "2.15. Tổng hợp công nghệ và vai trò thực tế", 2)
    table(doc, "2.9", "Công nghệ, phiên bản và cách dùng", ("Thành phần", "Công nghệ", "Cách dùng trong HomeX POS"), (
        ("Frontend", "Next.js 16.2.6, React 19.2.4, TypeScript", "28 trang, state/component, routing và type"),
        ("UI", "Tailwind 4, Radix UI, Lucide", "Responsive, dialog, tab và điều hướng"),
        ("Dữ liệu UI", "Axios, TanStack Table, RHF, Zod", "API, bảng, form và validation"),
        ("Backend", "Express 5.2.1, TypeScript 6, Zod", "117 khai báo HTTP trong router, middleware và rule"),
        ("CSDL", "PostgreSQL, Prisma 6.19.3", "22 model, transaction, row lock và migration"),
        ("Tích hợp", "PayOS 2.0.5, OpenAI SDK 6.45.0, Gemini, Groq", "Gemini là AI chính; Groq dự phòng; SDK chỉ dùng làm giao diện gửi yêu cầu"),
        ("Hiển thị", "Recharts, JsBarcode, html5-qrcode", "Biểu đồ, tem barcode và camera scan"),
    ), (3.0, 5.2, 7.2), 10.2)

def chapter3(doc: Document) -> None:
    heading(doc, "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1, True)

    heading(doc, "3.1. Quy trình nghiệp vụ tổng quát", 2)
    body(doc, "Quy trình vận hành bắt đầu khi nhân viên đăng nhập và mở ca. Thu ngân tìm hoặc quét sản phẩm, tạo giỏ, chọn khách hàng/khuyến mãi rồi thanh toán tiền mặt hoặc PayOS. Một giao dịch hoàn tất làm phát sinh Payment, StockTransaction, điểm khách hàng và Warranty nếu sản phẩm có bảo hành. Cuối ca, tiền dự kiến được đối chiếu với tiền thực tế.")
    picture(doc, "2_Sơ đồ/01_use_case_tong_quat.png", "3.1", "Sơ đồ ca sử dụng tổng quát")
    body(doc, "Cashier tập trung bán tại quầy, đơn hàng, khách hàng, bảo hành và ca. Admin quản lý thêm sản phẩm, danh mục, nhà cung cấp, kho, phiếu nhập, hoàn trả, khuyến mãi, VAT, người dùng, báo cáo, audit và settings. Khách hàng chỉ tương tác với hóa đơn/tra cứu bảo hành công khai. PayOS và mô hình AI là tác nhân ngoài, không được quyền truy cập trực tiếp CSDL.")

    heading(doc, "3.2. Tác nhân và ranh giới quyền", 2)
    table(doc, "3.1", "Tác nhân của hệ thống", ("Tác nhân", "Mục tiêu", "Giới hạn quyền"), (
        ("Admin", "Quản trị dữ liệu, kho, báo cáo và cấu hình", "Mọi chức năng nội bộ theo ADMIN"),
        ("Cashier", "Bán hàng, khách hàng, bảo hành và ca", "Không quản lý kho/cấu hình/người dùng nhạy cảm"),
        ("Khách hàng", "Thanh toán và tra cứu chứng từ/bảo hành", "Chỉ endpoint công khai, dữ liệu tối thiểu"),
        ("PayOS", "Tạo link và thông báo kết quả thanh toán", "Chỉ qua API/webhook có chữ ký"),
        ("Dịch vụ AI", "Xếp hạng, diễn giải và đọc nhãn hỗ trợ", "Chỉ nhận ứng viên; không ghi giao dịch"),
        ("Điện thoại quét", "Gửi barcode về phiên POS", "Không có quyền checkout hoặc xem dữ liệu quản trị"),
    ), (3.0, 7.1, 5.3), 10.3)
    picture(doc, "2_Sơ đồ/02_use_case_admin_chi_tiet.png", "3.2", "Sơ đồ ca sử dụng chi tiết của Admin", max_height=15.5)

    heading(doc, "3.3. Yêu cầu chức năng", 2)
    table(doc, "3.2", "Danh sách yêu cầu chức năng", ("Mã", "Yêu cầu", "Ưu tiên"), (
        ("FR-01", "Đăng nhập và phân quyền ADMIN/CASHIER", "Bắt buộc"),
        ("FR-02", "Sản phẩm, danh mục, nhà cung cấp, barcode và làm giàu dữ liệu", "Bắt buộc/Nâng cao"),
        ("FR-03", "Mở/đóng ca và đối soát", "Bắt buộc"),
        ("FR-04", "Giỏ, đơn nháp, khuyến mãi và checkout", "Bắt buộc"),
        ("FR-05", "Thanh toán tiền mặt và PayOS", "Bắt buộc/Nâng cao"),
        ("FR-06", "Nhập, bán, điều chỉnh, phục hồi và lịch sử kho", "Bắt buộc"),
        ("FR-07", "Khách hàng, điểm và hạng thành viên", "Nâng cao"),
        ("FR-08", "Tạo, quản lý và tra cứu bảo hành điện tử", "Đặc thù"),
        ("FR-09", "Trả hàng và hóa đơn VAT", "Nâng cao"),
        ("FR-10", "Dashboard, doanh thu, lợi nhuận và top sản phẩm", "Bắt buộc"),
        ("FR-11", "Nhật ký, thông báo và cấu hình", "Nâng cao"),
        ("FR-12", "AI gợi ý bán hàng và ước lượng nhu cầu nhập", "Nâng cao"),
        ("FR-13", "Quét barcode bằng điện thoại", "Nâng cao"),
    ), (2.0, 10.4, 3.0), 10.2)

    heading(doc, "3.4. Danh mục đầy đủ các ca sử dụng", 2)
    body(doc, "Để thể hiện toàn bộ phạm vi thay vì chỉ mô tả một vài ca tiêu biểu, bảng dưới đây lập danh mục đầy đủ theo chức năng đang tồn tại trong route, trang giao diện và schema. Mã UC được dùng nhất quán trong phần đặc tả và kiểm thử.")
    table(doc, "3.3", "Danh mục ca sử dụng", ("Mã", "Tên ca sử dụng", "Tác nhân chính", "Liên hệ yêu cầu"), (
        ("UC-01", "Đăng nhập/đăng xuất và kiểm tra quyền", "Admin, Cashier", "FR-01"),
        ("UC-02", "Mở, theo dõi và đóng ca", "Admin, Cashier", "FR-03"),
        ("UC-03", "Quản lý sản phẩm, barcode và làm giàu dữ liệu", "Admin", "FR-02"),
        ("UC-04", "Quản lý danh mục và nhà cung cấp", "Admin", "FR-02"),
        ("UC-05", "Bán hàng và thanh toán tiền mặt", "Admin, Cashier", "FR-04, FR-05"),
        ("UC-06", "Thanh toán PayOS và xử lý webhook", "Admin, Cashier, PayOS", "FR-05"),
        ("UC-07", "Tra cứu, in và hủy đơn theo quyền", "Admin, Cashier", "FR-04"),
        ("UC-08", "Quản lý khách hàng, điểm và hạng", "Admin, Cashier", "FR-07"),
        ("UC-09", "Quản lý và áp dụng khuyến mãi", "Admin, Cashier", "FR-04"),
        ("UC-10", "Theo dõi, nhập và điều chỉnh tồn", "Admin", "FR-06"),
        ("UC-11", "Ước lượng nhu cầu và tạo phiếu nhập", "Admin", "FR-06, FR-12"),
        ("UC-12", "Hoàn trả và phục hồi tồn", "Admin", "FR-09"),
        ("UC-13", "Tạo, xử lý và tra cứu bảo hành", "Admin, Cashier, Khách", "FR-08"),
        ("UC-14", "Yêu cầu, duyệt và gửi hóa đơn VAT", "Admin, Khách", "FR-09"),
        ("UC-15", "Xem dashboard và báo cáo", "Admin", "FR-10"),
        ("UC-16", "Quản lý người dùng, settings và audit", "Admin", "FR-01, FR-11"),
        ("UC-17", "Quét barcode từ điện thoại", "Cashier, Điện thoại", "FR-13"),
        ("UC-18", "AI gợi ý sản phẩm/bán kèm", "Admin, Cashier", "FR-12"),
    ), (1.7, 6.6, 4.1, 3.0), 9.2)

    heading(doc, "3.5. Tiêu chí chọn ca sử dụng để đặc tả chi tiết", 2)
    body(doc, "Use case được mô tả theo cách tiếp cận UML [19]. Không phải mọi CRUD đều cần đặc tả dài như nhau. Tám use case dưới đây được chọn vì đại diện ít nhất một tiêu chí: tạo giá trị cốt lõi, cập nhật nhiều bảng trong transaction, chứa rủi ro bảo mật/tài chính, tích hợp dịch vụ ngoài hoặc thể hiện đặc thù cửa hàng đồ gia dụng. Các use case còn lại vẫn có trong danh mục đầy đủ và được kiểm tra ở Chương 5.")
    table(doc, "3.4", "Tiêu chí lựa chọn use case trọng yếu", ("Use case", "Cốt lõi", "Transaction", "Rủi ro/tích hợp", "Đặc thù"), (
        ("UC-01 Đăng nhập", "Nền tảng", "Không", "JWT, role", "Không"),
        ("UC-05 Bán tiền mặt", "Có", "Có", "Tồn, tiền, ca", "Bán tại quầy"),
        ("UC-06 PayOS", "Có", "Có", "Chữ ký, replay", "Thanh toán QR"),
        ("UC-11 Ước lượng nhập", "Có", "Khi tạo phiếu", "Dữ liệu ít/LLM", "NCC, nhập hàng"),
        ("UC-12 Hoàn trả", "Có", "Có", "Không hoàn vượt", "Đổi trả"),
        ("UC-13 Bảo hành", "Có", "Có", "Quyền/dữ liệu công khai", "Bảo hành"),
        ("UC-17 Remote scan", "Hỗ trợ", "Không", "Phiên/timeout", "Barcode"),
        ("UC-18 AI bán hàng", "Hỗ trợ", "Không", "ID giả/hết hàng", "Bán kèm"),
    ), (4.3, 2.3, 2.6, 3.5, 2.7), 9.5)

    heading(doc, "3.6. Đặc tả các ca sử dụng trọng yếu", 2)
    use_cases = (
        ("3.5", "3.6.1. UC-01 - Đăng nhập và phân quyền", (
            ("Tác nhân", "Admin, Cashier"), ("Tiền điều kiện", "Tài khoản tồn tại và ACTIVE"),
            ("Luồng chính", "Nhập email/mã nhân viên và mật khẩu; Zod kiểm tra; bcrypt xác minh; ký JWT; frontend tải profile và menu theo role"),
            ("Ngoại lệ", "Sai định dạng, sai mật khẩu, tài khoản bị khóa, token hết hạn, truy cập URL trái quyền"),
            ("Hậu điều kiện", "Client giữ token; mọi request nội bộ gắn Bearer; API vẫn kiểm tra role độc lập giao diện"),
        )),
        ("3.6", "3.6.2. UC-05 - Bán hàng và thanh toán tiền mặt", (
            ("Tác nhân", "Cashier/Admin; khách hàng"), ("Tiền điều kiện", "Đăng nhập; Cashier có ca OPEN; sản phẩm ACTIVE và đủ tồn"),
            ("Luồng chính", "Tìm/quét hàng; chọn khách/khuyến mãi; tạo DRAFT; backend tính lại giá; nhập tiền nhận; checkout; in hóa đơn"),
            ("Ngoại lệ", "Ca chưa mở, giỏ rỗng, vượt tồn, khuyến mãi sai, tiền nhận thiếu, khách hàng inactive"),
            ("Hậu điều kiện", "Order COMPLETED, Payment PAID, tồn/điểm/bảo hành/audit cập nhật trong transaction"),
        )),
        ("3.7", "3.6.3. UC-06 - Thanh toán PayOS", (
            ("Tác nhân", "Cashier/Admin, PayOS"), ("Tiền điều kiện", "Order DRAFT; chưa có Payment khác; khóa PayOS được cấu hình"),
            ("Luồng chính", "Backend tính tiền; tạo Payment PENDING và providerOrderCode; gọi PayOS; trả QR; PayOS gửi webhook; xác minh chữ ký; khóa Payment; đối chiếu mã/số tiền; hoàn tất transaction"),
            ("Ngoại lệ", "Payload sai chữ ký, mã không khớp, số tiền lệch, tồn thay đổi, webhook lặp, PayOS lỗi hoặc client đóng"),
            ("Hậu điều kiện", "Chỉ nguồn PayOS hợp lệ chuyển PAID; log PROCESSED/DUPLICATE/UNMATCHED/FAILED; client polling nhận trạng thái"),
        )),
        ("3.8", "3.6.4. UC-11 - Ước lượng nhu cầu và tạo phiếu nhập", (
            ("Tác nhân", "Admin"), ("Tiền điều kiện", "Có sản phẩm/NCC; dữ liệu bán 7/30 ngày có thể rỗng"),
            ("Luồng chính", "Chọn kỳ 7/15/30 ngày; công thức theo luật tính nhu cầu; LLM chỉ diễn giải nếu khả dụng; Admin duyệt SKU/số lượng; chọn NCC; tạo PurchaseOrder"),
            ("Ngoại lệ", "LLM rate-limit/JSON sai vẫn dùng kết quả công thức; sản phẩm mới/không bán trả confidence LOW; khuyến mãi tương lai chưa được mô hình hóa; dữ liệu âm hoặc NCC inactive bị từ chối"),
            ("Hậu điều kiện", "Gợi ý không tự đổi kho; chỉ phiếu được duyệt mới tăng tồn và ghi IMPORT"),
        )),
        ("3.9", "3.6.5. UC-12 - Hoàn trả", (
            ("Tác nhân", "Admin"), ("Tiền điều kiện", "Order COMPLETED; dòng hàng còn số lượng có thể hoàn"),
            ("Luồng chính", "Chọn đơn/dòng/số lượng; kiểm tra tổng đã hoàn; tạo ReturnOrder/Item; phục hồi tồn; cập nhật bảo hành liên quan"),
            ("Ngoại lệ", "Hoàn vượt số mua, đơn không hợp lệ, số lượng âm, sản phẩm/dòng không thuộc đơn"),
            ("Hậu điều kiện", "Chứng từ hoàn được giữ; StockTransaction RESTORE; không xóa hóa đơn gốc"),
        )),
        ("3.10", "3.6.6. UC-13 - Bảo hành", (
            ("Tác nhân", "Admin, Cashier, Khách hàng"), ("Tiền điều kiện", "Dòng hàng có warrantyMonths > 0 hoặc đủ điều kiện tạo thủ công"),
            ("Luồng chính", "Checkout tạo mã/ngày bắt đầu-kết thúc; nhân viên xem, claim, complete/reject; khách tra mã hoặc số điện thoại"),
            ("Ngoại lệ", "Hết hạn, đã hủy/hoàn, tạo trùng orderDetail, mã không tồn tại, dữ liệu công khai không hợp lệ"),
            ("Hậu điều kiện", "Warranty gắn duy nhất OrderDetail; trạng thái và ghi chú được lưu; endpoint công khai chỉ trả dữ liệu cần thiết"),
        )),
        ("3.11", "3.6.7. UC-17 - Quét barcode từ điện thoại", (
            ("Tác nhân", "Cashier, điện thoại"), ("Tiền điều kiện", "POS tạo sessionId còn hiệu lực; điện thoại mở URL QR"),
            ("Luồng chính", "Camera đọc CODE128/QR; điện thoại POST barcode; POS polling nhận mã; backend/frontend tìm sản phẩm và thêm giỏ"),
            ("Ngoại lệ", "Hết phiên, camera bị từ chối, barcode không tồn tại, sản phẩm inactive/hết tồn"),
            ("Hậu điều kiện", "Chỉ barcode được chuyển; checkout vẫn do người dùng POS xác nhận"),
        )),
        ("3.12", "3.6.8. UC-18 - AI gợi ý bán hàng", (
            ("Tác nhân", "Admin, Cashier"), ("Tiền điều kiện", "Có sản phẩm ACTIVE còn hàng; nhu cầu/ngân sách có thể để trống"),
            ("Luồng chính", "Lọc và chấm điểm ứng viên; gọi mô hình nếu có token; hậu kiểm ID; trả tối đa 5 gợi ý; người dùng chọn thêm giỏ"),
            ("Ngoại lệ", "AI lỗi/rate-limit/JSON sai dùng heuristic; ID lạ bị bỏ; không có ứng viên trả danh sách rỗng"),
            ("Hậu điều kiện", "AI không sửa giá, tồn, khuyến mãi hoặc tự thêm hàng"),
        )),
    )
    for number, title, rows in use_cases:
        heading(doc, title, 3)
        table(doc, number, f"Đặc tả {title.split(' - ', 1)[-1]}", ("Thuộc tính", "Nội dung"), rows, (3.5, 11.9), 10.2)

    heading(doc, "3.7. Quy trình hoạt động bán hàng", 2)
    picture(doc, "2_Sơ đồ/04_activity_ban_hang_tai_quay.png", "3.3", "Sơ đồ hoạt động bán hàng tại quầy", max_height=18)
    body(doc, "Sơ đồ chia trách nhiệm giữa thu ngân, hệ thống và CSDL. Các điểm quyết định nằm ở ca, tồn, giảm giá, phương thức và kết quả thanh toán. Dữ liệu chỉ cập nhật sau khi điều kiện backend được thỏa mãn; PayOS return page không phải bằng chứng trả tiền.")

    heading(doc, "3.8. Mô hình trạng thái", 2)
    picture(doc, "2_Sơ đồ/06_1_state_vong_doi_don_hang.png", "3.4", "Vòng đời đơn hàng", 11.5, 15)
    body(doc, "Order bắt đầu DRAFT, chuyển COMPLETED khi Payment hợp lệ hoặc CANCELLED theo quyền. Payment PENDING của PayOS chỉ chuyển PAID từ webhook đã xác minh hoặc truy vấn trực tiếp provider. Hoàn trả tạo ReturnOrder liên kết thay vì xóa lịch sử đơn gốc.")
    picture(doc, "2_Sơ đồ/06_2_state_canh_bao_ton_kho.png", "3.5", "Trạng thái cảnh báo tồn kho", 10.5, 13)

    heading(doc, "3.9. Yêu cầu phi chức năng", 2)
    table(doc, "3.13", "Yêu cầu phi chức năng", ("Mã", "Nhóm", "Yêu cầu kiểm tra"), (
        ("NFR-01", "Bảo mật", "bcrypt, JWT, role và xác minh webhook server"),
        ("NFR-02", "Toàn vẹn", "Transaction, Decimal, row lock và idempotency"),
        ("NFR-03", "Hiệu năng", "Tìm kiếm/phân trang; AI giới hạn 25/40 ứng viên"),
        ("NFR-04", "Khả dụng", "Loading, lỗi, fallback AI và polling PayOS"),
        ("NFR-05", "Dễ dùng", "POS thao tác nhanh, quét mã, tồn và tổng tiền rõ"),
        ("NFR-06", "Tương thích", "Desktop/tablet; mobile cho quét/tra cứu"),
        ("NFR-07", "Bảo trì", "TypeScript, router/component/service theo mô-đun"),
        ("NFR-08", "Truy vết", "StockTransaction, webhook log và AuditLog"),
        ("NFR-09", "An toàn AI", "Không bịa ID, hậu kiểm, không tự ghi giao dịch"),
    ), (2.0, 3.2, 10.2), 10.2)

    heading(doc, "3.10. Quy tắc nghiệp vụ", 2)
    rules = (
        "Chỉ tài khoản ACTIVE đăng nhập; endpoint quản trị chỉ cho ADMIN.", "Cashier phải có ca OPEN trước checkout.",
        "Backend tải lại sản phẩm và tính tiền; không tin giá/tổng tiền từ frontend.", "Không bán vượt tồn nếu allowOversell=false; số lượng và giá phải dương.",
        "Khuyến mãi phải còn hạn, đủ điều kiện và chưa vượt giới hạn.", "Webhook PayOS phải đúng chữ ký, mã đơn và số tiền; dòng Payment bị khóa khi xử lý.",
        "Payment PAID hoặc event đã PROCESSED không tạo hiệu ứng lần hai.", "Bảo hành chỉ tạo cho dòng hàng có warrantyMonths > 0 và không trùng orderDetail.",
        "Số lượng hoàn không vượt phần đã mua còn có thể hoàn; hoàn kho ghi RESTORE.", "Gợi ý AI chỉ dùng SKU/ID có thật; Admin/người bán phê duyệt trước tác vụ.",
        "Sản phẩm không có doanh số 7/30 ngày chỉ được đề xuất nhập bù minStock khi kho đang thiếu và có confidence LOW; nếu tồn đã đủ thì lượng nhập bằng 0.", "Đóng ca lưu tiền dự kiến, thực tế và chênh lệch.",
    )
    table(doc, "3.14", "Các quy tắc nghiệp vụ chính", ("Mã", "Quy tắc"), tuple((f"BR-{i:02}", rule) for i, rule in enumerate(rules, 1)), (2.1, 13.3), 10.3)

    heading(doc, "3.11. Phân tích nghiệp vụ đồ gia dụng", 2)
    body(doc, "Đặc thù đồ gia dụng được thể hiện sâu nhất ở nguồn hàng, bảo hành và hậu mãi. Product có supplierId, costPrice, salePrice, minStock và warrantyMonths; PurchaseOrder/Item lưu chứng từ nhập; Warranty gắn OrderDetail để xác định đúng sản phẩm đã mua; ReturnOrder/Item kiểm soát số lượng hoàn. Cách thiết kế này đủ cho cửa hàng một kho bán hàng theo số lượng, chưa phải quản lý tài sản theo từng serial.")
    table(doc, "3.15", "Nghiệp vụ đồ gia dụng đã triển khai", ("Nghiệp vụ", "Cách xử lý hiện tại", "Giới hạn"), (
        ("Thời hạn bảo hành", "warrantyMonths theo Product; ngày kết thúc tính lúc bán", "Không phiên bản chính sách theo thời gian"),
        ("Nơi bảo hành", "HomeX quản lý trạng thái/ghi chú", "Chưa phân biệt cửa hàng hay hãng"),
        ("Nhà cung cấp", "Supplier, lịch sử PurchaseOrder", "Chưa hợp đồng/công nợ"),
        ("Phiếu nhập", "NCC, item, số lượng, giá vốn và IMPORT", "Một kho; chưa lô/serial"),
        ("Đổi trả", "ReturnOrder theo OrderDetail, phục hồi tồn", "Chưa luồng hoàn tiền cổng tự động"),
        ("Sản phẩm lớn", "Có mô tả/ghi chú chung", "Chưa giao hàng/lắp đặt riêng"),
    ), (3.2, 7.1, 5.3), 10.2)
    table(doc, "3.16", "Mô hình mở rộng đề xuất, chưa có trong schema", ("Đối tượng", "Thuộc tính/quan hệ đề xuất", "Giá trị nghiệp vụ"), (
        ("ProductModel", "modelCode, brand, thông số kỹ thuật", "Phân biệt model cùng nhóm"),
        ("ProductSerial", "serial/IMEI, productId, orderDetailId, trạng thái", "Theo dõi từng thiết bị"),
        ("InventoryBatch", "lotCode, purchaseOrderItemId, ngày nhập", "Truy vết lô/giá vốn"),
        ("WarrantyProvider", "STORE/MANUFACTURER, trung tâm, SLA", "Phân tuyến bảo hành"),
        ("DeliveryOrder", "địa chỉ, hàng cồng kềnh, phí, lịch giao", "Theo dõi giao hàng"),
        ("InstallationJob", "kỹ thuật viên, lịch, kết quả", "Quản lý lắp đặt"),
        ("ProductAccessory", "sản phẩm chính-phụ kiện, số lượng", "Bán kèm và kiểm tra đủ bộ"),
    ), (3.2, 7.3, 5.1), 10.2)
    body_with_reference(doc, "Các đối tượng ở ", "bang_3_16", "Bảng 3.16", " là định hướng thiết kế, không được tính là chức năng đã hoàn thành. Việc tách rõ giúp báo cáo bám sát mã nguồn hiện tại đồng thời trả lời cách hệ thống có thể mở rộng cho nghiệp vụ đồ gia dụng sâu hơn.")

def chapter3_design(doc: Document) -> None:
    heading(doc, "3.12. Kiến trúc tổng thể", 2)
    body(doc, "Hệ thống áp dụng ba tầng logic: trình bày Next.js, dịch vụ Express và dữ liệu PostgreSQL. Trình duyệt quản lý trạng thái giao diện nhưng không nắm giữ khóa bí mật và không có quyền tự xác nhận giao dịch. Backend thực thi phân quyền, quy tắc, transaction và tích hợp ngoài; Prisma Client là cổng truy cập CSDL thống nhất.")
    picture(doc, "2_Sơ đồ/14_component_diagram.png", "3.6", "Sơ đồ thành phần hệ thống", max_height=18)
    body(doc, "Frontend gọi REST API; backend sử dụng Prisma, PayOS, Gemini, Groq và SMTP khi được cấu hình. Mỗi tích hợp ngoài được bao quanh bởi service hoặc route chuyên trách. Nếu mô hình AI lỗi, hệ thống dùng phương án dự phòng; nếu PayOS chưa xác nhận, đơn không được chuyển sang đã thanh toán.")

    heading(doc, "3.13. Thiết kế triển khai", 2)
    picture(doc, "2_Sơ đồ/15_deployment_diagram.png", "3.7", "Sơ đồ triển khai hệ thống", max_height=18)
    body(doc, "Môi trường phát triển dùng frontend cổng 3000 và backend cổng 5000. Khi triển khai, hai dịch vụ có thể tách riêng và giao tiếp HTTPS. Endpoint webhook PayOS phải công khai nhưng chỉ chấp nhận dữ liệu có chữ ký hợp lệ. Chuỗi kết nối, JWT secret, khóa PayOS và token mô hình được lưu bằng biến môi trường, không đưa xuống frontend.")

    heading(doc, "3.14. Thiết kế mô-đun", 2)
    table(doc, "3.17", "Ánh xạ mô-đun", ("Phân hệ", "Trang/Component", "API"), (
        ("Xác thực", "/, auth context, role guard", "/api/auth"),
        ("Bán hàng", "/pos, printable-invoice", "/api/orders, /payments, /pos"),
        ("Hàng hóa", "/products, /categories, /suppliers", "/api/products, /categories, /suppliers"),
        ("Kho", "/inventory, /purchase-orders", "/api/inventory, /purchase-orders"),
        ("Khách hàng", "/customers", "/api/customers"),
        ("Sau bán", "/warranties, /return-orders, /vat-invoices", "API tương ứng"),
        ("Vận hành", "/shifts, /settings, /audit-logs", "API tương ứng"),
        ("Báo cáo", "/dashboard, /reports", "/api/reports"),
    ), (3.0, 6.2, 6.2), 10.2)

    heading(doc, "3.15. Thiết kế cơ sở dữ liệu", 2)
    picture(doc, "2_Sơ đồ/13_erd_rut_gon.png", "3.8", "Sơ đồ thực thể - liên kết rút gọn", max_height=18.5)
    body(doc, "Thiết kế lược đồ quan hệ tuân theo nguyên tắc tách dữ liệu danh mục và giao dịch [20]. Schema hiện có 22 model. Bảng giao dịch tách khỏi danh mục để bảo toàn lịch sử. OrderDetail lưu unitPrice, lineTotal và unitCost tại thời điểm bán, nhờ đó hóa đơn và lợi nhuận không đổi khi giá sản phẩm được cập nhật. StockTransaction, AuditLog và PaymentWebhookLog tạo ba lớp dấu vết cho kho, thao tác người dùng và thanh toán ngoài.")
    table(doc, "3.18", "Nhóm bảng dữ liệu", ("Nhóm", "Bảng", "Mục đích"), (
        ("Người dùng", "Role, User", "Tài khoản và vai trò"),
        ("Hàng hóa", "Category, Supplier, Product", "Danh mục, nguồn hàng, giá và tồn"),
        ("Bán hàng", "Order, OrderDetail, Payment", "Đơn, dòng hàng và thanh toán"),
        ("Thanh toán ngoài", "PaymentWebhookLog", "Dấu vết, kết quả xử lý và idempotency"),
        ("Kho", "StockTransaction, PurchaseOrder, PurchaseOrderItem", "Biến động và phiếu nhập"),
        ("Khách hàng", "Customer, Warranty", "Thành viên và bảo hành"),
        ("Sau bán", "ReturnOrder, ReturnOrderItem, VatInvoiceRequest", "Hoàn trả và VAT"),
        ("Vận hành", "Shift, Setting, Promotion", "Ca, cấu hình và giảm giá"),
        ("Giám sát", "AuditLog, Notification", "Nhật ký và thông báo"),
    ), (2.7, 6.5, 6.2), 9.8)
    heading(doc, "3.15.1. Từ điển dữ liệu cốt lõi", 3)
    table(doc, "3.19", "Bảng Product", ("Thuộc tính", "Kiểu/ràng buộc", "Ý nghĩa"), (
        ("id", "Int, PK", "Định danh"), ("sku", "String, UNIQUE", "Mã quản lý"),
        ("categoryId, supplierId", "Int, FK", "Danh mục và nhà cung cấp"),
        ("costPrice, salePrice", "Decimal(12,2)", "Giá vốn và giá bán"),
        ("stockQuantity, minStock", "Int", "Tồn hiện tại và ngưỡng cảnh báo"),
        ("warrantyMonths", "Int", "Thời hạn bảo hành theo sản phẩm"),
        ("barcode, qrCode", "String, UNIQUE?", "Mã nhận dạng để quét"),
        ("status", "ProductStatus", "ACTIVE/INACTIVE"),
    ), (4.0, 4.5, 6.9), 10.2)
    table(doc, "3.20", "Nhóm bảng giao dịch Order", ("Bảng", "Thuộc tính quan trọng", "Ý nghĩa"), (
        ("Order", "orderCode, userId, customerId, totalAmount, shiftId, status", "Thông tin đầu đơn"),
        ("OrderDetail", "productId, quantity, unitPrice, unitCost, lineTotal", "Ảnh chụp dòng hàng"),
        ("Payment", "method, amount, status, providerOrderCode, paidAt", "Thanh toán"),
        ("PaymentWebhookLog", "eventId, paymentId, orderCode, payload, status, errorMessage", "Dấu vết webhook; chữ ký được SDK xác minh nhưng không lưu thành cột riêng"),
        ("StockTransaction", "productId, orderId, type, quantity", "Biến động tồn"),
        ("Warranty", "warrantyCode, orderDetailId, startDate, endDate", "Bảo hành theo dòng"),
    ), (3.0, 7.8, 4.6), 9.8)

    heading(doc, "3.16. Thiết kế lớp nghiệp vụ", 2)
    picture(doc, "2_Sơ đồ/11_class_nghiep_vu_cot_loi.png", "3.9", "Sơ đồ lớp nghiệp vụ cốt lõi", max_height=18)
    body(doc, "Role-User-Order-OrderDetail-Product tạo trục giao dịch. Một người dùng lập nhiều đơn; một đơn có nhiều dòng; mỗi dòng tham chiếu một sản phẩm. Payment, Customer và Warranty mở rộng vòng đời trước, trong và sau bán.")
    picture(doc, "2_Sơ đồ/12_class_kho_nhap_hang_hoan_tra.png", "3.10", "Sơ đồ lớp kho, nhập hàng và hoàn trả", max_height=18)
    body(doc, "PurchaseOrderItem ghi số lượng và giá vốn nhập; ReturnOrderItem tham chiếu OrderDetail để kiểm soát lượng đã mua và đã hoàn. StockTransaction là sổ biến động dùng chung, không thay thế chứng từ nguồn.")

    heading(doc, "3.17. Thiết kế tuần tự nghiệp vụ", 2)
    diagrams = (
        ("3.17.1. Checkout và tạo hóa đơn", "2_Sơ đồ/08_sequence_checkout_tao_hoa_don.png", "3.11", "Checkout và tạo hóa đơn"),
        ("3.17.2. Thanh toán PayOS", "2_Sơ đồ/05_2_sequence_thanh_toan_payos.png", "3.12", "Thanh toán PayOS"),
        ("3.17.3. Nhập kho", "2_Sơ đồ/09_sequence_nhap_kho.png", "3.13", "Tạo phiếu nhập kho"),
        ("3.17.4. Cấp bảo hành", "2_Sơ đồ/10_sequence_tao_bao_hanh_tu_dong.png", "3.14", "Tạo bảo hành tự động"),
        ("3.17.5. Trợ lý bán hàng", "2_Sơ đồ/05_4_sequence_ai_goi_y_ban_hang.png", "3.15", "AI gợi ý bán hàng"),
        ("3.17.6. Báo cáo doanh thu", "2_Sơ đồ/05_6_sequence_bao_cao_doanh_thu.png", "3.16", "Báo cáo doanh thu"),
        ("3.17.7. Cảnh báo tồn", "2_Sơ đồ/05_7_sequence_canh_bao_ton_kho.png", "3.17", "Cảnh báo tồn kho"),
    )
    for title, path, number, caption in diagrams:
        heading(doc, title, 3)
        picture(doc, path, number, caption, max_height=17)
        if number == "3.11":
            body(doc, "Backend tải lại sản phẩm và tính tổng thay vì tin giá client. Trong transaction, hệ thống cập nhật Order, Payment, Product, StockTransaction, Customer và Warranty; chỉ trả kết quả in hóa đơn sau khi transaction thành công.")
        elif number == "3.12":
            body(doc, "Việc frontend hiển thị QR hoặc trang return thành công không đồng nghĩa đã thanh toán. Chỉ webhook có chữ ký hợp lệ hoặc truy vấn PayOS xác nhận trạng thái PAID mới kích hoạt hoàn tất đơn.")
        elif number == "3.15":
            body(doc, "Backend giới hạn dữ liệu gửi mô hình, hậu kiểm ID, trạng thái, tồn và ngân sách. Khi thêm gợi ý, frontend vẫn áp dụng cùng kiểm tra như sản phẩm chọn trực tiếp.")

    heading(doc, "3.18. Thiết kế REST API", 2)
    table(doc, "3.21", "Endpoint đại diện", ("Method", "Endpoint", "Chức năng", "Quyền"), (
        ("POST", "/api/auth/login", "Đăng nhập", "Công khai"),
        ("GET/POST", "/api/products", "Danh sách/tạo sản phẩm", "Nội bộ/Admin ghi"),
        ("POST", "/api/orders/draft", "Tạo đơn nháp", "ADMIN, CASHIER"),
        ("PATCH", "/api/orders/:id/checkout", "Checkout", "ADMIN, CASHIER"),
        ("POST", "/api/payments/payos/create", "Tạo QR PayOS", "ADMIN, CASHIER"),
        ("POST", "/api/payments/webhook/payos", "Nhận và xác minh webhook", "PayOS/chữ ký"),
        ("POST", "/api/purchase-orders", "Tạo phiếu nhập", "ADMIN"),
        ("POST", "/api/return-orders", "Tạo phiếu hoàn", "ADMIN"),
        ("GET", "/api/reports/*", "Báo cáo", "ADMIN"),
        ("GET", "/api/warranties/public/*", "Tra cứu bảo hành", "Công khai"),
        ("POST/GET", "/api/pos/remote-scan*", "Mã quét từ xa", "Theo phiên"),
        ("POST", "/api/pos/sales-assistant", "Gợi ý bán hàng", "Nội bộ"),
    ), (2.0, 5.5, 5.2, 2.7), 9.5)

    heading(doc, "3.19. Thiết kế an toàn PayOS và webhook", 2)
    body(doc, "PayOS là nguồn xác nhận thanh toán ngoài hệ thống. Thiết kế áp dụng nguyên tắc không tin trạng thái từ frontend: return URL, nút đóng modal, nội dung QR và dữ liệu do trình duyệt gửi chỉ dùng để điều hướng hoặc yêu cầu backend kiểm tra; chúng không được phép đặt Payment thành PAID.")
    table(doc, "3.22", "Kiểm soát an toàn webhook PayOS", ("Kiểm soát", "Cách thực hiện trong HomeX POS", "Rủi ro được xử lý"), (
        ("Xác minh chữ ký", "payOS.webhooks.verify(req.body) bằng checksum key phía server", "Payload giả mạo hoặc bị sửa"),
        ("Chống replay", "Khóa Payment trong transaction và nhận diện trạng thái PAID/sự kiện đã PROCESSED", "Gửi lại thông báo cũ"),
        ("Idempotency", "Mỗi lần nhận ghi PaymentWebhookLog; bản lặp trả DUPLICATE", "Trừ kho, cộng điểm, cấp bảo hành hai lần"),
        ("Đối chiếu số tiền", "So sánh data.amount với Payment.amount trước hoàn tất", "Thanh toán thiếu hoặc nhầm số tiền"),
        ("Đối chiếu mã đơn", "So sánh orderCode đã xác minh với providerOrderCode", "Gắn giao dịch vào đơn khác"),
        ("Khóa đồng thời", "SELECT ... FOR UPDATE trên bản ghi Payment", "Hai webhook hợp lệ chạy đồng thời"),
        ("Transaction", "Payment, Order, kho, điểm, bảo hành và log cùng giao dịch CSDL", "Trạng thái cập nhật dở dang"),
        ("Nhật ký", "Lưu PROCESSED, DUPLICATE, FAILED hoặc UNMATCHED cùng payload", "Thiếu dấu vết đối soát"),
    ), (3.1, 8.2, 4.1), 8.9)
    heading(doc, "3.19.1. Xử lý thứ tự sự kiện bất định", 3)
    table(doc, "3.23", "Tình huống thời gian của thanh toán", ("Tình huống", "Cách xử lý", "Kết quả nhất quán"), (
        ("Webhook đến trước trang return", "Backend xác minh và hoàn tất; frontend mở lại sẽ đọc trạng thái từ API", "Đơn PAID, không phụ thuộc trang return"),
        ("Trang return đến trước webhook", "Frontend yêu cầu backend đồng bộ trạng thái PayOS; nếu chưa PAID tiếp tục chờ", "Không xác nhận sớm"),
        ("Client đóng trình duyệt", "Webhook server-to-server vẫn hoàn tất transaction", "Kho, điểm và bảo hành vẫn được cập nhật"),
        ("Webhook gửi nhiều lần", "Lần đầu PROCESSED; các lần sau DUPLICATE", "Không phát sinh tác dụng phụ lần hai"),
        ("Mã đơn/số tiền sai", "Ghi FAILED hoặc UNMATCHED, không hoàn tất", "Dữ liệu nội bộ không bị thay đổi"),
        ("Webhook và đồng bộ chủ động cùng lúc", "Cùng đi qua hàm hoàn tất và kiểm tra trạng thái trong transaction", "Một kết quả PAID duy nhất"),
    ), (3.8, 7.4, 4.2), 9.0)
    body(doc, "Benchmark tích hợp cục bộ tạo chữ ký HMAC theo thuật toán SDK, thử payload hợp lệ và payload bị sửa, sau đó gửi hai webhook hợp lệ đồng thời. Kết quả đạt 10/10 điều kiện: chữ ký sửa bị từ chối, chỉ một log PROCESSED và một DUPLICATE, tồn kho chỉ giảm một lần, chỉ có một giao dịch SALE, đơn/thanh toán hoàn tất và bảo hành không bị tạo lặp.")

    heading(doc, "3.20. Thiết kế các chức năng hỗ trợ thông minh", 2)
    body(doc, "HomeX POS tách ba chức năng theo nguồn quyết định. Gợi ý bán hàng dùng LLM trên tập ứng viên đã kiểm soát; nhập hàng dùng công thức xác định và chỉ nhờ LLM diễn giải; làm giàu sản phẩm kết hợp API/ảnh nhưng không tự lưu. Phiên bản hiện tại không đọc ảnh hóa đơn và không huấn luyện mô hình dự báo.")
    table(doc, "3.24", "Phạm vi ba chức năng hỗ trợ", ("Chức năng", "Đầu vào", "Đầu ra", "Quyền quyết định cuối"), (
        ("Gợi ý bán hàng", "Nhu cầu, ngân sách, hàng ACTIVE còn tồn", "Tối đa 5 ID và lý do", "Thu ngân chọn có thêm giỏ hay không"),
        ("Ước lượng nhập hàng", "Tồn, minStock, bán 7/30 ngày, danh mục, tháng", "Nhu cầu và lượng đề xuất theo luật", "Admin duyệt/tạo PurchaseOrder"),
        ("Làm giàu sản phẩm", "Barcode, API ngoài; ảnh khi thiếu nhận dạng", "Dữ liệu điền trước và missingFields", "Admin sửa/duyệt rồi mới lưu"),
    ), (3.4, 5.0, 4.2, 2.8), 8.8)

    heading(doc, "3.20.1. Luồng ước lượng nhu cầu", 3)
    table(doc, "3.25", "Thiết kế bộ ước lượng nhập hàng", ("Bước", "Xử lý", "Kiểm soát/giới hạn"), (
        ("1. Lấy lịch sử", "OrderDetail của đơn COMPLETED/PAID trong 30 ngày", "Loại CANCELLED/REFUNDED; chưa trừ hoàn một phần"),
        ("2. Tổng hợp", "Bán 7/30 ngày theo SKU và danh mục", "Dùng createdAt; không phải tập training"),
        ("3. Tính nền", "avg30, avg7, trendRatio, categoryTrendRatio", "Dữ liệu thưa làm tỷ lệ biến động mạnh"),
        ("4. Mùa vụ", "Hệ số 1,10–1,35 theo tháng/danh mục", "Luật chuyên gia, chưa học hoặc backtest nhiều năm"),
        ("5. Tồn an toàn", "20% nhu cầu; so với currentStock và minStock", "Tỷ lệ cố định, chưa tối ưu chi phí thiếu/tồn"),
        ("6. LLM", "Diễn giải priority, reason, action plan", "Không được tạo SKU hoặc thay kết quả nền tùy ý"),
        ("7. Phê duyệt", "Admin sửa số lượng/chọn NCC và tạo PO", "Không tự tăng tồn; nhận PO mới ghi IMPORT"),
    ), (2.2, 7.7, 5.5), 8.7)
    body(doc, "Sản phẩm không có bán trong cả hai cửa sổ nhận confidence LOW. Nếu tồn thấp hơn minStock, hệ thống chỉ đề xuất nhập phần còn thiếu để đủ minStock; nếu tồn đã đủ thì suggestedRestockQuantity bằng 0. Admin vẫn phải xem kế hoạch trưng bày và thông tin nhà cung cấp. Promotion trong Order chỉ phản ánh giảm giá đã xảy ra; endpoint chưa nhận lịch khuyến mãi tương lai nên không điều chỉnh uplift.")

    heading(doc, "3.20.2. Thiết kế backtest", 3)
    table(doc, "3.26", "Thiết kế đánh giá định lượng", ("Thành phần", "Dữ liệu demo", "Tập mô phỏng có kiểm soát"), (
        ("Quy mô", "Bản lưu 29/07: 11 đơn, 34 dòng, 4 SKU, 5 ngày bán/16 ngày lịch", "180 ngày, 8 SKU, seed cố định"),
        ("Chia thời gian", "7 ngày đầu; rolling-origin 1 ngày", "30 ngày đầu; rolling-origin horizon 7 ngày"),
        ("Baseline", "Trung bình trượt 7 và 30 ngày", "Cùng hai baseline"),
        ("Kịch bản", "Nhu cầu demo rất thưa; không có promotion", "Ổn định, tăng/giảm, mùa vụ, promotion, intermittent, sản phẩm mới"),
        ("Chỉ số", "MAE, RMSE, WAPE, MAPE khi actual > 0, bias", "Cùng chỉ số và phân tích theo lát cắt"),
        ("Mục đích", "Đánh giá thăm dò, không tổng quát hóa", "Kiểm tra hành vi kỹ thuật, không thay dữ liệu thật"),
    ), (3.1, 6.1, 6.2), 8.7)

    heading(doc, "3.20.3. Thiết kế độ tin cậy của barcode và ảnh", 3)
    table(doc, "3.27", "Ranh giới tin cậy khi làm giàu sản phẩm", ("Nguồn/bước", "Quyết định được phép", "Không được phép"), (
        ("CSDL nội bộ", "Trả sản phẩm hiện có và cảnh báo trùng", "Tạo bản sao cùng barcode"),
        ("API barcode", "Điền trường có nguồn vào ô trống", "Coi ảnh/giá bên thứ ba là quyền sở hữu của HomeX"),
        ("Hợp nhất", "Chuẩn hóa category/supplier và báo missingFields", "Khẳng định HYBRID là đúng khi nguồn xung đột"),
        ("Vision/LLM", "Gợi ý name/category/description nếu có căn cứ", "Ghi đè danh tính tin cậy hoặc tự lưu"),
        ("Frontend", "Hiển thị source, cho Admin sửa và bấm Lưu", "Tự submit sau khi enrich"),
    ), (3.3, 6.4, 5.7), 8.8)
    body(doc, "Service chỉ sử dụng các nguồn API được tài liệu hóa trong phiên bản này. Kết quả từ API ngoài có thể thiếu, sai hoặc chịu giấy phép riêng. Trọng số confidence hiện là quy ước nội bộ, không phải xác suất hiệu chuẩn; nếu triển khai thương mại cần lưu provenance theo từng field, ghi attribution và quản lý thời hạn/quyền lưu dữ liệu.")

    heading(doc, "3.20.4. Hậu kiểm và phương án dự phòng", 3)
    table(doc, "3.28", "Rào chắn đầu ra", ("Nhóm", "Hậu kiểm", "Fallback/giới hạn"), (
        ("Bán hàng", "ID thuộc tập ứng viên; còn hàng; không vượt ngân sách; loại trùng", "Chấm điểm heuristic khi dịch vụ lỗi"),
        ("Nhập hàng", "Zod kiểm tra JSON; SKU có thật; số không âm; không bán chỉ nhập bù minStock khi thiếu", "Công thức vẫn chạy khi thiếu token; không tự phê duyệt"),
        ("Barcode", "Ưu tiên CSDL; sanitize AI với external; chỉ điền ô trống", "Không tìm thấy thì trả missingFields để nhập tay"),
        ("Ảnh", "Tên phải đọc được; category trong taxonomy; trường không chắc để trống", "Độ chính xác ảnh còn thấp; Admin duyệt"),
    ), (3.2, 7.5, 4.7), 9.0)
    body(doc, "Đánh giá ảnh dùng 12 ảnh sản phẩm hiện có: đúng tên cụ thể 33,3%, đúng danh mục 50,0% và không điền các trường không quan sát được đạt 100%. Kết quả cho thấy ảnh chỉ phù hợp nhận dạng sơ bộ, không đủ để tự động tạo hồ sơ sản phẩm.")

    heading(doc, "3.21. Thiết kế bảo mật", 2)
    table(doc, "3.29", "Kiểm soát bảo mật theo lớp", ("Lớp", "Cơ chế", "Mục tiêu"), (
        ("Trình duyệt", "Role guard, ẩn menu, xóa phiên khi 401", "Giảm thao tác sai; không thay thế bảo vệ API"),
        ("API", "JWT, authorizeRoles, Zod, error middleware", "Xác thực, phân quyền và đầu vào"),
        ("Mật khẩu", "bcrypt", "Không lưu mật khẩu rõ"),
        ("Dữ liệu", "Prisma, FK, unique, transaction, row lock", "Toàn vẹn và chống cạnh tranh"),
        ("Tích hợp", "Biến môi trường, xác minh webhook, hậu kiểm AI", "Bảo vệ khóa và giao dịch"),
        ("Giám sát", "AuditLog, PaymentWebhookLog", "Điều tra và đối soát"),
    ), (2.8, 7.0, 5.6), 10.2)

    heading(doc, "3.22. Nguyên tắc thiết kế giao diện", 2)
    bullets(doc, (
        "Điều hướng theo vai trò; thao tác chính và trạng thái trang nhất quán.",
        "Bảng có tìm kiếm, lọc, phân trang, trạng thái và menu thao tác.",
        "POS ưu tiên mật độ thông tin, giỏ cố định, tổng tiền rõ và quét mã.",
        "Có xác nhận tác vụ quan trọng; loading, empty, error và toast đầy đủ.",
        "Quản lý tập trung ngôn ngữ VI/EN và định dạng ngày dd/mm/yyyy.",
        "Kết quả AI và thanh toán được trình bày như dữ liệu cần kiểm chứng, không thay thế quyết định nghiệp vụ phía server.",
    ))

def chapter4(doc: Document) -> None:
    heading(doc, "CHƯƠNG 4. XÂY DỰNG HỆ THỐNG", 1, True)
    heading(doc, "4.1. Môi trường và cấu trúc mã nguồn", 2)
    table(doc, "4.1", "Môi trường phát triển", ("Hạng mục", "Cấu hình"), (
        ("Hệ điều hành", "Windows, PowerShell"), ("Runtime", "Node.js"),
        ("Ngôn ngữ", "TypeScript/TSX"), ("CSDL", "PostgreSQL"),
        ("Quản lý gói", "npm"), ("ORM/migration", "Prisma"),
        ("Kiểm tra", "TypeScript compiler, Next.js production build, kiểm thử chấp nhận"),
    ), (5.0, 10.4))
    body(doc, "Dự án tách frontend và backend. Backend có routes, services, middlewares, lib và prisma. Frontend dùng App Router với nhóm dashboard/public, invoice, payment, mobile-scan; component dùng chung đặt trong components. Cấu trúc hỗ trợ phát triển độc lập và giảm trùng lặp.")
    table(doc, "4.2", "Quy mô hiện trạng", ("Chỉ số", "Giá trị", "Ý nghĩa"), (
        ("Model Prisma", "22", "Dữ liệu nền, giao dịch, giám sát"),
        ("REST handler nghiệp vụ", "117 khai báo HTTP trong router", "Ngoài 3 endpoint hệ thống/health"),
        ("Trang Next.js", "28", "Nội bộ, công khai và thanh toán"),
        ("Kho ảnh giao diện", "74", "68 Admin, 6 Cashier; chọn 45 ảnh tiêu biểu"),
        ("Sơ đồ", "17", "UML, ERD, component, deployment"),
    ), (5.2, 2.5, 7.7))
    groups = (
        ("4.2. Xác thực và dashboard", (
            ("Admin/00.1_dang_nhap_admin.jpg", "4.1", "Giao diện đăng nhập HomeX POS", 15.7),
            ("Admin/01.1_tong_quan_dashboard.jpg", "4.2", "Dashboard tổng quan dành cho Admin", 15.7),
            ("Cashier/01.1_tong_quan_dashboard.jpg", "4.3", "Dashboard giới hạn dành cho Cashier", 15.7),
        ), "Form đăng nhập kiểm tra email/mật khẩu; backend xác minh bcrypt, trạng thái ACTIVE và ký JWT. Dashboard được điều chỉnh theo vai trò, trong khi API vẫn kiểm tra quyền phía server."),
        ("4.3. Quy trình bán hàng tại quầy", (
            ("Admin/02.1_pos_giao_dien_chinh.jpg", "4.4", "Giao diện chính POS", 15.7),
            ("Admin/02.2_pos_ket_noi_may_quet.jpg", "4.5", "Kết nối điện thoại làm máy quét", 9.5),
            ("Admin/02.3_pos_them_san_pham.jpg", "4.6", "Thêm sản phẩm vào giỏ", 15.7),
            ("Admin/02.4_pos_tro_ly_ai_goi_y.jpg", "4.7", "Trợ lý AI gợi ý bán hàng", 14.0),
            ("Admin/02.5_pos_xac_nhan_don_hang.jpg", "4.8", "Xác nhận đơn trước thanh toán", 12.0),
            ("Admin/02.6_pos_thanh_toan_payos.jpg", "4.9", "Thanh toán QR PayOS", 12.0),
            ("Admin/02.7_pos_thanh_toan_tien_mat.jpg", "4.10", "Thanh toán tiền mặt", 10.5),
            ("Admin/02.8_pos_hoa_don_hoan_thanh.jpg", "4.11", "Hóa đơn hoàn tất", 11.0),
        ), "POS bố trí sản phẩm và giỏ hàng trên cùng màn hình, hỗ trợ tìm, lọc, máy quét, khách hàng, khuyến mãi và trạng thái ca. Trợ lý AI chỉ đề xuất; người dùng quyết định thêm sản phẩm. Checkout tiền mặt tính tiền thừa; PayOS chờ xác nhận trước khi hoàn tất."),
        ("4.4. Đơn hàng và hoàn trả", (
            ("Admin/03.1_don_hang_danh_sach.jpg", "4.12", "Danh sách đơn hàng", 15.7),
            ("Admin/03.2_don_hang_chi_tiet.jpg", "4.13", "Chi tiết đơn và thanh toán", 15.7),
            ("Admin/03.6_don_hang_tao_phieu_tra_hang.jpg", "4.14", "Tạo phiếu trả hàng", 15.7),
        ), "Danh sách hỗ trợ tìm/lọc và xem người bán, khách hàng, dòng hàng, giảm giá, Payment. Hoàn trả tạo chứng từ riêng, kiểm tra số lượng và ghi RESTORE thay vì xóa lịch sử đơn gốc."),
        ("4.5. Khách hàng và bảo hành", (
            ("Admin/04.1_khach_hang_danh_sach.jpg", "4.15", "Khách hàng và hạng thành viên", 15.7),
            ("Admin/05.1_bao_hanh_danh_sach.jpg", "4.16", "Quản lý bảo hành nội bộ", 15.7),
            ("Admin/05.4_bao_hanh_tra_cuu_cong_khai.jpg", "4.17", "Tra cứu bảo hành công khai", 15.7),
            ("Admin/05.5_bao_hanh_ket_qua_tra_cuu.jpg", "4.18", "Kết quả tra cứu bảo hành", 15.7),
        ), "Khách hàng được định danh bằng số điện thoại, có điểm, hạng và lịch sử đơn hàng. Trang bảo hành công khai chỉ trả dữ liệu cần thiết; thay đổi trạng thái vẫn thuộc API nội bộ có quyền."),
    )
    for title, images, description in groups:
        heading(doc, title, 2)
        body(doc, description)
        for relative, number, caption, width in images:
            picture(doc, f"1_UI website/{relative}", number, caption, max_width=width, max_height=16)
    more_groups = (
        ("4.6. Sản phẩm và mã vạch", (
            ("Admin/07.1_san_pham_danh_sach.jpg", "4.19", "Danh sách và bộ lọc sản phẩm", 15.7),
            ("Admin/07.5_san_pham_them_moi.jpg", "4.20", "Biểu mẫu thêm sản phẩm", 12.0),
            ("Admin/07.4_san_pham_in_ma_vach.jpg", "4.21", "In nhãn mã vạch CODE128", 9.0),
        ), "Biểu mẫu kiểm tra SKU, giá, tồn, danh mục, nhà cung cấp và bảo hành. Sản phẩm hỗ trợ xóa mềm/khôi phục, import dữ liệu, barcode và ảnh; backend không cho trùng mã."),
        ("4.7. Kho và nhập hàng", (
            ("Admin/10.1_kho_hang_tong_quan.jpg", "4.22", "Tổng quan tồn kho", 15.7),
            ("Admin/10.2_kho_hang_tro_ly_ai.jpg", "4.23", "Trợ lý phân tích nhập hàng", 15.7),
            ("Admin/10.5_kho_hang_ai_khuyen_nghi_ton.jpg", "4.24", "Khuyến nghị số lượng nhập", 15.7),
            ("Admin/10.8_kho_hang_phieu_nhap.jpg", "4.25", "Tạo phiếu nhập từ đề xuất", 13.5),
            ("Admin/10.10_kho_hang_dieu_chinh_ton.jpg", "4.26", "Điều chỉnh tồn có ghi chú", 10.5),
        ), "Phân tích nhập hàng kết hợp tồn, minStock, tốc độ bán 7/30 ngày và kỳ dự phòng. Đây là công thức theo luật; LLM chỉ diễn giải. Admin duyệt đề xuất rồi tạo PurchaseOrder. Nhập/điều chỉnh đều ghi StockTransaction để truy vết."),
        ("4.8. Khuyến mãi, VAT và ca", (
            ("Admin/12.1_khuyen_mai_danh_sach.jpg", "4.27", "Danh sách khuyến mãi", 15.7),
            ("Admin/11.3_thanh_toan_vat_danh_sach.jpg", "4.28", "Danh sách yêu cầu VAT", 15.7),
            ("Admin/06.1_ca_lam_viec_danh_sach.jpg", "4.29", "Đối soát ca làm việc", 15.7),
        ), "Khuyến mãi có loại giảm, mức tối đa, đơn tối thiểu, thời gian và giới hạn dùng. VAT gắn với Order và có luồng duyệt/từ chối/email. Ca lưu tiền đầu, tiền dự kiến, thực tế và chênh lệch."),
        ("4.9. Báo cáo và giám sát", (
            ("Admin/13.1_bao_cao_loi_nhuan.jpg", "4.30", "Báo cáo doanh thu và lợi nhuận", 15.7),
            ("Admin/14.1_nguoi_dung_danh_sach.jpg", "4.31", "Quản lý người dùng", 15.7),
            ("Admin/15.1_lich_su_he_thong_audit.jpg", "4.32", "Nhật ký hệ thống", 15.7),
            ("Admin/16.3_cai_dat_cau_hinh_he_thong.jpg", "4.33", "Cấu hình hệ thống", 15.7),
        ), "Báo cáo tổng hợp doanh thu, giá vốn, lợi nhuận và top sản phẩm theo khoảng ngày. Người dùng, audit và settings giúp Admin kiểm soát quyền, truy vết và điều chỉnh quy tắc vận hành mà không sửa mã."),
    )
    for title, images, description in more_groups:
        heading(doc, title, 2)
        body(doc, description)
        for relative, number, caption, width in images:
            picture(doc, f"1_UI website/{relative}", number, caption, max_width=width, max_height=16)
    heading(doc, "4.9.1. Các giao diện quản trị bổ sung", 3)
    body(doc, "Các màn hình dưới đây hoàn thiện chuỗi dữ liệu nền và vận hành. Danh mục, nhà cung cấp và khuyến mãi cung cấp dữ liệu cho POS; VAT và cài đặt điều khiển quy trình sau bán, in ấn, máy quét và các giới hạn nghiệp vụ.")
    extra_admin = (
        ("Admin/08.1_danh_muc_danh_sach.jpg", "4.34", "Quản lý danh mục sản phẩm"),
        ("Admin/09.1_nha_cung_cap_danh_sach.jpg", "4.35", "Quản lý nhà cung cấp"),
        ("Admin/09.4_nha_cung_cap_lich_su_nhap_hang.jpg", "4.36", "Lịch sử nhập hàng theo nhà cung cấp"),
        ("Admin/11.1_thanh_toan_vat_yeu_cau.jpg", "4.37", "Tiếp nhận yêu cầu hóa đơn VAT"),
        ("Admin/11.4_thanh_toan_vat_chi_tiet.jpg", "4.38", "Chi tiết và xử lý yêu cầu VAT"),
        ("Admin/12.2_khuyen_mai_them_moi.jpg", "4.39", "Thiết lập điều kiện khuyến mãi"),
        ("Admin/16.1_cai_dat_van_hanh_chung.jpg", "4.40", "Cài đặt vận hành chung"),
        ("Admin/16.2_cai_dat_van_hanh_may_quet.jpg", "4.41", "Cài đặt máy quét và màn hình POS"),
    )
    for relative, number, caption in extra_admin:
        picture(doc, f"1_UI website/{relative}", number, caption, max_height=15.5)
    heading(doc, "4.9.2. Đối chiếu giao diện Cashier", 3)
    body(doc, "Cashier sử dụng cùng dữ liệu giao dịch nhưng phạm vi điều hướng được rút gọn. Giao diện giữ nhất quán với Admin ở POS, đơn hàng và khách hàng, nhờ đó giảm thời gian học; các trang cấu hình, kho, người dùng và báo cáo nhạy cảm không xuất hiện.")
    picture(doc, "1_UI website/Cashier/02.1_pos_ban_hang.jpg", "4.42", "Màn hình bán hàng của Cashier", max_height=15.5)
    picture(doc, "1_UI website/Cashier/06.1_ca_lam_viec_thu_ngan.jpg", "4.43", "Màn hình ca làm việc của Cashier", max_height=15.5)
    heading(doc, "4.9.3. Chứng từ và thông tin hậu mãi", 3)
    body(doc, "Hóa đơn in và màn hình chi tiết bảo hành là hai đầu ra quan trọng sau checkout. Chứng từ hiển thị đúng dòng hàng, tổng tiền và thông tin cửa hàng; hồ sơ bảo hành liên kết ngược về đơn, khách hàng và sản phẩm để nhân viên xác minh nhanh khi tiếp nhận.")
    picture(doc, "1_UI website/Admin/03.4_don_hang_in_hoa_don.jpg", "4.44", "Bản in hóa đơn bán hàng", max_width=11.5, max_height=15.5)
    picture(doc, "1_UI website/Admin/05.3_bao_hanh_chi_tiet.jpg", "4.45", "Chi tiết hồ sơ bảo hành", max_height=15.5)
    heading(doc, "4.10. Triển khai và cấu hình", 2)
    body(doc, "Backend cần DATABASE_URL và JWT_SECRET; PayOS, Gemini, Groq và SMTP là các tích hợp cấu hình bằng biến môi trường. Frontend cần NEXT_PUBLIC_API_URL. Prisma migration tạo schema và seed tạo dữ liệu mẫu. Mặc định backend chạy http://localhost:5000, frontend http://localhost:3000.")
    table(doc, "4.3", "Biến môi trường", ("Biến", "Thành phần", "Mục đích"), (
        ("DATABASE_URL", "Backend", "Kết nối PostgreSQL"), ("JWT_SECRET", "Backend", "Khóa ký JWT"),
        ("PAYOS_CLIENT_ID/API_KEY/CHECKSUM_KEY", "Backend", "Xác thực PayOS"),
        ("PAYOS_RETURN_URL/CANCEL_URL", "Backend", "Điều hướng thanh toán"),
        ("GEMINI_API_KEY/GEMINI_MODEL", "Backend", "Gemini chính; model mặc định gemini-flash-latest"),
        ("GROQ_API_KEY/GROQ_MODEL", "Backend", "Groq dự phòng; model chữ mặc định openai/gpt-oss-120b"),
        ("GROQ_VISION_MODEL", "Backend", "Model ảnh Groq mặc định qwen/qwen3.6-27b"),
        ("BARCODE_LOOKUP_API_KEY", "Backend", "Nguồn làm giàu barcode đang cấu hình"),
        ("UPCITEMDB_API_KEY/BARCODE_SPIDER_API_KEY", "Backend", "Nguồn barcode tùy chọn theo thuê bao"),
        ("NEXT_PUBLIC_API_URL", "Frontend", "Base URL API"),
    ), (6.1, 3.0, 6.3), 10.4)


def chapter5(doc: Document) -> None:
    heading(doc, "CHƯƠNG 5. KIỂM THỬ VÀ ĐÁNH GIÁ", 1, True)
    heading(doc, "5.1. Chiến lược, môi trường và mức bằng chứng", 2)
    body(doc, "Kiểm thử được tổ chức theo rủi ro nghiệp vụ thay vì chỉ chọn một vài màn hình thuận lợi. Các nhóm ưu tiên gồm thanh toán/webhook, checkout, kho, nhập hàng, bảo hành, hoàn trả và AI. Mỗi kết luận ghi rõ loại bằng chứng để không đồng nhất việc đọc mã với việc đã kiểm thử tự động hay đã vận hành thực tế.")
    table(doc, "5.1", "Phân loại bằng chứng kiểm thử", ("Mức", "Loại bằng chứng", "Ý nghĩa"), (
        ("A", "Benchmark/tập lệnh tái lập", "Đã chạy tự động với dữ liệu kiểm thử và lưu JSON kết quả"),
        ("B", "Kiểm tra mã và ràng buộc", "Đã đối chiếu route/service/schema/transaction; chưa thay thế test tải thực tế"),
        ("C", "Kiểm thử chấp nhận/giao diện", "Đã thực hiện theo luồng và có ảnh trạng thái hoặc dữ liệu demo"),
        ("D", "Chưa đủ bằng chứng thực tế", "Cần sandbox/production, người dùng hoặc bộ dữ liệu lớn hơn"),
    ), (1.5, 5.4, 8.5), 9.6)
    table(doc, "5.2", "Môi trường và dữ liệu đánh giá", ("Hạng mục", "Thiết lập"), (
        ("Backend", "Node.js, Express, TypeScript, Prisma; PostgreSQL dữ liệu demo"),
        ("Frontend", "Next.js 16.2.6 production build"),
        ("Dữ liệu hiện tại 11/08", "34 sản phẩm ACTIVE, 8 danh mục ACTIVE; 33 barcode và 34 ảnh"),
        ("Lịch sử 30 ngày hiện tại", "36 đơn hợp lệ, 84 dòng/103 đơn vị, 14 SKU, 9 ngày có bán; gồm 25 đơn demo tháng 8"),
        ("Backtest lưu ngày 29/07", "Tập cũ: 11 đơn; rolling-origin; 36 product-day; không đại diện dữ liệu hiện tại"),
        ("Tập kiểm soát", "180 ngày, 8 SKU, seed 23052084; horizon 7 ngày; 168 quan sát"),
        ("AI ảnh", "12 ảnh; đánh giá tên, danh mục và không bịa trường không quan sát được"),
        ("POS cục bộ 11/08", "8 tình huống, 56/56 ràng buộc; khóa AI được đặt rỗng, không gửi dữ liệu ra ngoài"),
        ("PayOS webhook", "Fixture HMAC, payload sửa và hai request hợp lệ đồng thời; tự hoàn nguyên"),
    ), (4.2, 11.2), 9.1)

    heading(doc, "5.2. Kiểm thử xác thực, phân quyền và bán hàng", 2)
    table(doc, "5.3", "Test case xác thực và bán hàng cốt lõi", ("TC", "Kịch bản", "Kết quả mong đợi", "Kết quả/bằng chứng"), (
        ("AUTH-01", "Đăng nhập đúng", "Trả JWT và thông tin vai trò", "Đạt (B,C)"),
        ("AUTH-02", "Sai mật khẩu/tài khoản inactive", "Từ chối, không cấp phiên", "Đạt (B,C)"),
        ("AUTH-03", "Cashier gọi API chỉ dành Admin", "API trả lỗi phân quyền", "Đạt (B)"),
        ("SALE-01", "Tạo đơn với sản phẩm hợp lệ", "Backend nạp lại giá và tồn", "Đạt (B,C)"),
        ("SALE-02", "Client sửa giá thấp hơn", "Giá client không được dùng làm nguồn tin", "Đạt (B)"),
        ("SALE-03", "Checkout khi chưa mở ca", "Yêu cầu mở ca", "Đạt (B,C)"),
        ("SALE-04", "Bán vượt tồn", "Từ chối nếu cấu hình không cho oversell", "Đạt (B)"),
        ("SALE-05", "Thanh toán tiền mặt", "PAID, trừ tồn, ghi SALE và tiền thừa", "Đạt (B,C)"),
        ("SALE-06", "Mã giảm giá hết hạn/không đủ điều kiện", "Từ chối và nêu lý do", "Đạt (B,C)"),
        ("SALE-07", "Checkout lỗi giữa transaction", "Không để đơn/kho/thanh toán cập nhật dở", "Đạt theo transaction (B)"),
    ), (1.7, 4.6, 6.0, 3.1), 9.5)

    heading(doc, "5.3. Kiểm thử PayOS và webhook", 2)
    body(doc, "Nhóm thanh toán được kiểm thử cả tính hợp lệ mật mã và tác dụng phụ nghiệp vụ. Benchmark không gọi thanh toán tiền thật; nó ký fixture bằng cùng thuật toán HMAC của SDK và gửi request vào endpoint cục bộ để kiểm chứng hành vi có thể tái lập.")
    table(doc, "5.4", "Test case PayOS/webhook", ("TC", "Điều kiện/kích thích", "Kết quả mong đợi", "Kết quả/bằng chứng"), (
        ("PAY-01", "Tạo link cho Payment hợp lệ", "Mã PayOS gắn đúng Payment", "Đạt theo mã (B)"),
        ("PAY-02", "Payload có chữ ký hợp lệ", "verify chấp nhận", "Đạt benchmark (A)"),
        ("PAY-03", "Sửa amount sau khi ký", "Chữ ký bị từ chối", "Đạt benchmark (A)"),
        ("PAY-04", "Sai amount nhưng chữ ký hợp lệ", "Ghi FAILED, không hoàn tất", "Đạt theo đối chiếu (B)"),
        ("PAY-05", "Sai providerOrderCode", "UNMATCHED/FAILED, không đổi đơn", "Đạt theo đối chiếu (B)"),
        ("PAY-06", "Hai webhook hợp lệ đồng thời", "Một PROCESSED, một DUPLICATE", "Đạt benchmark (A)"),
        ("PAY-07", "Webhook phát lại sau PAID", "Không trừ kho/cộng điểm/cấp bảo hành lần hai", "Đạt benchmark (A)"),
        ("PAY-08", "Webhook đến trước trang return", "Frontend đọc trạng thái đã hoàn tất từ backend", "Đạt theo thiết kế/mã (B)"),
        ("PAY-09", "Trang return đến trước webhook", "Chưa PAID thì chờ hoặc backend đồng bộ PayOS", "Đạt theo thiết kế/mã (B)"),
        ("PAY-10", "Client đóng trình duyệt", "Webhook server-to-server vẫn hoàn tất", "Đạt theo kiến trúc (B); cần sandbox (D)"),
        ("PAY-11", "Lỗi trong transaction hoàn tất", "Rollback toàn bộ tác dụng phụ", "Đạt theo transaction (B)"),
        ("PAY-12", "Webhook không tìm thấy Payment", "Ghi UNMATCHED để đối soát", "Đạt theo mã (B)"),
    ), (1.7, 4.8, 6.0, 2.9), 9.3)
    body(doc, "Kết quả benchmark PayOS đạt 10/10 điều kiện kiểm tra. Sau hai request đồng thời, hệ thống chỉ giảm tồn một lần, chỉ ghi một StockTransaction loại SALE, Order và Payment đều hoàn tất, Warranty tối đa một bản ghi; dữ liệu fixture được dọn và tồn kho được khôi phục sau phép thử.")

    heading(doc, "5.4. Kiểm thử nhập kho và quản lý tồn", 2)
    table(doc, "5.5", "Test case nhập kho và tồn", ("TC", "Kịch bản", "Kết quả mong đợi", "Kết quả/bằng chứng"), (
        ("INV-01", "Tạo phiếu nhập hợp lệ", "Lưu NCC, dòng hàng, giá vốn và tổng", "Đạt (B,C)"),
        ("INV-02", "Cashier tạo phiếu nhập", "API từ chối", "Đạt (B)"),
        ("INV-03", "Nhận phiếu nhập", "Tăng tồn và ghi IMPORT trong transaction", "Đạt (B,C)"),
        ("INV-04", "Nhận lại phiếu đã hoàn tất", "Không tăng tồn lần hai", "Đạt (B)"),
        ("INV-05", "Điều chỉnh âm vượt tồn", "Từ chối hoặc tuân theo rule oversell", "Đạt (B)"),
        ("INV-06", "Điều chỉnh có ghi chú", "Ghi ADJUSTMENT và người thao tác", "Đạt (B,C)"),
        ("INV-07", "Sản phẩm dưới minStock", "Xuất hiện cảnh báo tồn thấp", "Đạt (B,C)"),
        ("INV-08", "Bộ ước lượng tạo đề xuất", "Không tự làm tồn tăng; Admin phải duyệt PO", "Đạt rào chắn (A,B); sai số xem 5.6.2"),
    ), (1.7, 4.7, 6.0, 3.0), 9.4)

    heading(doc, "5.5. Kiểm thử bảo hành và hoàn trả", 2)
    table(doc, "5.6", "Test case bảo hành", ("TC", "Kịch bản", "Kết quả mong đợi", "Kết quả/bằng chứng"), (
        ("WAR-01", "Bán sản phẩm warrantyMonths > 0", "Tạo mã, ngày bắt đầu/kết thúc đúng", "Đạt (B,C)"),
        ("WAR-02", "Bán sản phẩm không bảo hành", "Không tạo hồ sơ không cần thiết", "Đạt (B)"),
        ("WAR-03", "Webhook bị lặp", "Không tạo Warranty lần hai", "Đạt benchmark (A)"),
        ("WAR-04", "Tra cứu mã đúng", "Chỉ trả dữ liệu công khai cần thiết", "Đạt (B,C)"),
        ("WAR-05", "Tra cứu mã sai", "Không tìm thấy, không lộ hồ sơ khác", "Đạt (B,C)"),
        ("WAR-06", "Cập nhật trạng thái bảo hành", "Chỉ tài khoản nội bộ có quyền", "Đạt (B)"),
    ), (1.7, 4.7, 6.0, 3.0), 9.4)
    table(doc, "5.7", "Test case hoàn trả", ("TC", "Kịch bản", "Kết quả mong đợi", "Kết quả/bằng chứng"), (
        ("RET-01", "Hoàn một phần số đã mua", "Tạo chứng từ và RESTORE đúng lượng", "Đạt (B,C)"),
        ("RET-02", "Hoàn vượt số đã mua trừ số đã hoàn", "Từ chối, không đổi dữ liệu", "Đạt (B)"),
        ("RET-03", "Hoàn sản phẩm không thuộc đơn", "Từ chối", "Đạt (B)"),
        ("RET-04", "Hai yêu cầu hoàn cùng dòng", "Tổng hoàn không vượt số mua", "Có kiểm tra tổng (B); cần tải đồng thời (D)"),
        ("RET-05", "Transaction lỗi khi phục hồi kho", "Rollback phiếu hoàn và tồn", "Đạt theo transaction (B)"),
        ("RET-06", "Xem lịch sử hoàn", "Truy được chứng từ, người xử lý và lý do", "Đạt (B,C)"),
    ), (1.7, 4.7, 6.0, 3.0), 9.4)

    heading(doc, "5.6. Kiểm thử các chức năng AI", 2)
    heading(doc, "5.6.1. Trợ lý gợi ý bán hàng", 3)
    table(doc, "5.8", "Kết quả kiểm tra trợ lý bán hàng", ("Nội dung", "Kết quả", "Diễn giải"), (
        ("Cách gợi ý cục bộ", "8 tình huống; 56/56 kiểm tra đạt", "ID có thật, còn tồn, không trùng, đúng ngân sách và tối đa 5 sản phẩm"),
        ("Gemini", "Đạt ở lần kiểm tra đầu", "Trả JSON hợp lệ cho dữ liệu giả lập của POS và kho"),
        ("Gemini ở lần kiểm tra sau", "HTTP 429", "Khóa hoặc dự án đã chạm giới hạn tần suất/hạn mức tại thời điểm thử"),
        ("Groq", "Đạt", "GPT-OSS trả JSON hợp lệ cho POS/Kho; Qwen nhận ảnh giả hợp lệ"),
        ("Phương án cuối", "Đạt", "POS dùng cách chấm điểm cục bộ; kho giữ kết quả công thức"),
    ), (4.7, 3.0, 7.7), 9.5)
    body(doc, "Thước đo trên đánh giá tính an toàn và hợp lệ của đề xuất, chưa chứng minh mức độ thuyết phục khách hàng hoặc tăng doanh số. Muốn đánh giá chất lượng kinh doanh cần A/B test hoặc đánh giá của nhân viên trên tập nhu cầu thực tế.")
    body(doc, "Kết quả kết nối chỉ dùng dữ liệu và ảnh giả lập, không gửi dữ liệu sản phẩm hoặc tồn kho thật ra ngoài. Gemini đã hoạt động đúng ở lần kiểm tra đầu nhưng lần chạy tiếp theo nhận HTTP 429 do hạn mức tại thời điểm thử. Groq hoạt động thành công: model openai/gpt-oss-120b trả JSON hợp lệ cho POS và Kho, còn qwen/qwen3.6-27b nhận ảnh giả hợp lệ. Nếu cả hai dịch vụ không dùng được, hệ thống vẫn chuyển sang phương án cục bộ.")

    heading(doc, "5.6.2. Ước lượng nhu cầu và đề xuất nhập hàng", 3)
    body(doc, "Phép thử tách hai mục tiêu: 24 kiểm tra bất biến hiện tại xác nhận nguồn công thức, SKU và số lượng hợp lệ; backtest định lượng mới đánh giá sai số nhu cầu. Dự án không có dữ liệu huấn luyện và không huấn luyện LLM. Công thức theo luật được so với baseline trung bình trượt 7 ngày và 30 ngày bằng rolling-origin, tuyệt đối không dùng dữ liệu tương lai tại thời điểm dự đoán.")
    table(doc, "5.9", "Kiểm tra tính hợp lệ kỹ thuật", ("Nội dung", "Kết quả", "Ý nghĩa"), (
        ("Kỳ phân tích API", "7, 15, 30 ngày", "Kiểm tra cấu trúc ở ba tham số"),
        ("Kiểm tra bất biến hiện tại", "24/24 đạt", "ID thật, tối đa 12 dòng, số không âm; không bán chỉ nhập bù minStock khi thiếu"),
        ("Gemini/GitHub không dùng được", "Có fallback công thức", "Không làm endpoint mất khả dụng; kết quả ghi source = FORMULA"),
        ("Quyền phê duyệt", "Không thuộc LLM", "Admin quyết định tạo/duyệt PurchaseOrder"),
    ), (4.3, 3.0, 8.1), 9.4)
    table(doc, "5.10", "Thiết kế hai tầng backtest", ("Tiêu chí", "Dữ liệu demo quan sát", "Tập mô phỏng có kiểm soát"), (
        ("Nguồn", "Bản PostgreSQL lưu ngày 29/07", "Sinh bằng script, seed cố định"),
        ("Quy mô", "Bản lưu 29/07: 11 đơn; 34 dòng; 4 SKU; 16 ngày", "180 ngày; 8 SKU; 1.885 đơn vị holdout"),
        ("Độ thưa", "36 quan sát, 3 actual > 0; zero rate 91,67%", "168 quan sát, 135 actual > 0; zero rate 19,64%"),
        ("Horizon", "1 ngày; bước 1 ngày", "7 ngày; bước 7 ngày"),
        ("Kịch bản", "Không có promotion; chưa đủ mùa", "Ổn định, tăng/giảm, mùa, promotion, intermittent, new/cold-start"),
        ("Cách dùng kết quả", "Chỉ thăm dò", "Kiểm tra hành vi; không đại diện cửa hàng thật"),
    ), (3.1, 6.0, 6.3), 8.7)
    table(doc, "5.11", "Sai số trên dữ liệu demo", ("Phương pháp", "MAE", "RMSE", "WAPE", "MAPE*", "Bias"), (
        ("Công thức HomeX", "1,033", "2,172", "743,9%", "100,0%", "+0,755"),
        ("Baseline 7 ngày", "0,421", "0,702", "302,9%", "100,0%", "+0,143"),
        ("Baseline 30 ngày", "0,331", "0,496", "238,0%", "78,3%", "+0,103"),
    ), (4.2, 2.0, 2.1, 2.3, 2.3, 2.1), 8.8)
    body(doc, "* MAPE chỉ tính ba quan sát có actual > 0. Tổng nhu cầu holdout chỉ 5 đơn vị nên WAPE rất lớn và không ổn định. Công thức HomeX có MAE/RMSE cao hơn hai baseline và bias dương, cho thấy xu hướng ước lượng dư. Đây không phải bằng chứng đủ để suy rộng cho cửa hàng.")
    table(doc, "5.12", "Sai số trên tập mô phỏng 180 ngày", ("Phương pháp", "MAE", "RMSE", "WAPE", "MAPE*", "Bias"), (
        ("Công thức HomeX", "7,120", "12,156", "63,5%", "91,3%", "+3,675"),
        ("Baseline 7 ngày", "4,429", "6,865", "39,5%", "52,4%", "-0,191"),
        ("Baseline 30 ngày", "4,105", "6,778", "36,6%", "49,3%", "-0,738"),
    ), (4.2, 2.0, 2.1, 2.3, 2.3, 2.1), 8.8)
    table(doc, "5.13", "Sai số ở các trường hợp khó", ("Lát cắt", "Số quan sát", "HomeX MAE/WAPE", "Baseline 7 ngày MAE/WAPE", "Kết luận"), (
        ("Có promotion", "6", "21,158 / 74,2%", "14,167 / 49,7%", "Chưa có biến kế hoạch promotion nên sai số tăng"),
        ("Trong mùa vụ", "95", "9,070 / 61,6%", "5,400 / 36,7%", "Hệ số mùa cố định khuếch đại quá mức"),
        ("Cold-start", "33", "0,546 / 100%", "0,546 / 100%", "Không lịch sử nên cả hai dự đoán 0, bỏ lỡ nhu cầu mới"),
        ("Lịch sử < 5 đơn vị", "34", "0,643 / 81,0%", "0,706 / 88,9%", "Kết quả rất nhạy; cần Admin/minStock"),
    ), (3.0, 2.2, 3.5, 4.0, 3.0), 8.4)
    body(doc, "Tập mô phỏng cho phép lặp lại phép thử nhưng không phải dữ liệu cửa hàng thật. Trên cả hai tầng, công thức hiện tại không vượt baseline và có xu hướng ước lượng dư; riêng promotion, mùa vụ và cold-start còn yếu. Vì vậy báo cáo không gọi đây là mô hình dự báo đã được kiểm chứng. Chức năng chỉ nên cảnh báo/đề xuất tham khảo, còn Admin phê duyệt. Hướng tiếp theo là thu thập tối thiểu 90–180 ngày bán thực, lưu lịch promotion, hiệu chỉnh hệ số theo backtest và chỉ triển khai công thức nếu vượt baseline ổn định.")

    heading(doc, "5.6.3. Nhận dạng ảnh hỗ trợ làm giàu sản phẩm", 3)
    table(doc, "5.14", "Độ chính xác theo trường trên 12 ảnh", ("Trường/tiêu chí", "Kết quả", "Nhận xét"), (
        ("Tên sản phẩm cụ thể", "4/12 = 33,3%", "Thấp; không dùng để tự động ghi đè tên"),
        ("Danh mục", "6/12 = 50,0%", "Chỉ phù hợp gợi ý sơ bộ"),
        ("Không bịa NCC/giá/tồn/bảo hành", "12/12 = 100%", "Đúng nguyên tắc để null khi không quan sát được"),
    ), (5.0, 3.2, 7.2), 9.5)
    body(doc, "Đánh giá theo từng trường cho thấy ảnh không đủ tin cậy để tự động nhập đầy đủ hồ sơ sản phẩm. Dự án vì vậy ưu tiên barcode và nguồn dữ liệu ngoài; ảnh chỉ bổ trợ khi thiếu nhận dạng, còn nhà cung cấp, giá, tồn và thời hạn bảo hành phải do dữ liệu có nguồn hoặc người dùng xác nhận.")

    heading(doc, "5.7. Kiểm tra chất lượng xây dựng", 2)
    table(doc, "5.15", "Kiểm tra tĩnh và build", ("Hạng mục", "Lệnh/điều kiện", "Kết quả"), (
        ("Backend TypeScript", "npx tsc --noEmit", "Không lỗi sau khi sửa benchmark script"),
        ("Frontend production", "npm run build", "Build thành công, tạo đủ route"),
        ("Tệp benchmark", "JSON trong docs/benchmarks", "Lưu đầu vào, kết quả và thời điểm chạy"),
        ("Dữ liệu PayOS tạm", "Kiểm tra hậu benchmark", "Không còn fixture, tồn được phục hồi"),
    ), (4.2, 6.0, 5.2), 9.5)
    body(doc, "Build thành công chứng minh mã biên dịch và các route được tạo, nhưng không thay thế kiểm thử đơn vị, tích hợp, hiệu năng, bảo mật xâm nhập hoặc khả năng sử dụng với người dùng thật.")

    heading(doc, "5.8. Đối chiếu mục tiêu - kết quả - minh chứng", 2)
    body(doc, "Bảng dưới đây thay cho cách tự chấm điểm. Mỗi mục tiêu được đối chiếu với sản phẩm quan sát được, bằng chứng và giới hạn còn lại; báo cáo không tự gán điểm hoặc khẳng định mức hoàn thành tuyệt đối.")
    table(doc, "5.16", "Đối chiếu khách quan mục tiêu đề tài", ("Mục tiêu", "Kết quả hiện có", "Minh chứng", "Giới hạn"), (
        ("Bán hàng POS", "Giỏ, ca, khuyến mãi, tiền mặt/PayOS, hóa đơn", "Route, schema, UI, TC SALE", "Một cửa hàng; chưa offline"),
        ("Quản lý hàng hóa", "Sản phẩm, danh mục, NCC, barcode", "UI, API, Product/Supplier", "Chưa model/serial/lô"),
        ("Kho và nhập", "PO, nhập, điều chỉnh, sổ kho, cảnh báo", "PurchaseOrder, StockTransaction, TC INV", "Chưa vị trí/chuyển kho"),
        ("Sau bán", "Bảo hành, hoàn trả, VAT", "UI, API, TC WAR/RET", "Chưa quy trình hãng/lắp đặt"),
        ("Thanh toán an toàn", "Xác minh, đối chiếu, khóa, idempotency, log", "10/10 benchmark webhook", "Chưa test PayOS tiền thật/sandbox"),
        ("Hỗ trợ thông minh", "Gợi ý bán, ước lượng nhu cầu nhập và làm giàu sản phẩm", "4 JSON benchmark", "Ước lượng nhập chưa vượt baseline; ảnh còn thấp"),
        ("Báo cáo/phân quyền", "Dashboard, doanh thu, lợi nhuận, Admin/Cashier", "UI, API, middleware", "Chưa test tải dữ liệu lớn"),
        ("Khả năng xây dựng", "Frontend/backend biên dịch thành công", "tsc và production build", "Chưa CI/test suite đầy đủ"),
    ), (3.3, 5.0, 3.8, 3.3), 9.0)

    heading(doc, "5.9. Hạn chế và rủi ro còn lại", 2)
    bullets(doc, (
        "Chưa có bộ unit, integration và end-to-end bao phủ toàn bộ 117 khai báo HTTP trong router; benchmark hiện tập trung vào các rủi ro mà giảng viên yêu cầu.",
        "PayOS mới được kiểm tra bằng chữ ký và endpoint cục bộ, chưa thực hiện giao dịch tiền thật hoặc thử nghiệm mất mạng dài hạn.",
        "Schema chưa quản lý serial/IMEI, lô/model riêng, phụ kiện đi kèm, giao hàng cồng kềnh, lắp đặt và phân biệt bảo hành cửa hàng với hãng.",
        "AI ảnh có độ chính xác tên/danh mục thấp trên tập 12 ảnh. Backtest nhập hàng trên dữ liệu demo rất thưa và phép thử mô phỏng không thay thế dữ liệu cửa hàng thật; công thức hiện tại còn kém hai baseline nên chỉ được dùng như cảnh báo tham khảo.",
        "Hệ thống hướng tới một cửa hàng; chưa có đa chi nhánh, tồn theo vị trí, thanh toán kết hợp, hàng đợi offline và kiểm thử tải.",
        "Bảo mật production cần bổ sung rate limiting, CORS cụ thể, refresh/thu hồi phiên, CSP, quản lý secret, backup và quan sát tập trung.",
    ))

def chapter6(doc: Document) -> None:
    heading(doc, "CHƯƠNG 6. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1, True)
    heading(doc, "6.1. Kết luận", 2)
    body(doc, "Đề tài đã phân tích, thiết kế và xây dựng HomeX POS cho chu trình của một cửa hàng đồ gia dụng: dữ liệu hàng hóa và nhà cung cấp, bán tại quầy, tiền mặt/PayOS, tồn kho, nhập hàng, khách hàng, bảo hành, hoàn trả, VAT, ca làm việc và báo cáo. Next.js, Express, Prisma và PostgreSQL tạo ranh giới rõ giữa giao diện, quy tắc nghiệp vụ và dữ liệu; barcode, quét từ xa, Gemini và Groq bổ sung khả năng hỗ trợ thao tác; mọi kết quả AI vẫn được hậu kiểm hoặc cần người dùng duyệt.")
    body(doc, "Kết quả kiểm thử cho thấy cơ chế webhook đạt 10/10 điều kiện benchmark về chữ ký, idempotency và tác dụng phụ; trợ lý bán hàng cục bộ đạt 56/56 ràng buộc trên 8 tình huống; đầu ra ước lượng nhập đạt 24/24 kiểm tra bất biến. Tuy nhiên, đánh giá sai số cho thấy công thức nhập hàng chưa vượt baseline: trên tập mô phỏng WAPE là 63,46%, trong khi trung bình trượt 7 ngày đạt 39,47% và 30 ngày đạt 36,58%. Backtest ngày 29/07 chỉ có 11 đơn, 4 sản phẩm phát sinh bán và 5 ngày bán trong 16 ngày nên chưa đủ đại diện; dữ liệu hiện tại đã bổ sung 25 đơn demo tháng 8 nhưng vẫn chưa phải dữ liệu cửa hàng thật. Nhận dạng ảnh cũng chỉ đạt 33,3% theo tên và 50,0% theo danh mục trên 12 ảnh. Vì vậy hai chức năng này chỉ hỗ trợ người dùng và không tự động ghi dữ liệu thương mại.")
    body(doc, "Sản phẩm đáp ứng luồng cốt lõi trong phạm vi một cửa hàng và môi trường demo. Kết luận này dựa trên mã nguồn, giao diện, build và benchmark tái lập; không được hiểu là hệ thống đã sẵn sàng production hoặc đã thay thế kiểm thử với giao dịch tiền thật, dữ liệu lớn và người dùng thực tế.")

    heading(doc, "6.2. Hướng phát triển", 2)
    bullets(doc, (
        "Chuẩn hóa nghiệp vụ đồ gia dụng bằng ProductModel, ProductSerial/IMEI, InventoryBatch, phụ kiện đi kèm và chính sách bảo hành theo model hoặc từng thiết bị.",
        "Phân biệt bảo hành tại cửa hàng và bảo hành hãng; quản lý nhiều lần tiếp nhận, chuyển hãng, linh kiện thay thế, chi phí và lịch sử trạng thái.",
        "Bổ sung DeliveryOrder và InstallationJob cho hàng cồng kềnh: địa chỉ, khung giờ, phí, đội giao/lắp, nghiệm thu và đổi trả sau lắp đặt.",
        "Mở rộng kho theo chi nhánh/vị trí, chuyển kho, kiểm kê, theo dõi lô và đối soát tồn bằng sổ StockTransaction.",
        "Hỗ trợ thanh toán kết hợp, hoàn tiền qua cổng, hàng đợi webhook, retry có kiểm soát và quy trình đối soát tự động.",
        "Triển khai unit, integration, end-to-end, concurrency, security và load test; chạy tự động trong CI/CD với CSDL kiểm thử biệt lập.",
        "Thu thập tối thiểu 90–180 ngày dữ liệu bán theo SKU, lưu lịch khuyến mãi và biến mùa vụ; backtest theo thời gian bằng MAE, RMSE, WAPE/MAPE và chỉ nâng thành dự báo tự động khi vượt baseline ổn định.",
        "Cải thiện nhận dạng ảnh bằng taxonomy đóng, OCR/vision có tập gán nhãn lớn hơn và ngưỡng tin cậy; trường không chắc chắn luôn cần người dùng duyệt.",
        "Bổ sung PWA/offline queue, rate limit, refresh/thu hồi phiên, CSP, secret manager, backup và quan sát tập trung trước khi production.",
    ))

    heading(doc, "6.3. Kiến nghị", 2)
    body(doc, "Trước khi vận hành thực tế, cửa hàng cần chuẩn hóa SKU, model, serial, giá vốn, nhà cung cấp, ngưỡng tồn, thời hạn và đơn vị bảo hành; đồng thời ban hành quy trình giao/lắp, đổi trả và đối soát thanh toán. PayOS phải thử ở môi trường phù hợp với webhook công khai và kịch bản mất mạng. AI chỉ nên tạo gợi ý, mọi thay đổi giá, tồn, phiếu nhập và dữ liệu bảo hành phải có nguồn xác thực hoặc phê duyệt của người có trách nhiệm.")

    heading(doc, "DANH MỤC TÀI LIỆU THAM KHẢO", 1, True)
    refs = (
        "[1] Microsoft, ‘TypeScript Documentation,’ [Online]. Available: https://www.typescriptlang.org/docs/. [Accessed: Jul. 26, 2026].",
        "[2] Meta Platforms, Inc., ‘React Documentation,’ [Online]. Available: https://react.dev/. [Accessed: Jul. 26, 2026].",
        "[3] Vercel, ‘Next.js Documentation,’ [Online]. Available: https://nextjs.org/docs. [Accessed: Jul. 26, 2026].",
        "[4] OpenJS Foundation, ‘Express Documentation,’ [Online]. Available: https://expressjs.com/. [Accessed: Jul. 26, 2026].",
        "[5] PostgreSQL Global Development Group, ‘PostgreSQL Documentation,’ [Online]. Available: https://www.postgresql.org/docs/. [Accessed: Jul. 26, 2026].",
        "[6] Prisma Data, Inc., ‘Prisma ORM Documentation,’ [Online]. Available: https://www.prisma.io/docs/orm. [Accessed: Jul. 26, 2026].",
        "[7] M. Jones, J. Bradley, and N. Sakimura, ‘JSON Web Token (JWT),’ RFC 7519, May 2015.",
        "[8] Zod, ‘Zod Documentation,’ [Online]. Available: https://zod.dev/. [Accessed: Jul. 26, 2026].",
        "[9] PayOS, ‘PayOS API and Webhook Documentation,’ [Online]. Available: https://payos.vn/docs/api/. [Accessed: Jul. 26, 2026].",
        "[10] Google, ‘OpenAI compatibility for the Gemini API,’ [Online]. Available: https://ai.google.dev/gemini-api/docs/openai. [Accessed: Aug. 11, 2026].",
        "[11] Groq, ‘OpenAI Compatibility,’ [Online]. Available: https://console.groq.com/docs/openai. [Accessed: Aug. 11, 2026].",
        "[12] UPCitemdb, ‘API Documentation and Terms of Service,’ [Online]. Available: https://www.upcitemdb.com/api and https://www.upcitemdb.com/terms. [Accessed: Jul. 29, 2026].",
        "[13] Barcode Lookup, ‘API Documentation and Terms and Conditions,’ [Online]. Available: https://www.barcodelookup.com/api-documentation and https://www.barcodelookup.com/terms-and-conditions. [Accessed: Jul. 29, 2026].",
        "[14] Open Food Facts, ‘Data, contents and image licensing,’ [Online]. Available: https://openfoodfacts.github.io/documentation/docs/Product-Opener/api/tutorials/license-be-on-the-legal-side/. [Accessed: Jul. 29, 2026].",
        "[15] Open Food Facts, ‘Scanning cosmetics, pet food and other products,’ [Online]. Available: https://openfoodfacts.github.io/documentation/docs/Product-Opener/api/tutorials/scanning-cosmetics-pet-food-and-other-products/. [Accessed: Jul. 29, 2026].",
        "[16] Barcode Spider, ‘Terms of Service,’ [Online]. Available: https://devapi.barcodespider.com/terms. [Accessed: Jul. 29, 2026].",
        "[17] S. Minhazav, ‘html5-qrcode Documentation,’ GitHub. [Online]. Available: https://github.com/mebjas/html5-qrcode. [Accessed: Jul. 26, 2026].",
        "[18] OWASP Foundation, ‘Application Security Verification Standard,’ [Online]. Available: https://owasp.org/www-project-application-security-verification-standard/. [Accessed: Jul. 26, 2026].",
        "[19] I. Jacobson, G. Booch, and J. Rumbaugh, The Unified Software Development Process. Addison-Wesley, 1999.",
        "[20] R. Elmasri and S. B. Navathe, Fundamentals of Database Systems, 7th ed. Pearson, 2016.",
    )
    for ref in refs:
        paragraph = body(doc, ref, False)
        paragraph.paragraph_format.left_indent = Cm(0.8)
        paragraph.paragraph_format.first_line_indent = Cm(-0.8)

def appendices(doc: Document) -> None:
    heading(doc, "PHỤ LỤC A. HƯỚNG DẪN CÀI ĐẶT", 1, True)
    heading(doc, "A.1. Chuẩn bị", 2)
    bullets(doc, (
        "Cài Node.js, npm và PostgreSQL; tạo database rỗng.",
        "Backend: tạo .env với DATABASE_URL, JWT_SECRET; thêm GEMINI_API_KEY (chính), GEMINI_MODEL nếu muốn đổi model, GROQ_API_KEY (dự phòng), GROQ_MODEL/GROQ_VISION_MODEL nếu muốn đổi model, khóa PayOS và các khóa barcode tùy chọn khi dùng.",
        "Frontend: tạo .env.local với NEXT_PUBLIC_API_URL=http://localhost:5000/api.",
    ))
    heading(doc, "A.2. Chạy hệ thống", 2)
    body(doc, "Backend: cd backend; npm install; npx prisma migrate dev; npx prisma db seed; npm run dev.", False)
    body(doc, "Frontend: cd frontend; npm install; npm run dev. Truy cập http://localhost:3000.", False)

    heading(doc, "PHỤ LỤC B. TÀI KHOẢN VÀ KỊCH BẢN DEMO", 1, True)
    table(doc, "B.1", "Tài khoản demo", ("Vai trò", "Email", "Mật khẩu"), (
        ("ADMIN", "admin@homex.com", "123456"), ("CASHIER", "cashier@homex.com", "123456"),
    ), (4.0, 7.0, 4.4))
    heading(doc, "B.1. Kịch bản bảo vệ đề tài", 2)
    bullets(doc, (
        "Đăng nhập Cashier và mở ca.", "Mở POS, quét/chọn hàng, tìm khách và áp dụng khuyến mãi.",
        "Dùng trợ lý gợi ý và thêm một sản phẩm phù hợp.", "Thanh toán tiền mặt/PayOS, xem hóa đơn và tồn đã giảm.",
        "Tra cứu bảo hành ở trang công khai.", "Đăng nhập Admin, xem báo cáo và cảnh báo tồn.",
        "Duyệt gợi ý nhập, tạo phiếu và xem StockTransaction.", "Mở audit log để chứng minh truy vết.",
    ), True)
    heading(doc, "PHỤ LỤC C. MA TRẬN PHÂN QUYỀN", 1, True)
    table(doc, "C.1", "Ma trận quyền", ("Chức năng", "Admin", "Cashier", "Khách"), (
        ("Đăng nhập", "Có", "Có", "Không"), ("Dashboard", "Toàn bộ", "Giới hạn", "Không"),
        ("Bán hàng/đơn", "Có", "Có", "Không"), ("Khách hàng/bảo hành nội bộ", "Có", "Có", "Không"),
        ("Sản phẩm/NCC", "Có", "Không", "Không"), ("Kho/nhập/hoàn", "Có", "Không", "Không"),
        ("Khuyến mãi/VAT", "Có", "Giới hạn", "Gửi yêu cầu"), ("Người dùng/cài đặt/audit", "Có", "Không", "Không"),
        ("Tra cứu bảo hành", "Có", "Có", "Công khai"),
    ), (7.0, 2.8, 2.8, 2.8), 10.5)
    heading(doc, "PHỤ LỤC D. NHÓM REST API", 1, True)
    table(doc, "D.1", "Base path theo phân hệ", ("Base path", "Nội dung"), (
        ("/api/auth", "Xác thực"), ("/api/products, /categories, /suppliers", "Hàng hóa"),
        ("/api/customers, /warranties", "Khách hàng và bảo hành"),
        ("/api/orders, /payments", "Đơn và thanh toán"),
        ("/api/inventory, /purchase-orders, /return-orders", "Kho, nhập và hoàn"),
        ("/api/promotions, /vat-invoices", "Khuyến mãi và VAT"),
        ("/api/shifts, /settings, /notifications", "Vận hành"),
        ("/api/reports, /audit-logs", "Báo cáo và giám sát"),
        ("/api/pos", "Quét từ xa và trợ lý"), ("/api/invoices/public", "Hóa đơn/VAT công khai"),
    ), (7.4, 8.0), 10.5)
    heading(doc, "PHỤ LỤC E. MINH CHỨNG KIỂM THỬ TÁI LẬP", 1, True)
    body(doc, "Các tập lệnh dưới đây tạo tệp JSON ghi thời điểm, dữ liệu đầu vào, từng phép kiểm tra và kết quả tổng hợp. Khóa bí mật không được ghi vào báo cáo hoặc tệp kết quả.")
    table(doc, "E.1", "Tập lệnh và tệp kết quả", ("Nhóm", "Tập lệnh", "Tệp kết quả"), (
        ("PayOS/webhook", "backend/scripts/benchmark-payos-webhook.ts", "docs/benchmarks/payos-webhook-benchmark.json"),
        ("Trợ lý bán hàng hiện tại", "backend/scripts/benchmark-sales-ai-current.ts", "docs/benchmarks/sales-ai-current-benchmark.json"),
        ("Kết nối Gemini (dữ liệu giả)", "backend/scripts/verify-gemini-ai.ts", "Kết quả tại cửa sổ lệnh"),
        ("Kết nối Groq (dữ liệu giả)", "backend/scripts/verify-groq-ai.ts", "POS, Kho và ảnh đều đạt"),
        ("Bất biến nhập hàng hiện tại", "backend/scripts/benchmark-inventory-ai-current.ts", "docs/benchmarks/inventory-ai-current-benchmark.json"),
        ("Sai số ước lượng nhập", "backend/scripts/benchmark-inventory-forecast-accuracy.ts", "docs/benchmarks/inventory-forecast-accuracy.json"),
        ("AI ảnh", "backend/scripts/benchmark-ai-images.ts", "docs/benchmarks/ai-image-benchmark.json"),
    ), (3.3, 6.2, 5.9), 9.1)
    table(doc, "E.2", "Mã nguồn làm căn cứ đối chiếu", ("Nội dung", "Vị trí"), (
        ("Webhook, đồng bộ PayOS và transaction hoàn tất", "backend/src/routes/payment.routes.ts"),
        ("Trợ lý bán hàng và hậu kiểm ứng viên", "backend/src/services/sales-assistant.service.ts"),
        ("Ước lượng nhu cầu nhập hàng", "backend/src/services/inventory-ai.service.ts"),
        ("Chọn Gemini/GitHub và chuyển phương án", "backend/src/services/ai-provider.service.ts"),
        ("Làm giàu barcode/ảnh", "backend/src/services/barcode-enrichment.service.ts"),
        ("Lược đồ dữ liệu", "backend/prisma/schema.prisma"),
    ), (6.0, 9.4), 9.5)


def build_report() -> Path:
    doc = Document()
    configure(doc)
    cover(doc)
    doc.add_page_break()
    cover(doc)
    front_matter(doc)
    opening(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter3_design(doc)
    chapter4(doc)
    chapter5(doc)
    chapter6(doc)
    appendices(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_report().name)
