from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from generate_ai_analysis_doc import (
    BLACK,
    BLUE,
    LIGHT_BLUE,
    LIGHT_GREEN,
    LIGHT_RED,
    LIGHT_YELLOW,
    MID_GRAY,
    NAVY,
    add_bullets,
    add_callout,
    add_code,
    add_flow,
    add_numbered,
    add_plain,
    add_table,
    set_cell_margins,
    set_cell_text,
    setup_document,
    shade,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Phan_tich_hoat_dong_AI_trong_HomeX_POS.docx"


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("HOMEX POS — SỔ TAY HỌC NHANH")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("NẮM CÁCH AI HOẠT ĐỘNG\nTRONG TRANG POS VÀ KHO HÀNG")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("Tài liệu để đọc code, học luồng chạy, debug và trả lời khi được hỏi")
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(MID_GRAY)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Dành cho", "Người cần hiểu dự án, không phải nội dung chép vào báo cáo"),
        ("Phạm vi chính", "AI tại trang POS và trang Kho hàng"),
        ("Mô hình", "openai/gpt-oss-120b qua Groq"),
        ("Hai service", "sales-assistant.service.ts và inventory-ai.service.ts"),
        ("Cập nhật code", "11/08/2026 — sau khi tách riêng service AI Kho"),
    ]
    for i, (key, value) in enumerate(rows):
        shade(table.rows[i].cells[0], BLUE)
        set_cell_text(table.rows[i].cells[0], key, bold=True, color="FFFFFF", size=10.5)
        set_cell_text(table.rows[i].cells[1], value, size=10.5)
        set_cell_margins(table.rows[i].cells[0], 120, 140, 120, 140)
        set_cell_margins(table.rows[i].cells[1], 120, 140, 120, 140)

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph("Nguyễn Đức Thịnh — 2305CT2084 — CT07PM")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.runs[0].bold = True
    doc.add_page_break()


def build():
    doc = setup_document()
    add_cover(doc)

    doc.add_heading("CÁCH DÙNG TÀI LIỆU NÀY", level=1)
    add_plain(doc, "Đừng cố học thuộc từng dòng code. Hãy nắm một đường đi cố định: người dùng bấm nút ở frontend → frontend service gọi API → route kiểm tra quyền → backend service lấy dữ liệu và tính toán → nếu có token thì gọi mô hình → kiểm tra hoặc fallback → trả kết quả về giao diện.")
    add_callout(doc, "Mục tiêu sau khi đọc", "Bạn có thể tự mở đúng file khi có lỗi, giải thích phần nào là thuật toán và phần nào thật sự gọi AI, biết hệ thống chạy thế nào khi mất token, và trình bày được mỗi chức năng trong khoảng 60–90 giây.", LIGHT_GREEN)
    add_numbered(doc, [
        "Đọc phần 1 và 2 để có bản đồ tổng thể.",
        "Học phần 3 cho POS và phần 4 cho Kho theo đúng thứ tự file.",
        "Làm ví dụ tính tay ở phần Kho để hiểu các con số.",
        "Dùng phần 6 khi debug và phần 7 để luyện trả lời.",
    ])

    doc.add_heading("1. TRẢ LỜI NGẮN GỌN: AI ĐANG Ở ĐÂU?", level=1)
    add_callout(doc, "Câu trả lời đúng", "Hai chức năng AI chính trên giao diện là: (1) Trợ lý gợi ý bán hàng ở trang POS; (2) Phân tích và đề xuất nhập hàng ở trang Kho. Ngoài ra trang Sản phẩm có AI hỗ trợ bù thông tin barcode, nhưng đó là chức năng phụ trợ nhập liệu, không phải một màn hình phân tích AI chính.", LIGHT_BLUE)
    add_table(doc, ["Trang", "Người dùng làm gì?", "Service backend chính", "Nếu AI lỗi"], [
        ("POS", "Nhập nhu cầu, ngân sách, ưu tiên; nhận gợi ý sản phẩm.", "sales-assistant.service.ts", "Chấm điểm heuristic và vẫn trả gợi ý."),
        ("Kho hàng", "Chọn số ngày dự phòng; nhận đề xuất nhập và phân tích rủi ro.", "inventory-ai.service.ts", "Công thức toán học và fallback vẫn chạy."),
    ], [2.4, 5.3, 4.0, 3.2])
    add_plain(doc, "Điểm quan trọng sau khi refactor: AI Kho không còn nằm trực tiếp trong inventory.routes.ts. Route chỉ tiếp nhận request và gọi inventoryAiService.forecast(days). Nhờ vậy cấu trúc POS và Kho đã dễ nhìn hơn: mỗi chức năng có một service AI riêng.")

    doc.add_heading("2. BẢN ĐỒ LUỒNG CHẠY CHUNG", level=1)
    add_flow(doc, [
        "1. Người dùng bấm nút trên trang Next.js",
        "2. Hàm xử lý trong page.tsx tạo payload/tham số",
        "3. homex.service.ts gửi request tới Express API",
        "4. Route xác thực JWT và vai trò",
        "5. Service lấy dữ liệu thật bằng Prisma",
        "6. Thuật toán lọc/tính toán trước khi gọi AI",
        "7. Có GROQ_API_KEY? Có → gọi openai/gpt-oss-120b | Không → dùng fallback",
        "8. Parse/kiểm tra JSON và ghép lại dữ liệu thật",
        "9. Response quay về giao diện để người dùng xem và quyết định",
    ], "Bản đồ này áp dụng cho cả POS và Kho hàng", [LIGHT_BLUE, LIGHT_GREEN])
    add_callout(doc, "Nhớ một câu", "Database và thuật toán tạo nền dữ liệu; AI giúp chọn và diễn giải; con người quyết định hành động cuối cùng.", LIGHT_YELLOW)

    doc.add_heading("3. AI TRÊN TRANG POS CHẠY NHƯ THẾ NÀO?", level=1)
    doc.add_heading("3.1. Đọc code theo thứ tự này", level=2)
    add_table(doc, ["Bước", "File và vị trí", "Tìm hàm/dòng", "Ý nghĩa"], [
        ("1", "frontend/app/(dashboard)/pos/page.tsx", "handleRequestSalesSuggestion — khoảng dòng 244", "Lấy nhu cầu, ngân sách, khách hàng và giỏ hàng."),
        ("2", "frontend/services/homex.service.ts", "getSalesAssistantSuggestions — dòng 381", "POST /pos/sales-assistant."),
        ("3", "backend/src/routes/sales-assistant.routes.ts", "POST /sales-assistant — dòng 30", "Zod + JWT + quyền ADMIN/CASHIER."),
        ("4", "backend/src/services/sales-assistant.service.ts", "getSuggestions — dòng 211", "Toàn bộ logic gợi ý và gọi AI."),
    ], [1.2, 5.3, 4.2, 4.1], font_size=9)

    doc.add_heading("3.2. Frontend gửi gì?", level=2)
    add_code(doc, '''{
  need: "Khách cần nồi cơm điện cho gia đình 4 người",
  budgetMin: 300000,
  budgetMax: 1200000,
  customerId: 15,
  cartItems: [{ productId: 8, name: "Ấm đun", quantity: 1 }],
  preferences: {
    preferPromotion: false,
    preferWarranty: true,
    preferHighStock: false,
    crossSellFromCart: true
  }
}''', "Ví dụ payload của Sales Assistant")
    add_plain(doc, "Frontend không tự chạy AI. Nó chỉ gom dữ liệu từ form và giỏ hàng rồi gọi API. Khi response về, frontend hiển thị summary, recommendations, bundleSuggestion, cashierTips và cho thu ngân bấm thêm sản phẩm vào giỏ.")

    doc.add_heading("3.3. Backend làm gì trước khi gọi AI?", level=2)
    add_numbered(doc, [
        "Lấy Product có status ACTIVE và stockQuantity > 0.",
        "Lọc giá theo budgetMin và budgetMax.",
        "Chuẩn hóa nhu cầu: bỏ dấu tiếng Việt, chuyển chữ thường và dò cụm từ cốt lõi.",
        "Tính relevance giữa nhu cầu với tên, danh mục và mô tả sản phẩm.",
        "Cộng điểm khuyến mãi, bảo hành, tồn kho cao và cùng danh mục với giỏ hàng.",
        "Sắp xếp và chỉ gửi tối đa 25 candidateProducts cho mô hình.",
    ])
    add_table(doc, ["Tín hiệu", "Ví dụ", "Tác động"], [
        ("Đúng nhu cầu", "‘nồi cơm’ khớp Rice cooker/Nồi cơm điện", "Cộng nền 1000 điểm nên ưu tiên rất cao."),
        ("Ưu đãi", "originalPrice > salePrice", "Cộng điểm khuyến mãi."),
        ("Bảo hành", "warrantyMonths > 0", "Cộng điểm, nhiều hơn khi khách ưu tiên."),
        ("Mua kèm", "Cùng category với sản phẩm trong giỏ", "Cộng điểm cross-sell."),
        ("Phụ kiện sai ý", "Khách tìm thiết bị chính nhưng kết quả là bóng đèn/filter", "Trừ 300 điểm."),
    ], [3.0, 6.1, 5.5])
    add_callout(doc, "Điều dễ hiểu nhầm", "AI không tìm kiếm trực tiếp toàn bộ database. Backend đã lọc và chọn ứng viên trước. AI chỉ được lựa chọn trong danh sách tối đa 25 sản phẩm đã gửi.", LIGHT_YELLOW)

    doc.add_heading("3.4. Phần nào thật sự gọi AI?", level=2)
    add_code(doc, '''if (process.env.GROQ_API_KEY) {
  const openai = new OpenAI({
    baseURL: "https://api.groq.com/openai/v1",
    apiKey: process.env.GROQ_API_KEY,
  });

  const response = await openai.chat.completions.create({
    model: "openai/gpt-oss-120b",
    response_format: { type: "json_object" },
    messages: [systemPrompt, userPayload]
  });
}''', "sales-assistant.service.ts — khối gọi mô hình khoảng dòng 297–349")
    add_plain(doc, "System prompt buộc mô hình chỉ chọn sản phẩm phù hợp trong candidateProducts và trả JSON. Mô hình trả productId, reason, type và confidence. Backend không lấy tên/giá/tồn kho do AI tự nói; nó dùng productId để tìm lại sản phẩm thật trong candidatesToSend.")
    add_flow(doc, [
        "AI trả productId",
        "Backend tìm ID trong candidatesToSend",
        "Không tìm thấy → loại bỏ",
        "Tìm thấy → lấy name/price/stock/image từ Product thật",
        "Chỉ giữ tối đa 5 recommendations",
    ], "Cơ chế chống AI bịa sản phẩm ở POS", [LIGHT_GREEN, LIGHT_BLUE])

    doc.add_heading("3.5. Nếu không có token thì sao?", level=2)
    add_plain(doc, "Service bỏ qua phần gọi mô hình và dùng chính scoredCandidates để tạo tối đa 5 gợi ý. Response có source = HEURISTIC. Nếu token sai, model timeout, JSON lỗi hoặc không có recommendation hợp lệ thì catch cũng đưa luồng về heuristic.")
    add_table(doc, ["Trường hợp", "Kết quả"], [
        ("Token hợp lệ, AI trả hợp lệ", "source = AI"),
        ("Không có GROQ_API_KEY", "source = HEURISTIC"),
        ("Token hết hạn/401", "Ghi cảnh báo và dùng HEURISTIC"),
        ("AI trả ID không tồn tại", "ID bị loại; nếu không còn kết quả thì tiếp tục fallback"),
        ("Không có sản phẩm đúng nhu cầu", "Thông báo không tìm thấy, không cố gợi ý sản phẩm sai"),
    ], [5.2, 9.4])

    doc.add_heading("4. AI TRÊN TRANG KHO HÀNG CHẠY NHƯ THẾ NÀO?", level=1)
    doc.add_heading("4.1. Cấu trúc code sau khi tách service", level=2)
    add_table(doc, ["Bước", "File và vị trí", "Trách nhiệm"], [
        ("1", "frontend/app/(dashboard)/inventory/page.tsx — runAiAnalysis, dòng 94", "Đọc forecastDays và bắt đầu trạng thái loading."),
        ("2", "frontend/services/homex.service.ts — aiForecast, dòng 198", "GET /inventory/ai-forecast?days=N, timeout 120 giây."),
        ("3", "backend/src/routes/inventory.routes.ts — dòng 320–336", "JWT, quyền ADMIN, đọc days và gọi service."),
        ("4", "backend/src/services/inventory-ai.service.ts — forecast, dòng 136", "Lấy dữ liệu, tính toán, gọi AI, fallback và trả kết quả."),
    ], [1.2, 7.1, 6.3], font_size=9.2)
    add_code(doc, '''// inventory.routes.ts — route giờ rất ngắn
const days = Math.max(1, Number(req.query.days || 15));
const result = await inventoryAiService.forecast(days);

return res.json({
  success: true,
  message: "Phân tích kho hàng và đề xuất nhập hàng thành công.",
  data: result,
});''', "Route chỉ điều phối; nghiệp vụ AI nằm trong inventory-ai.service.ts")

    doc.add_heading("4.2. Service lấy dữ liệu gì từ database?", level=2)
    add_bullets(doc, [
        "Danh sách Product ACTIVE: id, sku, name, imageUrl, stockQuantity, minStock, categoryId và category.name.",
        "OrderDetail trong 30 ngày gần nhất.",
        "Chỉ tính đơn COMPLETED hoặc payment PAID.",
        "Loại đơn CANCELLED và payment REFUNDED.",
        "Tách doanh số thành cửa sổ 7 ngày và 30 ngày cho từng sản phẩm, đồng thời tổng hợp theo danh mục.",
    ])

    doc.add_heading("4.3. Hiểu từng con số", level=2)
    add_table(doc, ["Biến", "Cách tính", "Nói dễ hiểu"], [
        ("avgDailySales7", "soldLast7Days / 7", "Mỗi ngày gần đây bán trung bình bao nhiêu."),
        ("avgDailySales30", "soldLast30Days / 30", "Mức bán nền trong một tháng."),
        ("trendRatio", "avg7 / avg30", "> 1 nghĩa là gần đây bán nhanh hơn mức tháng."),
        ("categoryTrendRatio", "avg7 danh mục / avg30 danh mục", "Cả nhóm sản phẩm có đang tăng không."),
        ("seasonBoost", "Luật theo tháng và tên danh mục", "Hệ số nhu cầu mùa nóng hoặc dịp Tết."),
        ("stockCoverageDays", "currentStock / avg7", "Tồn hiện tại đủ bán khoảng bao nhiêu ngày."),
    ], [3.2, 5.2, 6.2])

    doc.add_heading("4.4. Công thức dự báo và ví dụ tính tay", level=2)
    add_callout(doc, "Công thức", "predictedDailySales = avgDailySales30 × trendRatio × categoryTrendRatio × seasonBoost\nexpectedDemand = predictedDailySales × days\nsafetyStock = expectedDemand × 20%\nsuggestedRestock = max(minStock − currentStock, ceil(expectedDemand + safetyStock − currentStock), 0)", LIGHT_BLUE)
    add_table(doc, ["Dữ liệu ví dụ", "Giá trị"], [
        ("Bán 7 ngày", "14 sản phẩm → avg7 = 2/ngày"),
        ("Bán 30 ngày", "30 sản phẩm → avg30 = 1/ngày"),
        ("trendRatio", "2 / 1 = 2"),
        ("Xu hướng danh mục", "1,1"),
        ("Mùa vụ", "1,25"),
        ("Dự báo mỗi ngày", "1 × 2 × 1,1 × 1,25 = 2,75"),
        ("Dự báo 15 ngày", "2,75 × 15 = 41,25"),
        ("Tồn an toàn", "41,25 × 20% = 8,25"),
        ("Tồn hiện tại", "10"),
        ("Nên nhập", "ceil(41,25 + 8,25 − 10) = 40 sản phẩm"),
    ], [6.5, 8.1])
    add_plain(doc, "Nếu sản phẩm không bán trong cả 7 và 30 ngày, service không dùng dự báo nhu cầu để nhập nhiều. Nó chỉ nhập bù đến minStock khi currentStock < minStock; nếu tồn đã đủ minStock thì đề xuất 0.")

    doc.add_heading("4.5. Hệ thống chọn ứng viên như thế nào?", level=2)
    add_plain(doc, "Một sản phẩm được đưa vào candidates nếu có ít nhất một tín hiệu: cần nhập > 0, tồn ≤ minStock, trendRatio ≥ 1,3, danh mục tăng ≥ 1,2, mùa vụ ≥ 1,2 hoặc số ngày tồn có thể bán không vượt quá hai lần thời gian dự phòng. Sau đó danh sách được sắp xếp theo: hết hàng → tồn thấp → sắp cạn theo ngày → bán 30 ngày cao → xu hướng cao → số lượng cần nhập.")
    add_plain(doc, "Service chỉ lấy 40 ứng viên đầu cho mô hình và cũng giới hạn fallbackRestockList ở 40 sản phẩm.")

    doc.add_heading("4.6. Phần nào là AI, phần nào là toán học?", level=2)
    add_table(doc, ["Xử lý", "Thuật toán cục bộ", "AI"], [
        ("Lấy doanh số và tồn", "Có", "Không"),
        ("Tính avg7, avg30, trend, mùa vụ", "Có", "Không"),
        ("Tính expectedDemand và số lượng fallback", "Có", "Không"),
        ("Lọc/sắp xếp top 40", "Có", "Không"),
        ("Viết overview, reason, risks, actionPlan", "Có bản fallback mẫu", "Có thể diễn giải tự nhiên hơn"),
        ("Phân loại và suggestedRestockQuantity cuối", "Có giá trị nền", "Code hiện cho phép AI trả lại các giá trị này"),
        ("Ghép tên/ảnh/mùa vụ chính thức", "Có", "Không"),
    ], [5.0, 4.7, 4.9])
    add_callout(doc, "Cách nói chính xác", "Dự báo Kho là hệ thống lai: công thức và luật nghiệp vụ tạo dự báo nền; mô hình ngôn ngữ phân tích/diễn giải danh sách ứng viên. Không nên nói toàn bộ số lượng nhập là do AI tự học từ dữ liệu.", LIGHT_YELLOW)

    doc.add_heading("4.7. Gọi mô hình và fallback", level=2)
    add_code(doc, '''if (process.env.GROQ_API_KEY && topCandidates.length > 0) {
  const response = await openai.chat.completions.create({
    model: "openai/gpt-oss-120b",
    response_format: { type: "json_object" },
    messages: [
      { role: "system", content: systemPrompt },
      { role: "user", content: JSON.stringify(userPrompt) }
    ]
  });
}''', "inventory-ai.service.ts — khối gọi mô hình khoảng dòng 391–459")
    add_plain(doc, "Trước khi gọi model, fallbackResponse đã được tạo sẵn. Nếu token thiếu, topCandidates rỗng, API lỗi, parse JSON lỗi hoặc Zod không hợp lệ thì finalResult vẫn giữ fallbackResponse. Đây là lý do màn hình Kho vẫn hoạt động khi AI ngoài hệ thống không khả dụng.")
    add_plain(doc, "Sau khi nhận kết quả, service chỉ giữ SKU có trong topCandidates. Tiếp theo service ghép productId, name, imageUrl và thông tin mùa vụ từ forecastList. Nếu AI gắn SEASONAL_HOT nhưng không có seasonName/seasonReason thật, code hạ loại xuống RISING_TREND hoặc CATEGORY_MOMENTUM.")

    doc.add_heading("5. SO SÁNH POS VÀ KHO ĐỂ KHÔNG NHẦM", level=1)
    add_table(doc, ["Điểm", "POS", "Kho hàng"], [
        ("Người sử dụng", "ADMIN/CASHIER", "ADMIN"),
        ("Input chính", "Nhu cầu, ngân sách, giỏ, sở thích", "Số ngày dự phòng"),
        ("Dữ liệu DB", "Sản phẩm còn hàng", "Sản phẩm + OrderDetail 7/30 ngày"),
        ("Tiền xử lý", "Lọc ngân sách và chấm relevance", "Công thức dự báo, xu hướng, mùa vụ"),
        ("Số ứng viên gửi AI", "Tối đa 25", "Tối đa 40"),
        ("AI trả", "Sản phẩm, lý do, combo, mẹo", "Overview, phân loại, số lượng, rủi ro, kế hoạch"),
        ("Fallback", "HEURISTIC", "MATH/RULE fallbackResponse"),
        ("Hành động cuối", "Thu ngân thêm vào giỏ", "ADMIN duyệt/từ chối đề xuất"),
        ("Service", "sales-assistant.service.ts", "inventory-ai.service.ts"),
    ], [3.5, 5.6, 5.5])

    doc.add_heading("6. CÁCH DEBUG KHI AI KHÔNG CHẠY", level=1)
    doc.add_heading("6.1. Kiểm tra chung", level=2)
    add_numbered(doc, [
        "Kiểm tra backend có GROQ_API_KEY hay không. Không in giá trị token ra log.",
        "Kiểm tra request có JWT hợp lệ và đúng role.",
        "Mở Network tab xem endpoint, status code và response data.",
        "Xem terminal backend: lỗi 401 token, timeout, parse JSON hoặc Zod.",
        "Phân biệt ‘AI ngoài không chạy’ với ‘chức năng hỏng’: nếu fallback trả dữ liệu thì chức năng vẫn đang chạy đúng thiết kế.",
    ])
    doc.add_heading("6.2. Checklist POS", level=2)
    add_table(doc, ["Triệu chứng", "Mở file/hàm", "Kiểm tra"], [
        ("Bấm nhưng không gửi request", "pos/page.tsx → handleRequestSalesSuggestion", "need có rỗng không; loading/error state."),
        ("400", "sales-assistant.routes.ts", "Payload có đúng Zod schema không."),
        ("403", "auth middleware/route", "Role phải là ADMIN hoặc CASHIER."),
        ("Chỉ hiện fallback", "sales-assistant.service.ts", "GROQ_API_KEY, lỗi model hoặc JSON."),
        ("Gợi ý sai", "calculateRelevance, CORE_TERM_MAP", "Từ khóa cốt lõi, danh mục và mô tả sản phẩm."),
        ("Không thêm vào giỏ", "handleAddSuggestionToCart", "productId có trong products hoặc detail API không."),
    ], [4.0, 5.2, 5.4], font_size=9.2)
    doc.add_heading("6.3. Checklist Kho", level=2)
    add_table(doc, ["Triệu chứng", "Mở file/hàm", "Kiểm tra"], [
        ("Không gửi request", "inventory/page.tsx → runAiAnalysis", "forecastDays và trạng thái loading."),
        ("403", "inventory.routes.ts", "Chỉ ADMIN được gọi."),
        ("Danh sách rỗng", "inventory-ai.service.ts → candidates", "Có OrderDetail hợp lệ? tồn/minStock và ngưỡng lọc?"),
        ("Số lượng lạ", "forecast() → công thức", "sold7, sold30, trendRatio, category ratio, seasonBoost."),
        ("Luôn fallback", "forecast() → GROQ_API_KEY block", "Token/topCandidates/model/Zod."),
        ("Mùa vụ sai", "SEASONAL_RULES/getSeasonalRule", "Tháng hiện tại và từ khóa category.name."),
    ], [4.0, 5.3, 5.3], font_size=9.2)

    doc.add_heading("7. CÁCH TRÌNH BÀY KHI ĐƯỢC HỎI", level=1)
    doc.add_heading("7.1. Câu trả lời POS trong 60 giây", level=2)
    add_callout(doc, "Bài nói mẫu", "Ở trang POS, thu ngân nhập nhu cầu, ngân sách và các ưu tiên. Frontend gửi dữ liệu cùng giỏ hàng tới API sales-assistant. Backend lấy sản phẩm đang hoạt động và còn tồn, lọc ngân sách, chấm độ liên quan rồi chọn tối đa 25 ứng viên. Khi có token, openai/gpt-oss-120b chỉ được chọn trong danh sách này và trả JSON gồm sản phẩm, lý do, combo và mẹo bán hàng. Backend kiểm tra lại productId và lấy giá, tồn kho từ database. Nếu AI lỗi thì hệ thống dùng heuristic, nên tính năng vẫn hoạt động.", LIGHT_GREEN)
    doc.add_heading("7.2. Câu trả lời Kho trong 90 giây", level=2)
    add_callout(doc, "Bài nói mẫu", "Ở trang Kho, quản trị viên chọn số ngày dự phòng. Route gọi inventoryAiService.forecast. Service lấy tồn kho và lịch sử bán hợp lệ trong 7/30 ngày, tính tốc độ bán, xu hướng sản phẩm, xu hướng danh mục, mùa vụ, nhu cầu dự kiến và tồn an toàn. Từ đó hệ thống tạo danh sách cần nhập và fallback toán học. Nếu có token, openai/gpt-oss-120b phân tích tối đa 40 ứng viên, diễn giải lý do, rủi ro và kế hoạch. Đầu ra được kiểm tra bằng Zod và đối chiếu SKU. Nếu AI lỗi thì dùng kết quả toán học. ADMIN vẫn là người duyệt, AI không tự nhập kho.", LIGHT_GREEN)

    doc.add_heading("7.3. Các câu hỏi dễ bị hỏi", level=2)
    add_table(doc, ["Câu hỏi", "Ý trả lời cần nhớ"], [
        ("AI dùng mô hình gì?", "openai/gpt-oss-120b qua Groq; code dùng OpenAI SDK với baseURL Groq API."),
        ("Tại sao gọi là AI nếu có công thức?", "Đây là hybrid AI: luật/toán tạo dữ liệu nền, LLM chọn và diễn giải."),
        ("Mất Internet/token có chạy không?", "Có, nhờ heuristic POS và fallback toán học Kho."),
        ("AI có bịa sản phẩm không?", "Giảm rủi ro bằng candidateProducts và đối chiếu ID/SKU với dữ liệu thật."),
        ("AI có tự nhập kho không?", "Không. Endpoint chỉ phân tích; ADMIN phê duyệt và nghiệp vụ nhập là bước khác."),
        ("Vì sao tách service?", "Route chỉ điều phối; service chứa nghiệp vụ, dễ đọc, test và bảo trì."),
        ("Dữ liệu huấn luyện ở đâu?", "Dự án không tự huấn luyện model; chỉ gửi ngữ cảnh hiện tại cho model có sẵn."),
        ("AI học từ database không?", "Không fine-tune. Backend truy vấn DB rồi đưa một phần dữ liệu vào prompt mỗi lần gọi."),
    ], [6.0, 8.6], font_size=9.2)

    doc.add_heading("8. LỘ TRÌNH ĐỌC CODE TRONG 30 PHÚT", level=1)
    add_table(doc, ["Phút", "Việc cần làm", "Kết quả cần đạt"], [
        ("0–5", "Đọc hai handler frontend", "Biết nút nào gọi hàm nào và payload có gì."),
        ("5–8", "Đọc homex.service.ts", "Nhớ hai endpoint."),
        ("8–12", "Đọc hai route", "Biết quyền và chỗ gọi service."),
        ("12–20", "Đọc sales-assistant.service.ts từ getSuggestions", "Nắm filter → score → AI → validate → heuristic."),
        ("20–28", "Đọc inventory-ai.service.ts từ forecast", "Nắm DB → map doanh số → công thức → candidate → AI/fallback."),
        ("28–30", "Đóng code và tự nói lại", "Nói được mỗi luồng không nhìn tài liệu."),
    ], [1.8, 7.0, 5.8])
    add_bullets(doc, [
        "Từ khóa POS để Ctrl+F: handleRequestSalesSuggestion, getSuggestions, calculateRelevance, candidatesToSend, GROQ_API_KEY, HEURISTIC.",
        "Từ khóa Kho để Ctrl+F: runAiAnalysis, inventoryAiService, forecast, predictedDailySales, fallbackResponse, aiResponseSchema.",
    ])

    doc.add_heading("9. BÀI TỰ KIỂM TRA", level=1)
    add_numbered(doc, [
        "Hãy kể đủ 9 bước của luồng chung mà không nhìn trang 2.",
        "Tại sao POS chỉ gửi tối đa 25 sản phẩm cho AI?",
        "Nếu AI trả productId không thuộc candidatesToSend thì chuyện gì xảy ra?",
        "Với sold7 = 14 và sold30 = 30, hãy tính avg7, avg30 và trendRatio.",
        "Tại sao forecast Kho vẫn chạy khi xóa GROQ_API_KEY?",
        "Khác nhau giữa source AI và HEURISTIC ở POS là gì?",
        "Hãy chỉ đúng file chứa công thức Kho sau khi refactor.",
        "AI có tự ghi Product.stockQuantity không? Vì sao?",
        "Điểm khác nhau giữa dữ liệu huấn luyện và dữ liệu prompt là gì?",
        "Hãy trình bày POS trong 60 giây và Kho trong 90 giây.",
    ])
    add_callout(doc, "Đáp án cốt lõi", "Nếu bạn trả lời được rằng dữ liệu thật được lấy từ Prisma, thuật toán tiền xử lý trước, model chỉ nhận danh sách đã giới hạn, đầu ra được đối chiếu, có fallback và người dùng quyết định cuối — bạn đã nắm đúng kiến trúc AI của dự án.", LIGHT_BLUE)

    doc.add_heading("10. DANH SÁCH FILE CẦN NHỚ", level=1)
    add_table(doc, ["Mức cần nhớ", "File", "Vai trò một câu"], [
        ("Bắt buộc", "backend/src/services/sales-assistant.service.ts", "Bộ não AI của POS."),
        ("Bắt buộc", "backend/src/services/inventory-ai.service.ts", "Bộ não dự báo/AI của Kho."),
        ("Bắt buộc", "backend/src/routes/sales-assistant.routes.ts", "Cổng API POS, Zod và quyền."),
        ("Bắt buộc", "backend/src/routes/inventory.routes.ts", "Cổng API Kho; gọi inventoryAiService."),
        ("Bắt buộc", "frontend/app/(dashboard)/pos/page.tsx", "Nơi người dùng nhập nhu cầu và xem gợi ý."),
        ("Bắt buộc", "frontend/app/(dashboard)/inventory/page.tsx", "Nơi chạy phân tích và duyệt đề xuất."),
        ("Nên biết", "frontend/services/homex.service.ts", "Nối frontend với hai endpoint."),
        ("Nên biết", "frontend/types/domain.ts", "Kiểu request/response của Sales Assistant."),
        ("Phụ trợ", "backend/src/services/barcode-enrichment.service.ts", "AI bù dữ liệu barcode ở trang Sản phẩm."),
    ], [2.3, 7.5, 4.8], font_size=9)

    doc.add_heading("11. ĐIỂM CẦN CẢI TIẾN SAU NÀY", level=1)
    add_bullets(doc, [
        "Kho: kiểm tra AI trả đủ toàn bộ SKU, không chỉ cần ít nhất một SKU hợp lệ.",
        "Kho: tái tính stats từ danh sách cuối và giới hạn mức AI được sửa suggestedRestockQuantity.",
        "POS: dùng Zod cho toàn bộ response AI, đặc biệt type và confidence.",
        "Tạo interface response Kho ở frontend thay cho getData<any>.",
        "Tách model/baseURL thành cấu hình dùng chung thay vì viết lặp.",
        "Thêm unit test cho công thức forecast và relevance scoring; mock lời gọi model trong integration test.",
    ])
    add_callout(doc, "Không cần học thuộc hạn chế", "Chỉ cần nhớ ba ý để thể hiện hiểu sâu: output AI phải được validate; số liệu quan trọng nên do backend ràng buộc; fallback giúp hệ thống không phụ thuộc hoàn toàn vào mô hình.", LIGHT_YELLOW)

    doc.core_properties.title = "Sổ tay nắm cách AI hoạt động trong HomeX POS"
    doc.core_properties.subject = "Tài liệu học hiểu luồng code AI POS và Kho"
    doc.core_properties.author = "Nguyễn Đức Thịnh – 2305CT2084"
    doc.core_properties.keywords = "HomeX POS, AI, POS, inventory, learning guide, openai/gpt-oss-120b"
    doc.core_properties.comments = "Tài liệu học nội bộ, không phải chương báo cáo."
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print("Created AI learning guide")


if __name__ == "__main__":
    build()
