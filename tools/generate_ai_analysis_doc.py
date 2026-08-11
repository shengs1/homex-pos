from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Phan_tich_hoat_dong_AI_trong_HomeX_POS.docx"

NAVY = "17324D"
BLUE = "1F4E78"
TEAL = "0F766E"
GREEN = "15803D"
LIGHT_BLUE = "DCEAF7"
LIGHT_GREEN = "DCFCE7"
LIGHT_YELLOW = "FEF3C7"
LIGHT_RED = "FEE2E2"
LIGHT_GRAY = "F3F4F6"
MID_GRAY = "64748B"
WHITE = "FFFFFF"
BLACK = "111827"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=BLACK, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def add_table(doc, headers, rows, widths=None, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        shade(hdr.cells[i], BLUE)
        set_cell_text(hdr.cells[i], h, bold=True, color=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_margins(hdr.cells[i])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        if ridx % 2:
            for c in cells:
                shade(c, "F8FAFC")
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
            set_cell_margins(cells[i])
    if widths:
        set_col_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_field(run, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Trang ")
    add_field(run, "PAGE")
    run.add_text(" / ")
    add_field(run, "NUMPAGES")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MID_GRAY)


def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(1.0)

    for style_name, size, color in (("Title", 24, NAVY), ("Heading 1", 16, NAVY), ("Heading 2", 14, BLUE), ("Heading 3", 12.5, TEAL)):
        st = doc.styles[style_name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.first_line_indent = Cm(0)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.space_before = Pt(12 if style_name != "Title" else 0)
        st.paragraph_format.space_after = Pt(6)

    code = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string("0F172A")
    code.paragraph_format.first_line_indent = Cm(0)
    code.paragraph_format.left_indent = Cm(0.4)
    code.paragraph_format.right_indent = Cm(0.2)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0

    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(10)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MID_GRAY)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)

    for section in doc.sections:
        hp = section.header.paragraphs[0]
        hp.text = "HOMEX POS  |  PHÂN TÍCH ỨNG DỤNG AI"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for r in hp.runs:
            r.font.name = "Arial"
            r.font.size = Pt(8)
            r.font.bold = True
            r.font.color.rgb = RGBColor.from_string(MID_GRAY)
        add_page_number(section.footer.paragraphs[0])
    return doc


def add_plain(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.7 + level * 0.5)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, 140, 180, 140, 180)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.first_line_indent = Cm(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p2.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc, code_text, caption=None):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, "EEF2F6")
    set_cell_margins(cell, 100, 120, 100, 120)
    cell.text = ""
    for idx, line in enumerate(code_text.strip("\n").splitlines()):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.style = "CodeBlock"
        p.paragraph_format.space_after = Pt(0)
        p.add_run(line)
    if caption:
        doc.add_paragraph(caption, style="Caption")


def add_flow(doc, steps, caption, colors=None):
    colors = colors or [LIGHT_BLUE] * len(steps)
    table = doc.add_table(rows=0, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, step in enumerate(steps):
        row = table.add_row()
        shade(row.cells[0], colors[idx % len(colors)])
        set_cell_text(row.cells[0], step, bold=True, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_margins(row.cells[0], 110, 140, 110, 140)
        if idx < len(steps) - 1:
            arrow = table.add_row()
            set_cell_text(arrow.cells[0], "↓", bold=True, color=TEAL, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph(caption, style="Caption")


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("TRƯỜNG/ĐƠN VỊ: ................................................")
    r.bold = True
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("ĐỒ ÁN CƠ SỞ")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("PHÂN TÍCH CÁCH HOẠT ĐỘNG CỦA AI\nÁP DỤNG TRONG DỰ ÁN HOMEX POS")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("Kiến trúc • Luồng dữ liệu • Thuật toán • Prompt • Kiểm soát đầu ra • File code liên quan")
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(MID_GRAY)

    info = doc.add_table(rows=4, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.style = "Table Grid"
    fields = [
        ("Dự án", "HomeX POS – Hệ thống quản lý bán hàng đồ gia dụng"),
        ("Sinh viên", "Nguyễn Đức Thịnh"),
        ("Mã sinh viên / Lớp", "2305CT2084 / CT07PM"),
        ("Ngày cập nhật", "11/08/2026"),
    ]
    for i, (k, v) in enumerate(fields):
        shade(info.rows[i].cells[0], BLUE)
        set_cell_text(info.rows[i].cells[0], k, bold=True, color=WHITE, size=11)
        set_cell_text(info.rows[i].cells[1], v, size=11)
        set_cell_margins(info.rows[i].cells[0], 120, 140, 120, 140)
        set_cell_margins(info.rows[i].cells[1], 120, 140, 120, 140)

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph("TP. Hồ Chí Minh, năm 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.runs[0].italic = True
    doc.add_page_break()


def build_document():
    doc = setup_document()
    add_cover(doc)

    doc.add_heading("MỤC LỤC NỘI DUNG", level=1)
    toc_items = [
        "1. Mục tiêu và phạm vi phân tích",
        "2. Kết luận nhanh về cách dự án dùng AI",
        "3. Kiến trúc tích hợp AI tổng thể",
        "4. Trợ lý gợi ý bán hàng tại POS",
        "5. Phân tích kho và đề xuất nhập hàng",
        "6. Làm giàu dữ liệu sản phẩm bằng mã vạch và AI",
        "7. Cơ chế dùng chung: xác thực, kiểm định, fallback và an toàn",
        "8. Danh mục file code liên quan",
        "9. Đánh giá ưu điểm, giới hạn và rủi ro",
        "10. Đề xuất cải tiến",
        "11. Kịch bản kiểm thử và minh chứng",
        "12. Kết luận",
        "Phụ lục A. Đoạn code trọng tâm",
        "Phụ lục B. Thuật ngữ",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.4 if item[0].isdigit() else 0.8)
        p.paragraph_format.space_after = Pt(3)
    add_callout(doc, "Phạm vi phiên bản", "Nội dung được đối chiếu trực tiếp với mã nguồn trong workspace HomeX POS tại ngày 11/08/2026. Nếu code thay đổi sau thời điểm này, cần cập nhật lại vị trí dòng và mô tả thuật toán.", LIGHT_YELLOW)
    doc.add_page_break()

    doc.add_heading("1. MỤC TIÊU VÀ PHẠM VI PHÂN TÍCH", level=1)
    add_plain(doc, "Tài liệu này giải thích rõ AI được tích hợp vào HomeX POS như thế nào, dữ liệu đi qua những lớp nào, mô hình nhận gì và trả về gì, hệ thống kiểm tra kết quả ra sao, cũng như các file code trực tiếp tham gia vào từng luồng. Trọng tâm là hành vi đang chạy trong mã nguồn, không chỉ là mô tả chức năng trên giao diện.")
    add_plain(doc, "Ba chức năng được phân tích gồm: (1) Trợ lý gợi ý bán hàng tại POS; (2) Phân tích kho và đề xuất nhập hàng; (3) Tra cứu mã vạch, kết hợp dữ liệu ngoài và AI để điền các trường còn thiếu của sản phẩm.")
    add_callout(doc, "Ranh giới trách nhiệm", "AI trong dự án là lớp hỗ trợ ra quyết định. AI không trực tiếp hoàn tất đơn hàng, không tự ý tạo sản phẩm và không tự nhập kho. Thu ngân hoặc quản trị viên vẫn là người xem, chọn và xác nhận hành động nghiệp vụ.", LIGHT_GREEN)

    doc.add_heading("2. KẾT LUẬN NHANH VỀ CÁCH DỰ ÁN DÙNG AI", level=1)
    add_plain(doc, "HomeX POS sử dụng kiến trúc lai (hybrid): dữ liệu thật và quy tắc nghiệp vụ được xử lý trước, mô hình ngôn ngữ chỉ nhận một tập ngữ cảnh đã thu gọn, sau đó đầu ra AI tiếp tục được parse và đối chiếu. Khi không có GROQ_API_KEY, token sai hoặc dịch vụ mô hình lỗi, hai chức năng gợi ý bán hàng và dự báo kho vẫn có kết quả nhờ thuật toán fallback.")
    add_table(doc,
        ["Chức năng", "Vai trò của thuật toán cục bộ", "Vai trò của mô hình AI", "Fallback"],
        [
            ("Gợi ý bán hàng", "Lọc hàng còn tồn; lọc ngân sách; chấm điểm độ liên quan, ưu đãi, bảo hành, tồn kho, mua kèm.", "Chọn tối đa 5 sản phẩm trong danh sách ứng viên và sinh lý do, combo, mẹo tư vấn.", "HEURISTIC – dùng chính điểm số cục bộ."),
            ("Dự báo nhập kho", "Tổng hợp bán 7/30 ngày; tính xu hướng, mùa vụ, tồn an toàn và số lượng nhập.", "Diễn giải nghiệp vụ, phân loại ưu tiên/rủi ro/kế hoạch và có thể điều chỉnh đề xuất trên tối đa 40 ứng viên.", "Công thức toán học và mô tả mẫu."),
            ("Làm giàu mã vạch", "Ưu tiên DB; gọi nhiều nguồn barcode; chuẩn hóa, ghép dữ liệu, bảo toàn dữ liệu tin cậy.", "Đọc nhãn từ ảnh hoặc bù trường còn thiếu khi đã có tên sản phẩm đáng tin cậy.", "Trả dữ liệu nguồn ngoài hoặc yêu cầu nhập thủ công."),
        ], [3.0, 4.2, 4.4, 3.2])
    add_callout(doc, "Công nghệ mô hình đang dùng", "Backend dùng thư viện openai, nhưng trỏ baseURL tới Groq/Groq API (https://api.groq.com/openai/v1), model openai/gpt-oss-120b và khóa môi trường GROQ_API_KEY. Vì vậy tài liệu cũ ghi Gemini không còn khớp với code hiện tại.", LIGHT_YELLOW)

    doc.add_heading("3. KIẾN TRÚC TÍCH HỢP AI TỔNG THỂ", level=1)
    add_flow(doc, [
        "Người dùng thao tác trên Next.js (POS / Kho / Sản phẩm)",
        "Frontend service gửi HTTP request kèm JWT",
        "Express route xác thực vai trò và kiểm tra dữ liệu",
        "Prisma lấy dữ liệu thật từ PostgreSQL + thuật toán tiền xử lý",
        "Backend gọi Groq (openai/gpt-oss-120b) khi có GROQ_API_KEY",
        "Parse JSON → kiểm định/đối chiếu → fallback nếu lỗi",
        "Frontend hiển thị gợi ý để người dùng duyệt và thực hiện nghiệp vụ",
    ], "Hình 1. Luồng tích hợp AI tổng quát trong HomeX POS", [LIGHT_BLUE, LIGHT_GREEN])
    add_table(doc, ["Lớp", "Thành phần", "Trách nhiệm"], [
        ("Giao diện", "Next.js/TypeScript", "Thu thập nhu cầu, ngân sách, thời gian dự phòng, mã vạch; hiển thị gợi ý và trạng thái nguồn AI/fallback."),
        ("API", "Express route + Zod + JWT/RBAC", "Bảo vệ endpoint, kiểm tra request, điều phối service."),
        ("Dữ liệu", "Prisma + PostgreSQL", "Cung cấp sản phẩm, tồn kho, danh mục, nhà cung cấp, đơn hàng và chi tiết bán."),
        ("Tiền xử lý", "TypeScript rule/math", "Lọc, chấm điểm, tính xu hướng/mùa vụ, chuẩn hóa dữ liệu barcode."),
        ("Suy luận", "Groq – openai/gpt-oss-120b", "Chọn/diễn giải đề xuất và bù trường dữ liệu thiếu theo prompt."),
        ("Hậu kiểm", "JSON parse + Zod/domain validation", "Loại ID/SKU không có thật, giới hạn kết quả, ghép lại thông tin chính thức."),
    ], [2.2, 4.2, 8.2])

    doc.add_heading("4. TRỢ LÝ GỢI Ý BÁN HÀNG TẠI POS", level=1)
    doc.add_heading("4.1. Mục tiêu nghiệp vụ", level=2)
    add_plain(doc, "Chức năng hỗ trợ thu ngân tìm sản phẩm phù hợp theo mô tả tự nhiên của khách, khoảng ngân sách, giỏ hàng hiện tại và bốn ưu tiên: khuyến mãi, bảo hành, tồn kho cao, mua chéo từ giỏ. Kết quả gồm tóm tắt, tối đa 5 sản phẩm, lý do, loại gợi ý, độ tin cậy, combo và mẹo tư vấn.")
    doc.add_heading("4.2. Luồng xử lý chi tiết", level=2)
    add_numbered(doc, [
        "Thu ngân mở hộp thoại AI trên trang POS, nhập nhu cầu hoặc chọn nhu cầu nhanh, đặt ngân sách và tùy chọn.",
        "Frontend ghép nhu cầu nhanh với nội dung tự nhập, lấy customerId và cartItems rồi gọi POST /api/pos/sales-assistant.",
        "Route yêu cầu JWT, chỉ cho ADMIN hoặc CASHIER và dùng Zod kiểm tra cấu trúc payload.",
        "Service lấy các sản phẩm ACTIVE có stockQuantity > 0; sau đó lọc theo budgetMin/budgetMax.",
        "Hệ thống chuẩn hóa tiếng Việt, dò cụm từ cốt lõi và chấm điểm độ liên quan. Sản phẩm phụ kiện bị trừ điểm mạnh nếu khách đang tìm thiết bị chính.",
        "Điểm bổ sung đến từ khuyến mãi, thời gian bảo hành, tồn kho cao và cùng danh mục với hàng trong giỏ. Tối đa 25 ứng viên được gửi cho mô hình.",
        "Nếu có GROQ_API_KEY, backend yêu cầu openai/gpt-oss-120b chỉ trả JSON và chỉ chọn productId thuộc candidateProducts.",
        "Backend parse JSON, tìm lại từng productId trong danh sách ứng viên, lấy tên/giá/tồn kho chính thức từ database và chỉ giữ tối đa 5 kết quả.",
        "Nếu AI lỗi, token thiếu/sai hoặc đầu ra không dùng được, hệ thống chạy heuristic và trả source = HEURISTIC.",
        "Thu ngân có thể bấm thêm sản phẩm gợi ý vào giỏ; đây mới là bước làm thay đổi phiên bán hàng.",
    ])
    doc.add_heading("4.3. Cơ chế chấm điểm trước AI", level=2)
    add_table(doc, ["Tín hiệu", "Điều kiện", "Tác động điểm"], [
        ("Khớp nhu cầu", "relevance > 0", "+1000 + relevance, bảo đảm sản phẩm đúng nhu cầu đứng trước."),
        ("Tên/danh mục/mô tả", "Khớp cụm từ cốt lõi", "+35 / +20 / +15 cho mỗi term; khớp trực tiếp còn cộng thêm."),
        ("Khuyến mãi", "originalPrice > salePrice", "+25 khi ưu tiên, nếu không +10."),
        ("Bảo hành", "warrantyMonths > 0", "+20 khi ưu tiên, nếu không +5."),
        ("Tồn kho cao", "stockQuantity > 10", "+15 khi ưu tiên, nếu không +5."),
        ("Mua chéo", "Cùng category với hàng trong giỏ", "+30 khi ưu tiên, nếu không +10."),
        ("Phụ kiện không đúng ý", "Tên là phụ kiện nhưng nhu cầu không tìm phụ kiện", "−300 để giảm gợi ý sai."),
    ], [3.0, 6.0, 5.6])
    add_code(doc, '''const response = await openai.chat.completions.create({
  model: "openai/gpt-oss-120b",
  response_format: { type: "json_object" },
  messages: [
    { role: "system", content: systemPrompt },
    { role: "user", content: JSON.stringify(userPayload) }
  ]
});''', "Mã 1. Gọi mô hình ở backend/src/services/sales-assistant.service.ts:342–349")
    doc.add_heading("4.4. Vì sao kết quả không dễ ‘bịa’ sản phẩm", level=2)
    add_plain(doc, "Mô hình không được truy cập toàn bộ database và không tự tạo bản ghi. Nó chỉ thấy tối đa 25 candidateProducts. Khi nhận kết quả, backend không tin trực tiếp tên, giá hoặc tồn kho do AI trả về; backend chỉ nhận productId rồi đối chiếu lại với candidatesToSend. ID lạ bị bỏ qua. Các trường name, price, stockQuantity và imageUrl được lấy lại từ đối tượng Product thật.")
    add_callout(doc, "Điểm cần lưu ý", "Trường type và confidence hiện được ép kiểu/chuyển số nhưng chưa kiểm tra enum/range bằng Zod ở service. Đây là điểm nên siết chặt để tránh type lạ hoặc confidence ngoài 0–1.", LIGHT_RED)

    doc.add_heading("5. PHÂN TÍCH KHO VÀ ĐỀ XUẤT NHẬP HÀNG", level=1)
    doc.add_heading("5.1. Dữ liệu đầu vào", level=2)
    add_plain(doc, "Endpoint GET /api/inventory/ai-forecast?days=N chỉ dành cho ADMIN. Backend lấy toàn bộ sản phẩm ACTIVE cùng tồn hiện tại, tồn tối thiểu và danh mục; đồng thời lấy OrderDetail trong 30 ngày gần nhất. Chỉ đơn hoàn thành hoặc đã thanh toán được tính, còn đơn hủy hoặc hoàn tiền bị loại.")
    add_table(doc, ["Biến", "Ý nghĩa", "Nguồn"], [
        ("soldLast7Days", "Tổng số lượng bán 7 ngày gần nhất", "OrderDetail hợp lệ"),
        ("soldLast30Days", "Tổng số lượng bán 30 ngày gần nhất", "OrderDetail hợp lệ"),
        ("avgDailySales7/30", "Tốc độ bán trung bình ngày", "Số bán chia 7 hoặc 30"),
        ("trendRatio", "Mức thay đổi tốc độ bán gần đây", "avgDailySales7 / avgDailySales30"),
        ("categoryTrendRatio", "Xu hướng chung của danh mục", "Doanh số danh mục 7 ngày so với 30 ngày"),
        ("seasonBoost", "Hệ số mùa vụ theo tháng và từ khóa danh mục", "SEASONAL_RULES"),
        ("stockCoverageDays", "Số ngày tồn kho có thể đáp ứng", "currentStock / avgDailySales7"),
    ], [3.4, 6.4, 4.8])
    doc.add_heading("5.2. Công thức dự báo", level=2)
    add_callout(doc, "Công thức đang dùng trong code", "predictedDailySales = avgDailySales30 × trendRatio × categoryTrendRatio × seasonBoost\nexpectedDemand = predictedDailySales × days\nsafetyStock = expectedDemand × 20%\nsuggestedRestock = max(minStock − currentStock, ceil(expectedDemand + safetyStock − currentStock), 0)", LIGHT_BLUE)
    add_plain(doc, "Do trendRatio = avgDailySales7 / avgDailySales30 khi mẫu số lớn hơn 0, phần avgDailySales30 × trendRatio thực chất đưa tốc độ dự báo về gần avgDailySales7. Sau đó hệ thống nhân thêm xu hướng danh mục và mùa vụ. Với sản phẩm không bán trong cả 7 và 30 ngày, hệ thống chỉ nhập bù đến minStock nếu tồn hiện tại thấp hơn mức tối thiểu.")
    doc.add_heading("5.3. Phân loại tín hiệu", level=2)
    add_table(doc, ["Loại", "Điều kiện chính", "Ý nghĩa"], [
        ("SEASONAL_HOT", "Có luật mùa vụ và có doanh số", "Sản phẩm đang có tín hiệu bán trong mùa cao điểm."),
        ("RISING_TREND", "trendRatio ≥ 1,25 hoặc categoryTrendRatio ≥ 1,2 và có bán", "Tốc độ bán gần đây tăng."),
        ("SEASONAL_WATCH", "Có luật mùa vụ nhưng chưa có doanh số", "Theo dõi tiềm năng, chưa khẳng định nóng."),
        ("CATEGORY_MOMENTUM", "Danh mục tăng trưởng và có bán", "Động lượng đến từ nhóm sản phẩm."),
        ("LOW_STOCK", "currentStock ≤ minStock hoặc hết hàng", "Cảnh báo tồn kho thấp."),
        ("NO_SIGNAL", "Không rơi vào điều kiện trên", "Chưa có tín hiệu rõ."),
    ], [3.3, 6.2, 5.1])
    doc.add_heading("5.4. Phần AI và phần fallback", level=2)
    add_plain(doc, "Từ danh sách đã lọc và sắp xếp, backend lấy tối đa 40 ứng viên. Prompt yêu cầu mô hình trả đủ tất cả SKU, phân loại, số lượng nhập, mức ưu tiên, độ tin cậy, lý do, rủi ro và kế hoạch hành động bằng tiếng Việt nghiệp vụ. Đầu ra được parse rồi kiểm tra bằng aiResponseSchema (Zod). Sau đó chỉ giữ item có SKU thật trong topCandidates.")
    add_plain(doc, "Nếu không gọi được AI, fallbackResponse đã được chuẩn bị sẵn từ công thức. Vì vậy tính năng vẫn hoạt động ngoại tuyến đối với dịch vụ mô hình. Cuối cùng, backend ghép name, imageUrl và thông tin mùa vụ từ forecastList, đồng thời hạ SEASONAL_HOT nếu thiếu bằng chứng mùa vụ.")
    add_callout(doc, "Rủi ro kiểm định hiện tại", "Prompt yêu cầu AI trả đủ 100% ứng viên, nhưng code chỉ kiểm tra filteredList.length > 0. Nếu mô hình trả thiếu SKU, hệ thống vẫn có thể chấp nhận danh sách thiếu. stats do AI trả cũng chưa được đối chiếu lại với danh sách cuối. Nên bổ sung kiểm tra tập SKU bằng nhau và tái tính stats ở backend.", LIGHT_RED)
    doc.add_heading("5.5. Quyền quyết định của người dùng", level=2)
    add_plain(doc, "Trang kho chỉ hiển thị phân tích và cho ADMIN phê duyệt hoặc từ chối từng gợi ý. Việc phê duyệt đưa item vào danh sách chuẩn bị nhập; bản thân endpoint ai-forecast không cập nhật Product.stockQuantity và không tạo StockTransaction. Tách biệt này giúp hạn chế AI tự động gây thay đổi tồn kho.")

    doc.add_heading("6. LÀM GIÀU DỮ LIỆU SẢN PHẨM BẰNG MÃ VẠCH VÀ AI", level=1)
    doc.add_heading("6.1. Mục tiêu và chiến lược ưu tiên nguồn", level=2)
    add_plain(doc, "Khi ADMIN nhập hoặc quét mã vạch dài từ 8 ký tự, frontend gọi POST /api/products/enrich. Chức năng ưu tiên dữ liệu có độ tin cậy cao hơn trước, sau đó AI chỉ bù phần còn thiếu. Đây là hướng Retrieval + Enrichment, không phải yêu cầu AI tự nghĩ toàn bộ sản phẩm từ một mã số.")
    add_flow(doc, [
        "1. Tìm barcode trong database nội bộ → nếu có, trả source DATABASE (confidence 1.0)",
        "2. Gọi tuần tự: UPCItemDB → Barcode Spider → Barcode Lookup → Open Products Facts → Open Food Facts",
        "3. Ghép dữ liệu ngoài; dừng sớm khi đã đủ trường và không cần AI",
        "4. Nếu thiếu danh tính nhưng có imageUrl → AI đọc nhãn trong ảnh",
        "5. Nếu có tên đáng tin cậy nhưng thiếu trường → AI bù trường còn thiếu",
        "6. Làm sạch, chuẩn hóa category/supplier, bảo toàn giá trị nguồn ngoài, gắn source HYBRID/AI",
        "7. Trả dữ liệu gợi ý vào form; ADMIN xem và lưu sản phẩm bằng thao tác riêng",
    ], "Hình 2. Luồng tra cứu và làm giàu mã vạch", [LIGHT_GREEN, LIGHT_BLUE])
    doc.add_heading("6.2. AI đọc ảnh và AI bù trường", level=2)
    add_table(doc, ["Nhánh", "Điều kiện gọi", "Dữ liệu gửi mô hình", "Kết quả"], [
        ("Đọc nhãn ảnh", "Không có dữ liệu nhận diện nhưng có imageUrl", "Barcode + URL ảnh; yêu cầu chỉ dựa trên chữ/nhãn nhìn thấy", "Tên, hãng, danh mục, đơn vị, giá demo, tồn/minStock, bảo hành, mô tả"),
        ("Bù trường thiếu", "Có GROQ_API_KEY, còn thiếu trường, tên sản phẩm đáng tin cậy", "Dữ liệu external + danh sách supplier đang ACTIVE + tên các field thiếu", "Chỉ gợi ý phần còn thiếu, tránh mâu thuẫn dữ liệu external"),
    ], [2.8, 4.0, 4.3, 3.5])
    doc.add_heading("6.3. Hậu kiểm và ghép dữ liệu", level=2)
    add_bullets(doc, [
        "cleanString giới hạn chiều dài; cleanImageUrl chỉ nhận HTTP/HTTPS; giá/số lượng được chuyển thành số nguyên không âm.",
        "category và supplier được đối chiếu với dữ liệu ACTIVE hiện có để tránh tạo tên tự do không khớp hệ thống.",
        "mergeEnrichedData ưu tiên externalValue; AI chỉ được dùng khi giá trị external trống.",
        "sanitizeAiDataAgainstExternal và applyAiOperationalSuggestions tiếp tục hạn chế xung đột và chuẩn hóa gợi ý vận hành.",
        "Nếu tất cả nguồn đều không có dữ liệu hữu ích, service trả null; giao diện yêu cầu nhập thủ công, không cố đoán từ barcode trống ngữ cảnh.",
    ])
    add_callout(doc, "Lưu ý về giá", "Prompt hiện cố ý giới hạn giá demo 1.000–9.500 VND và ước lượng giá nhập khoảng 70% giá bán. Đây là quy tắc phục vụ dữ liệu trình diễn, không phù hợp triển khai thực tế nếu chưa đổi sang bảng giá/đơn vị tiền chuẩn.", LIGHT_YELLOW)

    doc.add_heading("7. CƠ CHẾ DÙNG CHUNG: XÁC THỰC, KIỂM ĐỊNH, FALLBACK VÀ AN TOÀN", level=1)
    add_table(doc, ["Cơ chế", "Cách áp dụng", "Giá trị"], [
        ("JWT + RBAC", "Sales Assistant: ADMIN/CASHIER; Forecast và Product Enrich: ADMIN.", "Giới hạn người được gọi AI và xem dữ liệu nghiệp vụ."),
        ("Biến môi trường", "GROQ_API_KEY chỉ đọc ở backend; không gửi xuống frontend.", "Giữ bí mật khóa truy cập mô hình."),
        ("JSON mode", "response_format = json_object và prompt yêu cầu JSON thuần.", "Giảm lỗi parse so với văn bản tự do."),
        ("Domain grounding", "Chỉ gửi candidates/SKU/supplier có thật; đối chiếu lại ID/SKU.", "Giảm hallucination và bảo vệ tính đúng của dữ liệu."),
        ("Fallback", "Heuristic cho POS; toán học cho kho; dữ liệu ngoài/nhập tay cho barcode.", "Hệ thống không phụ thuộc tuyệt đối vào dịch vụ AI."),
        ("Human-in-the-loop", "Người dùng thêm vào giỏ, duyệt nhập hoặc lưu form.", "AI tư vấn nhưng không tự ghi nghiệp vụ quan trọng."),
        ("Timeout", "Frontend forecast cho phép 120 giây; external barcode có timeout 35 giây mỗi nguồn.", "Tránh chờ vô hạn, dù tổng thời gian barcode có thể còn dài."),
    ], [3.0, 7.0, 4.6])
    add_plain(doc, "Về bản chất, mức an toàn tốt nhất của dự án không đến từ prompt mà đến từ ràng buộc bằng code: truy vấn database trước, thu hẹp ứng viên, kiểm tra ID/SKU, ưu tiên dữ liệu external và yêu cầu người dùng xác nhận. Prompt chỉ là một lớp hướng dẫn bổ sung.")

    doc.add_heading("8. DANH MỤC FILE CODE LIÊN QUAN", level=1)
    add_table(doc, ["Nhóm", "File / vị trí chính", "Vai trò"], [
        ("Cấu hình API", "backend/src/index.ts:27, 100–122", "Nạp route và mount /api/products, /api/inventory, /api/pos."),
        ("POS route", "backend/src/routes/sales-assistant.routes.ts:10–43", "Zod schema, JWT/RBAC, endpoint POST /api/pos/sales-assistant."),
        ("POS AI service", "backend/src/services/sales-assistant.service.ts:1–475", "Tìm sản phẩm, chấm điểm, prompt, gọi model, hậu kiểm và heuristic fallback."),
        ("Kho AI", "backend/src/routes/inventory.routes.ts:199–318; 394–878", "Luật mùa vụ; truy vấn bán hàng; công thức; prompt; Zod; fallback và response."),
        ("Barcode route", "backend/src/routes/product.routes.ts:959–998", "Endpoint POST /api/products/enrich dành cho ADMIN."),
        ("Barcode service", "backend/src/services/barcode-enrichment.service.ts:1–825", "Tra cứu DB/API ngoài, chuẩn hóa, AI ảnh, AI bù thiếu và merge."),
        ("Frontend service", "frontend/services/homex.service.ts:150–174; 193–199; 378–382", "Khai báo ba lời gọi API từ frontend."),
        ("POS UI", "frontend/app/(dashboard)/pos/page.tsx:213–305; 2617–2923", "Thu thập payload, gọi AI, hiển thị và thêm gợi ý vào giỏ."),
        ("Kho UI", "frontend/app/(dashboard)/inventory/page.tsx:67–132; 608–913", "Chọn số ngày, chạy phân tích, lọc, duyệt/từ chối đề xuất."),
        ("Sản phẩm UI", "frontend/app/(dashboard)/products/page.tsx:500–542; 1335–1370", "Gọi enrich khi quét/nhập barcode và áp dữ liệu vào form."),
        ("Kiểu dữ liệu", "frontend/types/domain.ts:419–454", "Request/response của Sales Assistant."),
        ("Đa ngôn ngữ", "frontend/contexts/language-context.tsx:65 trở đi", "Nhãn, thông báo và trạng thái của trợ lý bán hàng/barcode."),
        ("Dependency", "backend/package.json", "openai ^6.45.0, zod ^4.4.3, Prisma và Express."),
        ("Benchmark", "backend/scripts/benchmark-sales-ai.ts; benchmark-inventory-ai.ts; benchmark-ai-images.ts", "Đo/kiểm tra các luồng AI trong môi trường phát triển."),
    ], [2.5, 6.2, 5.9], font_size=8.7)

    doc.add_heading("9. ĐÁNH GIÁ ƯU ĐIỂM, GIỚI HẠN VÀ RỦI RO", level=1)
    doc.add_heading("9.1. Ưu điểm", level=2)
    add_bullets(doc, [
        "Kiến trúc hybrid làm giảm phụ thuộc mô hình và vẫn trả kết quả khi mất dịch vụ AI.",
        "Dữ liệu đưa vào mô hình đã được giới hạn theo nghiệp vụ, giúp giảm chi phí và giảm gợi ý sai.",
        "Hậu kiểm ID/SKU và ưu tiên dữ liệu database/external hạn chế hallucination tác động vào dữ liệu thật.",
        "Phân quyền phù hợp: thu ngân dùng trợ lý bán hàng; chức năng kho và tạo dữ liệu sản phẩm dành cho quản trị viên.",
        "Người dùng giữ quyền quyết định cuối cùng đối với giỏ hàng, phiếu nhập và lưu sản phẩm.",
    ])
    doc.add_heading("9.2. Giới hạn/rủi ro hiện tại", level=2)
    add_table(doc, ["Mức", "Vấn đề", "Tác động"], [
        ("Cao", "Inventory chưa kiểm tra AI trả đủ toàn bộ SKU; stats không tái tính.", "Có thể thiếu đề xuất hoặc số liệu tổng quan không khớp."),
        ("Cao", "AI kho có thể trả suggestedRestockQuantity khác công thức mà chưa có biên giới hạn theo dữ liệu gốc.", "Đề xuất nhập quá cao/thấp nếu mô hình suy luận bất thường."),
        ("Trung bình", "Sales Assistant chưa Zod-validate toàn bộ response; type/confidence chưa giới hạn chặt.", "UI có thể nhận kiểu không chuẩn hoặc độ tin cậy không hợp lệ."),
        ("Trung bình", "Luật mùa vụ dựa vào từ khóa danh mục và tháng cố định.", "Khó phản ánh vùng địa lý, lịch lễ thay đổi hoặc xu hướng đột biến."),
        ("Trung bình", "Chuỗi gọi barcode external chạy tuần tự, mỗi nguồn timeout 35 giây.", "Trường hợp xấu có thể chờ lâu; trải nghiệm quét bị chậm."),
        ("Trung bình", "Giá do AI gợi ý bị giới hạn theo dữ liệu demo.", "Không dùng được làm giá thực tế nếu không cấu hình lại."),
        ("Thấp", "Tài liệu cũ nhắc Gemini nhưng code dùng Groq.", "Dễ gây nhầm khi thuyết trình/bảo trì."),
    ], [1.5, 7.4, 5.7], font_size=9.2)

    doc.add_heading("10. ĐỀ XUẤT CẢI TIẾN", level=1)
    add_numbered(doc, [
        "Tạo schema Zod dùng chung cho mọi response AI; kiểm tra enum, confidence 0–1, số không âm, độ dài chuỗi và giới hạn số phần tử.",
        "Ở dự báo kho, so sánh Set(SKU AI) với Set(SKU topCandidates); nếu thiếu/thừa thì fallback toàn bộ hoặc bù bằng kết quả toán học. Tái tính stats từ danh sách cuối.",
        "Giới hạn số lượng nhập AI trong một biên quanh suggestedRestockQuantity do công thức tạo ra; hoặc chỉ cho AI diễn giải, không cho sửa số lượng.",
        "Ghi log có cấu trúc gồm model, thời gian, nguồn AI/fallback, số ứng viên, lỗi parse và mã request; tuyệt đối không log token.",
        "Thêm cache ngắn hạn theo hash của payload và dữ liệu kho để giảm gọi lặp, chi phí và độ trễ.",
        "Chạy các API barcode độc lập song song có giới hạn concurrency, hoặc áp dụng timeout tổng; tiếp tục giữ thứ tự ưu tiên khi merge.",
        "Chuyển giá demo và luật mùa vụ sang bảng cấu hình/database để ADMIN có thể cập nhật mà không sửa code.",
        "Xây dựng bộ đánh giá định lượng: precision@5 cho gợi ý bán hàng, sai số dự báo MAE/MAPE, tỷ lệ trường barcode được điền đúng, tỷ lệ fallback và latency p95.",
        "Cập nhật tài liệu kiến trúc thống nhất tên Groq/openai/gpt-oss-120b; nếu đổi nhà cung cấp, ghi rõ qua adapter để tránh gắn cứng baseURL ở nhiều file.",
    ])

    doc.add_heading("11. KỊCH BẢN KIỂM THỬ VÀ MINH CHỨNG", level=1)
    add_table(doc, ["Mã", "Kịch bản", "Kết quả mong đợi"], [
        ("AI-POS-01", "Nhu cầu ‘nồi cơm điện’, ngân sách phù hợp, token hợp lệ", "Chỉ trả sản phẩm liên quan, ID tồn tại, tối đa 5, source AI."),
        ("AI-POS-02", "Xóa GROQ_API_KEY", "API vẫn thành công, source HEURISTIC, không có sản phẩm sai nhu cầu."),
        ("AI-POS-03", "AI giả lập trả productId không tồn tại", "ID lạ bị loại; nếu không còn kết quả hợp lệ thì fallback/không trả sản phẩm bịa."),
        ("AI-INV-01", "Sản phẩm bán nhanh 7 ngày và tồn thấp", "Được phân loại xu hướng/thiếu hàng, quantity > 0, có lý do và action plan."),
        ("AI-INV-02", "Không bán 30 ngày, tồn < minStock", "Chỉ đề xuất nhập bù minStock − currentStock."),
        ("AI-INV-03", "Token sai hoặc model timeout", "Trả fallback toán học, endpoint không lỗi 500 do AI."),
        ("AI-INV-04", "AI trả thiếu một SKU", "Sau cải tiến: từ chối response AI hoặc bù đúng bằng fallback."),
        ("AI-BC-01", "Barcode đã có trong DB", "Trả DATABASE, confidence 1, không gọi nguồn ngoài/AI."),
        ("AI-BC-02", "Nguồn ngoài đủ dữ liệu", "Trả nguồn external/HYBRID và không cần gọi AI."),
        ("AI-BC-03", "Có tên tin cậy nhưng thiếu category/giá", "AI chỉ bù trường thiếu; dữ liệu external không bị ghi đè."),
        ("AI-BC-04", "Không nguồn nào biết barcode", "Trả hướng dẫn nhập thủ công, không tự bịa sản phẩm."),
        ("AI-SEC-01", "CASHIER gọi forecast hoặc enrich", "HTTP 403; CASHIER chỉ được gọi Sales Assistant."),
    ], [1.8, 6.8, 6.0], font_size=8.9)
    add_plain(doc, "Các script benchmark hiện có trong backend/scripts có thể dùng làm nền cho kiểm thử hiệu năng/chất lượng. Tuy nhiên cần bổ sung assertion và dữ liệu chuẩn (ground truth) để biến benchmark thành kiểm thử hồi quy đáng tin cậy.")

    doc.add_heading("12. KẾT LUẬN", level=1)
    add_plain(doc, "AI trong HomeX POS được áp dụng đúng vai trò trợ lý: nâng chất lượng gợi ý, diễn giải dữ liệu và giảm thao tác nhập liệu, trong khi dữ liệu nghiệp vụ vẫn do database, thuật toán cục bộ và người dùng kiểm soát. Điểm nổi bật là ba lớp bảo vệ: giới hạn ngữ cảnh trước khi gọi mô hình, kiểm tra/ghép lại kết quả sau khi gọi và fallback khi dịch vụ AI không sẵn sàng.")
    add_plain(doc, "Để sẵn sàng hơn cho triển khai thực tế, ưu tiên tiếp theo nên là siết schema đầu ra, ràng buộc số lượng nhập kho, tái tính thống kê ở backend, tối ưu thời gian barcode và thiết lập bộ đo chất lượng định lượng. Sau các cải tiến này, hệ thống có thể giữ lợi ích của AI mà giảm đáng kể rủi ro kết quả thiếu nhất quán.")

    doc.add_page_break()
    doc.add_heading("PHỤ LỤC A. ĐOẠN CODE TRỌNG TÂM", level=1)
    add_code(doc, '''// Lọc sản phẩm thật mà AI được phép chọn
const product = candidatesToSend.find(p => p.id === Number(rec.productId));
if (product) {
  validatedRecs.push({
    productId: product.id,
    name: product.name,
    price: Number(product.salePrice),
    stockQuantity: product.stockQuantity,
    reason: rec.reason || "Sản phẩm phù hợp với nhu cầu khách hàng."
  });
}''', "Mã A1. Grounding kết quả POS – sales-assistant.service.ts:356–370")
    add_code(doc, '''const predictedDailySales =
  avgDailySales30 * trendRatio * categoryTrendRatio * seasonBoost;
const expectedDemand = predictedDailySales * days;
const safetyStock = expectedDemand * 0.2;
const demandGap = Math.max(
  0,
  Math.ceil(expectedDemand + safetyStock - currentStock)
);''', "Mã A2. Công thức dự báo – inventory.routes.ts:525–535")
    add_code(doc, '''for (const key of ["name", "category", "brand", "supplierName",
  "unit", "estimatedImportPrice", "estimatedSalePrice", "originalPrice",
  "warrantyMonths", "stockQuantity", "minStock", "imageUrl", "description"] as const) {
  const externalValue = external[key];
  const aiValue = ai[key];
  merged[key] = externalValue !== undefined && externalValue !== ""
    ? externalValue
    : aiValue;
}''', "Mã A3. Ưu tiên dữ liệu external, AI chỉ bù thiếu – barcode-enrichment.service.ts:419–439")
    add_code(doc, '''router.post(
  "/sales-assistant",
  authenticateToken,
  authorizeRoles(USER_ROLES.ADMIN, USER_ROLES.CASHIER),
  ...
);

router.get(
  "/ai-forecast",
  authenticateToken,
  authorizeRoles(USER_ROLES.ADMIN),
  ...
);''', "Mã A4. Phân quyền các endpoint AI")

    doc.add_heading("PHỤ LỤC B. THUẬT NGỮ", level=1)
    add_table(doc, ["Thuật ngữ", "Giải thích trong phạm vi dự án"], [
        ("AI/LLM", "Mô hình ngôn ngữ openai/gpt-oss-120b được gọi qua Groq để chọn, diễn giải hoặc bù dữ liệu."),
        ("Prompt", "Chỉ dẫn system + dữ liệu userPayload gửi cho mô hình."),
        ("Grounding", "Buộc kết quả gắn với sản phẩm/SKU/nhà cung cấp có thật trong dữ liệu đầu vào."),
        ("Hallucination", "Mô hình sinh thông tin không có căn cứ; dự án giảm rủi ro bằng đối chiếu ID/SKU và ưu tiên dữ liệu thật."),
        ("Heuristic", "Bộ quy tắc/chấm điểm thủ công dùng để xếp hạng hoặc fallback."),
        ("Fallback", "Kết quả thay thế khi AI không khả dụng hoặc đầu ra không hợp lệ."),
        ("Cross-sell", "Gợi ý sản phẩm mua kèm sản phẩm trong giỏ."),
        ("Up-sell", "Gợi ý lựa chọn có giá trị/tính năng cao hơn phù hợp nhu cầu."),
        ("Safety stock", "Tồn kho an toàn; code hiện lấy 20% nhu cầu dự kiến."),
        ("Human-in-the-loop", "Con người xem và xác nhận trước khi tác động nghiệp vụ."),
    ], [4.0, 10.6])

    # Core properties
    doc.core_properties.title = "Phân tích cách hoạt động của AI áp dụng trong dự án HomeX POS"
    doc.core_properties.subject = "Đồ án cơ sở – HomeX POS"
    doc.core_properties.author = "Nguyễn Đức Thịnh – 2305CT2084"
    doc.core_properties.keywords = "HomeX POS, AI, Groq, openai/gpt-oss-120b, dự báo kho, trợ lý bán hàng, barcode"
    doc.core_properties.comments = "Sinh tự động từ mã nguồn hiện hành và được trình bày theo dạng tài liệu học thuật."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(path)

