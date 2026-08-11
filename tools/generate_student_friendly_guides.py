from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from generate_ai_analysis_doc import (
    BLUE,
    LIGHT_BLUE,
    LIGHT_GREEN,
    LIGHT_RED,
    LIGHT_YELLOW,
    MID_GRAY,
    NAVY,
    add_bullets,
    add_callout,
    add_flow,
    add_numbered,
    add_plain,
    add_table,
    set_cell_margins,
    set_cell_text,
    setup_document,
    shade,
)


ROOT = Path(__file__).resolve().parents[1]
AI_OUTPUT = ROOT / "docs" / "Phan_tich_hoat_dong_AI_trong_HomeX_POS.docx"
WEBSITE_OUTPUT = ROOT / "docs" / "Huong_dan_thuyet_trinh_gioi_thieu_toan_bo_website_HomeX_POS.docx"


def add_simple_cover(doc, title, subtitle, purpose):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("HOMEX POS — TÀI LIỆU DÀNH CHO SINH VIÊN")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(23)
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(MID_GRAY)

    for _ in range(2):
        doc.add_paragraph()

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Mục đích", purpose),
        ("Cách viết", "Dễ hiểu, ít từ chuyên ngành, có câu nói mẫu"),
        ("Sinh viên", "Nguyễn Đức Thịnh — 2305CT2084 — CT07PM"),
        ("Cập nhật", "11/08/2026 — theo code hiện tại"),
    ]
    for i, (key, value) in enumerate(rows):
        shade(table.rows[i].cells[0], BLUE)
        set_cell_text(table.rows[i].cells[0], key, bold=True, color="FFFFFF", size=10.5)
        set_cell_text(table.rows[i].cells[1], value, size=10.5)
        set_cell_margins(table.rows[i].cells[0], 120, 140, 120, 140)
        set_cell_margins(table.rows[i].cells[1], 120, 140, 120, 140)

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph("Đây là tài liệu để học và tập nói, không phải chương viết vào báo cáo.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    doc.add_page_break()


def speech(doc, text, title="Bạn có thể nói với thầy"):
    add_callout(doc, title, text, LIGHT_GREEN)


def remember(doc, text):
    add_callout(doc, "Chỉ cần nhớ", text, LIGHT_YELLOW)


def build_ai_guide():
    doc = setup_document()
    add_simple_cover(
        doc,
        "HƯỚNG DẪN DỄ HIỂU\nCÁCH AI HOẠT ĐỘNG TRONG DỰ ÁN",
        "AI tại trang bán hàng và trang kho hàng",
        "Giúp sinh viên hiểu code và trình bày lại bằng lời của mình",
    )

    doc.add_heading("1. AI TRONG DỰ ÁN NÀY LÀ GÌ?", level=1)
    add_plain(doc, "Trong HomeX POS, AI không tự điều khiển cửa hàng. AI chỉ đóng vai trò người hỗ trợ: ở trang bán hàng, AI giúp chọn sản phẩm phù hợp với nhu cầu khách; ở trang kho, AI giúp đọc số liệu bán hàng và giải thích sản phẩm nào nên nhập thêm.")
    remember(doc, "Có hai nơi sử dụng AI chính: trang POS để gợi ý bán hàng và trang Kho để gợi ý nhập hàng. AI không tự thanh toán, không tự tạo đơn nhập và không tự thay đổi tồn kho.")
    add_table(doc, ["Nơi sử dụng", "AI giúp việc gì?", "Người quyết định cuối"], [
        ("Trang POS", "Gợi ý sản phẩm theo nhu cầu, giá tiền và giỏ hàng.", "Thu ngân chọn có thêm sản phẩm vào giỏ hay không."),
        ("Trang Kho", "Phân tích số lượng bán, tồn kho và đề xuất nhập hàng.", "Quản trị viên duyệt hoặc từ chối đề xuất."),
    ], [3.0, 7.0, 4.6])

    doc.add_heading("2. ĐƯỜNG ĐI CHUNG CỦA MỘT YÊU CẦU AI", level=1)
    add_flow(doc, [
        "Người dùng nhập thông tin và bấm nút",
        "Giao diện gửi yêu cầu đến máy chủ",
        "Máy chủ lấy dữ liệu thật từ cơ sở dữ liệu",
        "Hệ thống lọc hoặc tính toán trước",
        "Lần lượt thử Gemini → Groq",
        "Nếu nơi trước lỗi hoặc chờ quá thời gian quy định thì thử nơi tiếp theo",
        "Nếu tất cả đều lỗi thì dùng kết quả tính sẵn",
        "Kết quả quay lại giao diện để người dùng xem",
    ], "Sơ đồ chung, áp dụng cho cả POS và Kho", [LIGHT_BLUE, LIGHT_GREEN])
    add_plain(doc, "Hệ thống dùng Gemini làm dịch vụ AI chính. Nếu Gemini lỗi, chương trình thử Groq bằng GROQ_API_KEY. Groq dùng model openai/gpt-oss-120b cho POS và Kho, còn qwen/qwen3.6-27b dùng cho ảnh. Nếu cả hai dịch vụ không dùng được, chương trình chuyển sang cách tính cục bộ. Phần bán hàng chờ tối đa 15 giây và phần kho chờ tối đa 10 giây cho mỗi lần gọi. Dự án không dùng OPENAI_API_KEY; thư viện OpenAI chỉ là công cụ gửi yêu cầu theo định dạng tương thích.")
    speech(doc, "Trong dự án của em, Gemini là AI chính. Nếu Gemini lỗi hoặc quá thời gian chờ thì code chuyển sang Groq. Em đã kiểm tra Groq bằng dữ liệu giả và cả POS, Kho, đọc ảnh đều trả đúng dạng. Nếu cả hai dịch vụ lỗi thì POS dùng cách chấm điểm, còn Kho dùng công thức có sẵn. Dự án không dùng OPENAI_API_KEY.")

    doc.add_heading("3. AI GỢI Ý BÁN HÀNG Ở TRANG POS", level=1)
    doc.add_heading("3.1. Người dùng nhập những gì?", level=2)
    add_bullets(doc, [
        "Khách đang cần sản phẩm gì, ví dụ: nồi cơm điện cho gia đình 4 người.",
        "Khoảng tiền khách có thể chi.",
        "Các sản phẩm đang có trong giỏ hàng.",
        "Khách có ưu tiên khuyến mãi, bảo hành, hàng còn nhiều hoặc mua kèm hay không.",
    ])
    add_plain(doc, "Hàm handleRequestSalesSuggestion trong trang POS gom các thông tin trên rồi gửi đến đường dẫn /api/pos/sales-assistant.")

    doc.add_heading("3.2. Máy chủ xử lý theo 5 bước", level=2)
    add_numbered(doc, [
        "Lấy các sản phẩm đang bán và còn hàng. Sản phẩm hết hàng không được đưa vào gợi ý.",
        "Lọc theo khoảng tiền khách đã chọn.",
        "So sánh nhu cầu với tên, danh mục và mô tả sản phẩm rồi chấm điểm.",
        "Cộng thêm điểm nếu sản phẩm có khuyến mãi, bảo hành, còn nhiều hoặc phù hợp để mua kèm.",
        "Chọn tối đa 25 sản phẩm có điểm tốt để gửi cho AI.",
    ])
    remember(doc, "AI không được xem toàn bộ kho hàng. Máy chủ đã chọn trước tối đa 25 sản phẩm phù hợp, sau đó AI chỉ chọn trong danh sách này.")

    doc.add_heading("3.3. AI trả về những gì?", level=2)
    add_bullets(doc, [
        "Danh sách tối đa 5 sản phẩm nên giới thiệu cho khách.",
        "Lý do vì sao sản phẩm phù hợp.",
        "Gợi ý bộ sản phẩm có thể mua cùng nhau.",
        "Một số câu nhắc để thu ngân tư vấn khách.",
    ])
    add_plain(doc, "Sau khi nhận kết quả, máy chủ kiểm tra lại mã sản phẩm. Nếu AI trả một mã không có trong danh sách đã gửi thì kết quả đó bị bỏ. Tên, giá và tồn kho được lấy lại từ cơ sở dữ liệu thật, không lấy theo lời AI.")

    doc.add_heading("3.4. Nếu AI không chạy thì sao?", level=2)
    add_plain(doc, "Nếu các dịch vụ AI đều lỗi, chương trình dùng điểm đã tính ở bước trước để tự chọn sản phẩm. Trên giao diện, nguồn kết quả sẽ là HEURISTIC, có thể hiểu đơn giản là gợi ý bằng quy tắc chấm điểm của chương trình.")
    speech(doc, "Ở trang POS, hệ thống không đưa toàn bộ sản phẩm cho AI. Chương trình lọc hàng còn bán, kiểm tra ngân sách và chấm điểm trước. AI chỉ chọn trong tối đa 25 sản phẩm phù hợp. Khi AI trả về, máy chủ kiểm tra lại mã sản phẩm. Nếu các dịch vụ AI đều lỗi thì hệ thống vẫn gợi ý bằng cách chấm điểm có sẵn.")

    doc.add_heading("4. AI PHÂN TÍCH Ở TRANG KHO HÀNG", level=1)
    doc.add_heading("4.1. Khi bấm “Phân tích”, chương trình lấy dữ liệu gì?", level=2)
    add_bullets(doc, [
        "Số lượng đang còn của từng sản phẩm.",
        "Mức tồn kho tối thiểu do cửa hàng đặt ra.",
        "Số lượng đã bán trong 7 ngày gần nhất.",
        "Số lượng đã bán trong 30 ngày gần nhất.",
        "Danh mục sản phẩm và thời điểm trong năm để xét mùa vụ.",
        "Số ngày dự phòng do quản trị viên chọn, ví dụ 15 ngày.",
    ])
    add_plain(doc, "Chương trình chỉ tính các đơn đã hoàn tất hoặc đã thanh toán. Đơn đã hủy và khoản đã hoàn tiền không được tính vào doanh số.")

    doc.add_heading("4.2. Chương trình tính như thế nào?", level=2)
    add_plain(doc, "Trước tiên, chương trình xem gần đây sản phẩm bán nhanh hay chậm. Sau đó chương trình ước lượng trong số ngày dự phòng sẽ cần bao nhiêu sản phẩm, cộng thêm 20% dự phòng và trừ số lượng đang còn trong kho.")
    add_table(doc, ["Thông tin", "Hiểu đơn giản"], [
        ("Bán 7 ngày", "Cho biết tình hình bán gần đây."),
        ("Bán 30 ngày", "Cho biết mức bán bình thường trong một tháng."),
        ("Xu hướng", "Nếu 7 ngày gần đây bán nhanh hơn mức tháng thì sản phẩm đang tăng."),
        ("Số ngày đủ bán", "Ước lượng tồn hiện tại còn đủ bán được bao lâu."),
        ("Mùa vụ", "Một số nhóm hàng có thể bán tốt hơn vào mùa nóng hoặc dịp Tết."),
        ("Tồn an toàn", "Cộng thêm 20% để giảm nguy cơ thiếu hàng."),
    ], [4.0, 10.6])
    add_callout(doc, "Ví dụ rất đơn giản", "Một sản phẩm dự kiến cần 30 cái trong 15 ngày tới. Hệ thống cộng thêm 20% là 6 cái, tổng cần 36 cái. Nếu kho đang có 10 cái thì số lượng cần nhập dự kiến là 26 cái.", LIGHT_BLUE)

    doc.add_heading("4.3. Code mới tránh nhập dư như thế nào?", level=2)
    add_plain(doc, "Nếu tồn hiện tại đã cao hơn mức tối thiểu và đủ bán trên 60 ngày, hoặc tồn hiện tại đã đủ đáp ứng nhu cầu dự kiến cộng phần an toàn, chương trình đặt số lượng cần nhập về 0. Nếu sản phẩm không bán trong cả 7 ngày và 30 ngày thì chỉ nhập bù đến mức tối thiểu khi kho đang thiếu.")
    remember(doc, "Code không phải lúc nào cũng đề xuất nhập. Khi hàng đang còn nhiều và đủ bán lâu, kết quả sẽ là “chưa cần nhập thêm”.")

    doc.add_heading("4.4. Danh sách nào được gửi sang AI?", level=2)
    add_plain(doc, "Chương trình giữ các sản phẩm hết hàng, tồn thấp, bán tăng, có dấu hiệu mùa vụ hoặc sắp cạn. Sau khi sắp xếp mức quan trọng, tối đa 40 sản phẩm được đưa vào nội dung gửi AI. Tuy nhiên danh sách dự phòng hiển thị được giới hạn tối đa 12 sản phẩm để người quản trị dễ xem.")

    doc.add_heading("4.5. AI và kết quả dự phòng phối hợp ra sao?", level=2)
    add_numbered(doc, [
        "Chương trình luôn tính sẵn một kết quả bằng số liệu thật.",
        "Sau đó mới thử gọi Gemini; nếu Gemini lỗi thì thử Groq. Nếu cả hai đều lỗi, hệ thống dùng kết quả tính sẵn.",
        "Kết quả AI phải đúng mẫu dữ liệu và mã sản phẩm phải có trong danh sách đã gửi.",
        "Nếu AI trả thiếu sản phẩm, chương trình lấy kết quả tính sẵn để bù vào phần còn thiếu.",
        "Cuối cùng chương trình kiểm tra lại tên, hình ảnh, mùa vụ và số lượng cần nhập.",
    ])
    speech(doc, "Ở trang Kho, chương trình lấy tồn kho và số lượng bán trong 7 ngày, 30 ngày để tính trước. Hệ thống xem tốc độ bán, số ngày hàng còn đủ và mùa vụ, rồi tạo một kết quả dự phòng. Sau đó AI giúp diễn giải lý do, rủi ro và hướng xử lý. Nếu AI lỗi hoặc trả thiếu thì chương trình dùng kết quả đã tính sẵn, nên chức năng vẫn hoạt động.")

    doc.add_heading("5. SO SÁNH HAI CHỨC NĂNG AI", level=1)
    add_table(doc, ["Điểm so sánh", "Trang POS", "Trang Kho"], [
        ("Người dùng", "Thu ngân và quản trị viên", "Quản trị viên"),
        ("Thông tin đầu vào", "Nhu cầu, ngân sách, giỏ hàng", "Tồn kho, doanh số 7/30 ngày, số ngày dự phòng"),
        ("Chương trình làm trước", "Lọc và chấm điểm sản phẩm", "Tính nhu cầu và số lượng nên nhập"),
        ("AI giúp", "Chọn sản phẩm và viết lý do tư vấn", "Giải thích tình trạng kho, rủi ro và hướng xử lý"),
        ("Nếu AI lỗi", "Dùng điểm đã chấm", "Dùng kết quả tính sẵn"),
        ("Hành động cuối", "Thu ngân thêm vào giỏ", "Quản trị viên duyệt đề xuất"),
    ], [3.4, 5.6, 5.6])

    doc.add_heading("6. CÁC FILE CODE CẦN NHỚ", level=1)
    add_table(doc, ["File", "Vai trò"], [
        ("frontend/app/(dashboard)/pos/page.tsx", "Giao diện bán hàng, mở trợ lý và gửi nhu cầu."),
        ("backend/src/routes/sales-assistant.routes.ts", "Nhận yêu cầu AI POS và kiểm tra quyền."),
        ("backend/src/services/sales-assistant.service.ts", "Lọc, chấm điểm, gọi AI và tạo gợi ý POS."),
        ("frontend/app/(dashboard)/inventory/page.tsx", "Giao diện kho, chọn số ngày và xem phân tích."),
        ("backend/src/routes/inventory.routes.ts", "Nhận yêu cầu phân tích kho."),
        ("backend/src/services/inventory-ai.service.ts", "Lấy doanh số, tính dự báo, gọi AI và tạo kết quả kho."),
        ("frontend/services/homex.service.ts", "Nơi giao diện khai báo đường dẫn gọi hai chức năng AI."),
    ], [8.0, 6.6])

    doc.add_heading("7. CÂU HỎI THẦY CÓ THỂ HỎI", level=1)
    add_table(doc, ["Câu hỏi", "Cách trả lời ngắn"], [
        ("Dự án có tự huấn luyện AI không?", "Không. Dự án dùng mô hình có sẵn và gửi dữ liệu cần phân tích trong mỗi lần gọi."),
        ("Dữ liệu AI lấy từ đâu?", "Từ sản phẩm, tồn kho và đơn hàng trong cơ sở dữ liệu của dự án."),
        ("Mất mạng hoặc hết khóa AI có chạy không?", "Có. POS dùng cách chấm điểm; Kho dùng kết quả tính sẵn."),
        ("AI có tự thay đổi tồn kho không?", "Không. AI chỉ gợi ý, quản trị viên vẫn phải duyệt và tạo phiếu nhập."),
        ("Dự án dùng dịch vụ AI nào?", "Gemini là chính, Groq là dự phòng. Không dùng OPENAI_API_KEY; thư viện OpenAI chỉ giúp gửi yêu cầu theo định dạng tương thích."),
        ("Làm sao hạn chế AI bịa dữ liệu?", "Chỉ gửi danh sách đã lọc, kiểm tra lại mã sản phẩm và lấy tên/giá/tồn từ dữ liệu thật."),
        ("Phần Kho có phải hoàn toàn do AI tính không?", "Không. Công thức của chương trình tính trước; AI hỗ trợ phân tích và diễn giải."),
    ], [7.1, 7.5], font_size=9.5)
    remember(doc, "Khi trình bày, không cần đọc công thức dài. Hãy nói rõ: chương trình lấy dữ liệu thật → tính trước → AI hỗ trợ giải thích → có kết quả dự phòng → con người quyết định cuối.")

    doc.core_properties.title = "Hướng dẫn dễ hiểu cách AI hoạt động trong HomeX POS"
    doc.core_properties.subject = "Tài liệu học và tập trình bày dành cho sinh viên"
    doc.core_properties.author = "Nguyễn Đức Thịnh – 2305CT2084"
    AI_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(AI_OUTPUT)


def add_page_guide(doc, number, name, path, purpose, components, demo, talk, role="ADMIN"):
    doc.add_heading(f"{number}. {name.upper()}", level=1)
    add_table(doc, ["Đường dẫn", "Ai sử dụng", "Mục đích"], [(path, role, purpose)], [4.0, 3.2, 7.4], font_size=9.5)
    doc.add_heading("Trên trang có những gì?", level=2)
    add_bullets(doc, components)
    doc.add_heading("Khi trình diễn nên làm gì?", level=2)
    add_bullets(doc, demo)
    speech(doc, talk)


def build_website_guide():
    doc = setup_document()
    add_simple_cover(
        doc,
        "HƯỚNG DẪN THUYẾT TRÌNH\nGIỚI THIỆU TOÀN BỘ WEBSITE HOMEX POS",
        "Từng trang có gì, dùng để làm gì và nên nói như thế nào",
        "Giúp sinh viên trình diễn website theo một thứ tự rõ ràng",
    )

    doc.add_heading("CÁCH DÙNG FILE NÀY", level=1)
    add_plain(doc, "Mỗi phần bên dưới tương ứng với một trang của website. Bạn không cần học thuộc từng câu. Hãy đọc phần “Bạn có thể nói với thầy”, sau đó dùng lời của mình để trình bày. Khi demo, chỉ nên thực hiện một thao tác tiêu biểu trên mỗi nhóm chức năng để tránh mất thời gian.")
    add_callout(doc, "Thứ tự nên trình bày", "Đăng nhập → Tổng quan → POS → Hóa đơn → Khách hàng/Bảo hành → Kho/Sản phẩm → Báo cáo → Nhân viên/Ca làm → Cài đặt và nhật ký. Các trang công khai giới thiệu ở cuối.", LIGHT_BLUE)

    doc.add_heading("LỜI MỞ ĐẦU KHOẢNG 45 GIÂY", level=1)
    speech(doc, "Em xin giới thiệu HomeX POS, là hệ thống hỗ trợ quản lý bán hàng cho cửa hàng đồ gia dụng. Hệ thống có hai vai trò chính là quản trị viên và thu ngân. Thu ngân tập trung vào bán hàng, hóa đơn, khách hàng, bảo hành và ca làm. Quản trị viên có thêm các chức năng quản lý kho, sản phẩm, nhà cung cấp, nhân viên, báo cáo và cài đặt. Ngoài các nghiệp vụ cơ bản, dự án có AI hỗ trợ gợi ý sản phẩm tại quầy và phân tích nhập hàng trong kho.", "Lời mở đầu mẫu")

    doc.add_heading("PHÂN QUYỀN VÀ THANH ĐIỀU HƯỚNG", level=1)
    add_table(doc, ["Vai trò", "Các nhóm trang được thấy"], [
        ("Thu ngân", "Tổng quan, POS, hóa đơn, khách hàng, bảo hành và ca làm."),
        ("Quản trị viên", "Toàn bộ chức năng của thu ngân và thêm kho, sản phẩm, danh mục, nhà cung cấp, khuyến mãi, hóa đơn VAT, nhân viên, báo cáo, cài đặt, nhật ký."),
    ], [4.0, 10.6])
    add_plain(doc, "Bên trái là thanh menu. Phía trên có thông tin người dùng, thông báo, đổi ngôn ngữ và đăng xuất. Menu có thể thu gọn và trên điện thoại sẽ chuyển thành menu mở/đóng.")

    add_page_guide(doc, "1", "Trang đăng nhập", "/login", "Xác thực tài khoản và chuyển người dùng đến đúng khu vực theo vai trò.", [
        "Ô nhập mã nhân viên hoặc email.",
        "Ô mật khẩu và nút đăng nhập.",
        "Thông báo sai tài khoản, hết phiên đăng nhập hoặc đang xử lý.",
        "Nút đổi ngôn ngữ và thông tin tài khoản dùng thử.",
    ], [
        "Đăng nhập bằng tài khoản quản trị viên để có thể giới thiệu đầy đủ chức năng.",
        "Nếu còn thời gian, nói thêm thu ngân sẽ nhìn thấy menu ít hơn.",
    ], "Đây là trang đăng nhập. Hệ thống kiểm tra tài khoản và mật khẩu, sau đó lưu thông tin đăng nhập. Tùy vai trò là quản trị viên hay thu ngân, người dùng được chuyển vào khu vực phù hợp và chỉ nhìn thấy các chức năng được phép sử dụng.", "Tất cả")

    add_page_guide(doc, "2", "Trang tổng quan", "/dashboard", "Cho người dùng nhìn nhanh tình hình hoạt động của cửa hàng.", [
        "Bốn ô số liệu: doanh thu hôm nay, số đơn, số sản phẩm đã bán và cảnh báo tồn thấp.",
        "Biểu đồ doanh thu 7 ngày hoặc 30 ngày.",
        "Theo dõi doanh số theo danh mục sản phẩm.",
        "Danh sách sản phẩm sắp hết và sản phẩm bán chạy.",
        "Với thu ngân, trang được rút gọn thành ba lối tắt: POS, hóa đơn và bảo hành.",
    ], [
        "Chỉ vào bốn ô số liệu để giới thiệu tình hình trong ngày.",
        "Đổi biểu đồ từ 7 ngày sang 30 ngày.",
        "Chỉ vào danh sách sắp hết hàng để dẫn sang chức năng Kho.",
    ], "Đây là trang tổng quan dành cho quản trị viên. Bốn ô phía trên giúp xem nhanh doanh thu, số đơn, lượng sản phẩm đã bán và cảnh báo tồn kho. Bên dưới là biểu đồ doanh thu, doanh số theo danh mục, hàng sắp hết và sản phẩm bán chạy. Nhờ đó quản trị viên không cần mở từng trang vẫn nắm được tình hình chung.")

    add_page_guide(doc, "3", "Trang bán hàng tại quầy — POS", "/pos", "Thực hiện toàn bộ quy trình bán hàng và thanh toán.", [
        "Danh sách sản phẩm, tìm kiếm, lọc theo danh mục và xem tồn kho.",
        "Giỏ hàng: tăng giảm số lượng, xóa hàng và tính tổng tiền.",
        "Chọn hoặc tạo nhanh khách hàng.",
        "Áp mã khuyến mãi hoặc giảm giá thủ công theo quyền cài đặt.",
        "Quét mã vạch bằng máy quét hoặc điện thoại.",
        "Lưu đơn nháp, tiếp tục đơn nháp và hủy đơn.",
        "Thanh toán tiền mặt hoặc chuyển khoản PayOS.",
        "Trợ lý AI gợi ý sản phẩm theo nhu cầu và ngân sách.",
    ], [
        "Chọn một sản phẩm và thêm vào giỏ.",
        "Chọn khách hàng hoặc nhập mã khuyến mãi.",
        "Mở trợ lý AI, nhập nhu cầu ngắn và cho thấy kết quả gợi ý.",
        "Không cần thanh toán thật nếu thời gian ngắn; có thể giải thích hai phương thức.",
    ], "Đây là chức năng trung tâm của hệ thống. Thu ngân tìm hoặc quét sản phẩm, thêm vào giỏ, chọn khách hàng, áp khuyến mãi và thanh toán. Trang hỗ trợ tiền mặt, chuyển khoản PayOS, lưu đơn nháp và quét mã bằng điện thoại. Điểm nổi bật là trợ lý AI giúp thu ngân gợi ý sản phẩm phù hợp với nhu cầu và khoảng tiền của khách.", "ADMIN và CASHIER")

    add_page_guide(doc, "4", "Trang hóa đơn và giao dịch", "/orders", "Quản lý đơn hàng, thanh toán và trả hàng trên một trang.", [
        "Tab Tất cả: danh sách toàn bộ đơn hàng.",
        "Tab Đơn nháp: các đơn chưa thanh toán để tiếp tục bán.",
        "Tab Thanh toán: lịch sử thanh toán và chức năng hoàn tiền cho quản trị viên.",
        "Tab Trả hàng: tìm đơn đã hoàn tất và lập phiếu trả.",
        "Bộ lọc mã đơn, trạng thái, nhân viên và thời gian; hỗ trợ xuất CSV.",
        "Chi tiết đơn gồm sản phẩm, phương thức thanh toán và bảo hành.",
        "Có thể in hóa đơn, yêu cầu VAT, hủy đơn hoặc tạo trả hàng tùy trạng thái/quyền.",
    ], [
        "Tìm một mã đơn và mở chi tiết.",
        "Chỉ vào các tab để giải thích nghiệp vụ được gom chung.",
        "Mở phần sản phẩm và phương thức thanh toán trong chi tiết đơn.",
    ], "Trang này gom các nghiệp vụ sau bán hàng. Người dùng có thể xem đơn hoàn tất, tiếp tục đơn nháp, xem lịch sử thanh toán và lập phiếu trả hàng. Khi mở chi tiết, hệ thống hiển thị sản phẩm, tổng tiền, phương thức thanh toán và bảo hành liên quan. Quản trị viên có thêm quyền hủy đơn hoặc hoàn tiền.", "ADMIN và CASHIER")

    add_page_guide(doc, "5", "Trang khách hàng", "/customers", "Lưu thông tin khách mua hàng và theo dõi khách thân thiết.", [
        "Danh sách họ tên, số điện thoại, email, điểm và hạng khách hàng.",
        "Tìm kiếm và lọc trạng thái.",
        "Thêm mới, cập nhật, ngừng hoạt động hoặc khôi phục khách hàng.",
        "Xuất danh sách khách hàng ra CSV.",
    ], [
        "Tìm một khách bằng số điện thoại.",
        "Mở biểu mẫu thêm khách để giới thiệu các trường thông tin.",
    ], "Trang khách hàng giúp cửa hàng lưu thông tin liên hệ, điểm tích lũy và hạng thành viên. Dữ liệu này được sử dụng khi bán hàng, áp khuyến mãi và tra cứu bảo hành. Thu ngân có thể thêm hoặc cập nhật khách; các thao tác xóa hoặc khôi phục được kiểm soát theo quyền.", "ADMIN và CASHIER")

    add_page_guide(doc, "6", "Trang bảo hành", "/warranties", "Theo dõi bảo hành điện tử của sản phẩm đã bán.", [
        "Danh sách mã bảo hành, khách hàng, sản phẩm, ngày bắt đầu, ngày hết hạn và trạng thái.",
        "Tìm theo mã, khách hàng hoặc đơn hàng.",
        "Tạo bảo hành thủ công từ chi tiết đơn hàng.",
        "Xem chi tiết, gửi email, sao chép đường dẫn tra cứu.",
        "Hủy, khôi phục hoặc chuyển sang hết hạn theo quyền.",
    ], [
        "Tìm một mã bảo hành và mở chi tiết.",
        "Sao chép đường dẫn tra cứu để dẫn sang trang công khai.",
    ], "Trang bảo hành quản lý toàn bộ thời hạn bảo hành của sản phẩm đã bán. Mỗi bảo hành liên kết với sản phẩm trong một đơn hàng. Nhân viên có thể tra cứu, xem chi tiết, gửi email hoặc cung cấp đường dẫn để khách tự kiểm tra. Trạng thái bảo hành cũng được quản lý rõ ràng.", "ADMIN và CASHIER")

    add_page_guide(doc, "7", "Trang khuyến mãi", "/promotions", "Tạo và quản lý mã giảm giá cho hoạt động bán hàng.", [
        "Mã và tên chương trình khuyến mãi.",
        "Giảm theo số tiền hoặc phần trăm.",
        "Mức đơn tối thiểu, mức giảm tối đa và thời hạn.",
        "Giới hạn tổng lượt dùng, lượt dùng theo khách và hạng khách được áp dụng.",
        "Tìm kiếm, lọc trạng thái, thêm, sửa và ngừng chương trình.",
    ], [
        "Mở một khuyến mãi để giải thích điều kiện áp dụng.",
        "Liên hệ lại với trang POS: mã này được nhập khi thanh toán.",
    ], "Trang khuyến mãi cho phép quản trị viên tạo các mã giảm giá có điều kiện rõ ràng như thời gian, giá trị đơn tối thiểu, giới hạn lượt dùng và hạng khách hàng. Khi thu ngân nhập mã ở POS, máy chủ sẽ kiểm tra các điều kiện này trước khi giảm tiền.")

    add_page_guide(doc, "8", "Trang yêu cầu hóa đơn VAT", "/vat-invoices", "Tiếp nhận và xử lý yêu cầu hóa đơn VAT từ khách hàng.", [
        "Danh sách yêu cầu theo trạng thái chờ, đã duyệt hoặc từ chối.",
        "Thông tin công ty, mã số thuế, địa chỉ và email người mua.",
        "Duyệt, từ chối, nhập mã hóa đơn đỏ và ghi chú quản trị.",
        "Gửi lại email hoặc xóa yêu cầu theo quyền.",
    ], [
        "Mở một yêu cầu đang chờ.",
        "Chỉ vào nút duyệt/từ chối và giải thích đây là quy trình mô phỏng.",
    ], "Trang này xử lý các yêu cầu xuất hóa đơn VAT do khách gửi từ đường dẫn hóa đơn công khai. Quản trị viên kiểm tra thông tin doanh nghiệp, sau đó duyệt hoặc từ chối và có thể gửi lại kết quả qua email.")

    add_page_guide(doc, "9", "Trang kho hàng", "/inventory", "Theo dõi tồn kho, nhập hàng, điều chỉnh kho và phân tích bằng AI.", [
        "Danh sách sản phẩm sắp hết hàng.",
        "Lịch sử nhập, bán, điều chỉnh và khôi phục tồn kho.",
        "Nhập nhanh một sản phẩm hoặc điều chỉnh số lượng thực tế.",
        "Lập phiếu nhập gồm nhà cung cấp và nhiều dòng sản phẩm.",
        "Khu phân tích AI: chọn số ngày dự phòng, xem thống kê, đề xuất, lý do, rủi ro và kế hoạch.",
        "Duyệt hoặc từ chối đề xuất rồi chuyển các mục đã duyệt sang phiếu nhập.",
    ], [
        "Mở khu phân tích AI và chọn 15 ngày.",
        "Chỉ vào một sản phẩm để giải thích số đã bán, tồn hiện tại và số nên nhập.",
        "Duyệt một đề xuất và cho thấy cách chuyển sang phiếu nhập.",
    ], "Trang Kho vừa quản lý số lượng thực tế vừa hỗ trợ quyết định nhập hàng. Quản trị viên có thể nhập nhanh, điều chỉnh kho, xem lịch sử và lập phiếu nhập. Phần AI lấy doanh số 7 ngày, 30 ngày, tồn hiện tại và mùa vụ để gợi ý. AI chỉ hỗ trợ; quản trị viên vẫn là người duyệt trước khi tạo phiếu nhập.")

    add_page_guide(doc, "10", "Trang sản phẩm", "/products", "Quản lý thông tin hàng hóa đang bán trong cửa hàng.", [
        "Danh sách mã SKU, tên, danh mục, nhà cung cấp, giá, tồn kho, mức tồn tối thiểu và trạng thái.",
        "Tìm kiếm và lọc theo danh mục, nhà cung cấp, trạng thái hoặc tồn thấp.",
        "Thêm, sửa, ngừng hoạt động, khôi phục hoặc xóa theo quyền.",
        "Nhập nhiều sản phẩm từ dữ liệu mẫu/tệp và thao tác hàng loạt.",
        "Quét mã vạch bằng máy hoặc điện thoại.",
        "Tra cứu mã vạch từ nguồn ngoài và dùng AI bù thông tin còn thiếu.",
    ], [
        "Mở form thêm sản phẩm và giới thiệu các trường.",
        "Nhập/quét mã vạch để cho thấy hệ thống tự điền thông tin nếu tìm được.",
    ], "Trang Sản phẩm là nơi quản lý dữ liệu gốc của hàng hóa. Ngoài các thao tác thêm, sửa, lọc và quản lý tồn tối thiểu, trang còn hỗ trợ quét mã vạch. Hệ thống ưu tiên dữ liệu trong cửa hàng và các nguồn mã vạch, sau đó AI chỉ bù những thông tin còn thiếu để quản trị viên kiểm tra trước khi lưu.")

    add_page_guide(doc, "11", "Trang danh mục", "/categories", "Chia sản phẩm thành các nhóm để dễ tìm và báo cáo.", [
        "Danh sách mã, tên và trạng thái danh mục.",
        "Tìm kiếm, lọc trạng thái, thêm và cập nhật danh mục.",
        "Ngừng hoạt động hoặc khôi phục danh mục.",
    ], [
        "Mở form tạo danh mục và giải thích sản phẩm sẽ liên kết với danh mục này.",
    ], "Trang Danh mục dùng để chia sản phẩm thành từng nhóm như thiết bị nhà bếp, làm sạch hoặc đồ gia dụng. Việc phân nhóm giúp tìm kiếm, lọc sản phẩm và tổng hợp báo cáo dễ hơn.")

    add_page_guide(doc, "12", "Trang nhà cung cấp", "/suppliers", "Quản lý nơi cung cấp hàng và các lần nhập hàng.", [
        "Tab danh sách nhà cung cấp: tên, điện thoại, email, địa chỉ và trạng thái.",
        "Tab phiếu nhập: xem và tạo các phiếu nhập hàng.",
        "Tab lịch sử nhập: xem biến động kho loại nhập hàng.",
        "Tìm kiếm, thêm, sửa, ngừng hoặc khôi phục nhà cung cấp.",
    ], [
        "Mở ba tab để cho thấy dữ liệu nhà cung cấp liên kết với phiếu nhập và lịch sử kho.",
    ], "Trang Nhà cung cấp không chỉ lưu thông tin liên hệ mà còn liên kết với quá trình nhập hàng. Ba tab giúp quản trị viên xem danh sách nhà cung cấp, tạo hoặc xem phiếu nhập và kiểm tra lịch sử hàng đã nhập.")

    add_page_guide(doc, "13", "Trang ca làm việc", "/shifts", "Theo dõi tiền đầu ca, cuối ca và chênh lệch của thu ngân.", [
        "Ca hiện tại của người dùng.",
        "Mở ca với tiền đầu ca và loại ca.",
        "Đóng ca với số tiền thực tế cuối ca.",
        "Tính tiền dự kiến và chênh lệch.",
        "Quản trị viên xem lịch sử ca của nhiều nhân viên; thu ngân chủ yếu xem ca của mình.",
    ], [
        "Chỉ vào trạng thái ca hiện tại.",
        "Giải thích thu ngân phải mở ca trước khi thanh toán.",
    ], "Trang Ca làm việc giúp kiểm soát tiền mặt theo từng nhân viên. Thu ngân mở ca với số tiền ban đầu, thực hiện bán hàng và nhập số tiền thực tế khi đóng ca. Hệ thống tính số tiền dự kiến và chênh lệch để quản trị viên kiểm tra.", "ADMIN và CASHIER")

    add_page_guide(doc, "14", "Trang nhân viên", "/users", "Quản lý tài khoản và quyền sử dụng hệ thống.", [
        "Danh sách mã nhân viên, họ tên, email, số điện thoại, vai trò và trạng thái.",
        "Thêm tài khoản, cập nhật thông tin và đổi mật khẩu.",
        "Khóa, xóa hoặc khôi phục tài khoản.",
        "Các thao tác nhạy cảm yêu cầu xác nhận mật khẩu quản trị viên.",
    ], [
        "Mở form thêm nhân viên và chỉ vào lựa chọn ADMIN/CASHIER.",
        "Giải thích vì sao thao tác nhạy cảm cần xác nhận mật khẩu.",
    ], "Trang Nhân viên quản lý tài khoản đăng nhập và vai trò. Hai vai trò chính là quản trị viên và thu ngân. Các thao tác như đổi mật khẩu, khóa hoặc xóa tài khoản yêu cầu xác nhận của quản trị viên để tăng an toàn.")

    add_page_guide(doc, "15", "Trang báo cáo", "/reports", "Theo dõi doanh thu, chi phí, lợi nhuận và sản phẩm bán chạy theo thời gian.", [
        "Bộ lọc khoảng ngày.",
        "Các chỉ số doanh thu, giá vốn, lợi nhuận ròng và số đơn hoàn tất.",
        "Biểu đồ xu hướng tài chính theo ngày.",
        "Danh sách sản phẩm bán chạy cùng số lượng và doanh thu.",
    ], [
        "Chọn khoảng ngày có dữ liệu đơn ảo từ 05/08 đến 11/08.",
        "Giải thích sự khác nhau giữa doanh thu, giá vốn và lợi nhuận.",
    ], "Trang Báo cáo giúp quản trị viên đánh giá kết quả kinh doanh theo khoảng thời gian. Phía trên là doanh thu, giá vốn, lợi nhuận và số đơn. Bên dưới là biểu đồ thay đổi theo ngày và danh sách sản phẩm bán chạy. Dữ liệu được tổng hợp từ các đơn đã hoàn tất.")

    add_page_guide(doc, "16", "Trang cài đặt", "/settings", "Điều chỉnh cách hệ thống vận hành mà không cần sửa code.", [
        "Thông tin cửa hàng và thời gian kinh doanh.",
        "Thiết lập POS: giảm giá, xác nhận thanh toán, quét mã và cảnh báo tồn.",
        "Thiết lập in hóa đơn: khổ giấy, số bản và tự mở cửa sổ in.",
        "Thiết lập kho: tồn tối thiểu và cho phép bán vượt tồn.",
        "Cấu hình chuyển khoản/VietQR và phương thức thanh toán mặc định.",
        "Thiết lập ca làm, email VAT, dữ liệu cấu hình và bảng phân quyền.",
    ], [
        "Chỉ giới thiệu từng nhóm, không nên thay đổi nhiều giá trị khi đang demo.",
        "Mở bảng phân quyền để liên hệ với menu ADMIN/CASHIER.",
    ], "Trang Cài đặt tập trung các quy tắc vận hành của cửa hàng như in hóa đơn, thanh toán, cảnh báo tồn kho, quét mã, ca làm và email VAT. Nhờ đưa vào cài đặt, quản trị viên có thể thay đổi cách hệ thống hoạt động mà không cần sửa mã nguồn.")

    add_page_guide(doc, "17", "Trang nhật ký hoạt động", "/audit-logs", "Ghi lại các thao tác quan trọng để kiểm tra và truy vết.", [
        "Danh sách người thực hiện, hành động, loại dữ liệu, mã đối tượng, mô tả và thời gian.",
        "Bộ lọc theo nội dung, hành động, loại dữ liệu và thời gian.",
        "Hộp xem chi tiết một bản ghi nhật ký.",
    ], [
        "Tìm hành động tạo đơn hoặc thanh toán và mở chi tiết.",
    ], "Trang Nhật ký giúp quản trị viên biết ai đã thực hiện thao tác gì và vào thời điểm nào. Đây là phần hỗ trợ kiểm tra khi có sai sót hoặc cần truy vết thay đổi trong hệ thống.")

    doc.add_page_break()
    doc.add_heading("CÁC TRANG CÔNG KHAI VÀ TRANG HỖ TRỢ", level=1)
    add_plain(doc, "Các trang sau không nằm trực tiếp trong menu quản trị nhưng là một phần của website và hỗ trợ khách hàng hoặc quy trình thanh toán.")

    add_page_guide(doc, "18", "Trang hóa đơn công khai", "/invoice/[orderCode]", "Cho khách xem/in hóa đơn và gửi yêu cầu VAT.", [
        "Thông tin mã đơn, ngày tạo, thu ngân, tổng tiền và chi tiết sản phẩm.",
        "Nút in hóa đơn.",
        "Biểu mẫu tên công ty, mã số thuế, địa chỉ, email và ghi chú để yêu cầu VAT.",
        "Trạng thái xử lý yêu cầu VAT nếu đã gửi.",
    ], ["Mở bằng một mã đơn có thật và chỉ vào phần in/yêu cầu VAT."], "Đây là đường dẫn công khai gửi cho khách sau khi mua hàng. Khách có thể xem hoặc in hóa đơn mà không cần đăng nhập, đồng thời gửi yêu cầu xuất hóa đơn VAT. Yêu cầu đó sẽ xuất hiện tại trang quản lý VAT của quản trị viên.", "Khách hàng")

    add_page_guide(doc, "19", "Trang tra cứu bảo hành công khai", "/tra-cuu-bao-hanh", "Cho khách tự kiểm tra tình trạng bảo hành.", [
        "Hai cách tìm: số điện thoại hoặc mã bảo hành/mã hóa đơn.",
        "Kết quả gồm mã bảo hành, sản phẩm, tên khách đã che bớt, ngày mua, ngày hết hạn và trạng thái.",
        "Có thể tự tải kết quả nếu mã được truyền trong đường dẫn.",
    ], ["Nhập một mã bảo hành hoặc số điện thoại mẫu và mở kết quả."], "Trang này dành cho khách hàng nên không cần đăng nhập. Khách có thể tra cứu bằng số điện thoại hoặc mã. Hệ thống chỉ hiển thị thông tin cần thiết và che bớt tên khách để bảo vệ dữ liệu cá nhân.", "Khách hàng")

    add_page_guide(doc, "20", "Trang quét mã bằng điện thoại", "/mobile-scan", "Dùng camera điện thoại làm máy quét mã cho POS hoặc trang Sản phẩm.", [
        "Kết nối bằng mã phiên quét từ máy tính.",
        "Khung camera đọc mã vạch.",
        "Bật đèn, đổi camera, quét lại hoặc nhập mã thủ công.",
        "Gửi mã vừa quét về máy tính theo đúng phiên.",
    ], ["Từ POS mở mã QR phiên quét, dùng điện thoại truy cập và gửi một mã thử."], "Chức năng này biến điện thoại thành máy quét mã vạch từ xa. Máy tính tạo một phiên, điện thoại kết nối vào phiên đó, dùng camera đọc mã và gửi mã về đúng màn hình POS hoặc Sản phẩm.", "Nhân viên dùng điện thoại")

    add_page_guide(doc, "21", "Trang kết quả thanh toán PayOS", "/payment/payos/return và /cancel", "Thông báo thanh toán đã xử lý hoặc đã bị hủy.", [
        "Trang thành công nhắc người dùng quay lại POS.",
        "Trang hủy nhắc chọn lại phương thức thanh toán.",
        "POS vẫn kiểm tra trạng thái thanh toán từ máy chủ để hoàn tất đơn.",
    ], ["Chỉ cần giới thiệu, không bắt buộc thực hiện thanh toán thật khi thuyết trình."], "Hai trang này là điểm quay về sau khi khách thao tác trên PayOS. Một trang thông báo thanh toán đã được xử lý, trang còn lại thông báo người dùng đã hủy. Sau đó nhân viên quay lại POS để hệ thống kiểm tra trạng thái và tiếp tục đơn hàng.", "Khách/nhân viên")

    add_page_guide(doc, "22", "Trang không có quyền", "/unauthorized", "Thông báo khi tài khoản truy cập chức năng không được phép.", [
        "Thông báo không có quyền truy cập.",
        "Ngăn thu ngân mở các trang chỉ dành cho quản trị viên dù nhập đường dẫn trực tiếp.",
    ], ["Có thể giải thích bằng phân quyền, không cần cố tình tạo lỗi khi demo."], "Trang này xuất hiện khi người dùng cố vào chức năng không thuộc quyền của mình. Việc kiểm tra được thực hiện cả ở giao diện và máy chủ, nên không chỉ dựa vào việc ẩn menu.", "Tất cả")

    doc.add_heading("CÁC ĐƯỜNG DẪN CHỈ CHUYỂN TRANG", level=1)
    add_table(doc, ["Đường dẫn", "Chuyển đến đâu?", "Lý do"], [
        ("/payments", "/orders", "Lịch sử thanh toán đã nằm trong tab Thanh toán của trang Hóa đơn."),
        ("/return-orders", "/orders", "Trả hàng đã nằm trong tab Trả hàng của trang Hóa đơn."),
        ("/purchase-orders", "/inventory?tab=import", "Phiếu nhập được tích hợp vào trang Kho."),
        ("/", "/dashboard hoặc trang mặc định", "Điểm vào chính sau khi đã đăng nhập."),
    ], [4.0, 5.0, 5.6])

    doc.add_heading("KỊCH BẢN DEMO GỢI Ý TRONG 12–15 PHÚT", level=1)
    add_table(doc, ["Thời gian", "Nội dung"], [
        ("1 phút", "Mở đầu, đăng nhập và giải thích hai vai trò."),
        ("1 phút", "Trang Tổng quan: số liệu, biểu đồ, cảnh báo tồn."),
        ("4 phút", "POS: thêm hàng, khách hàng, khuyến mãi, AI và phương thức thanh toán."),
        ("2 phút", "Hóa đơn: chi tiết, thanh toán, trả hàng, bảo hành."),
        ("3 phút", "Kho và Sản phẩm: phân tích AI, duyệt nhập, mã vạch."),
        ("2 phút", "Báo cáo và quản lý nhân viên/ca làm."),
        ("1 phút", "Cài đặt, nhật ký và trang công khai."),
        ("1 phút", "Kết luận và mời thầy đặt câu hỏi."),
    ], [3.0, 11.6])

    doc.add_heading("LỜI KẾT KHOẢNG 30 GIÂY", level=1)
    speech(doc, "Qua phần trình bày, HomeX POS đã bao phủ quy trình từ bán hàng, thanh toán, hóa đơn, khách hàng, bảo hành đến quản lý kho, nhập hàng, nhân viên và báo cáo. Hệ thống có phân quyền rõ ràng và có các trang công khai cho khách tra cứu. Hai chức năng AI được đặt ở đúng nơi cần hỗ trợ quyết định là gợi ý bán hàng và phân tích nhập kho, nhưng người dùng vẫn là người quyết định cuối cùng. Em xin kết thúc phần giới thiệu và sẵn sàng trả lời câu hỏi của thầy.", "Lời kết mẫu")

    doc.core_properties.title = "Hướng dẫn thuyết trình giới thiệu toàn bộ website HomeX POS"
    doc.core_properties.subject = "Từng trang, thành phần và lời nói mẫu dành cho sinh viên"
    doc.core_properties.author = "Nguyễn Đức Thịnh – 2305CT2084"
    WEBSITE_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(WEBSITE_OUTPUT)


if __name__ == "__main__":
    build_ai_guide()
    build_website_guide()
    print("Created student-friendly AI guide and full website presentation guide")
